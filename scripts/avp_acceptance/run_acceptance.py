"""Acceptation RÉELLE du pack AVP sur une vraie maquette BIMData.

Génère le pack AVP depuis un modèle BIMData réel (config via l'environnement),
puis vérifie que **les 5 annexes** sont non vides et **habillées de la charte
BIMData**. **Read-only côté BIMData** (extraction seule ; aucune écriture, aucune
publication).

⚠️ Le pack généré (xlsx/docx) contient des **données client** : il ne doit JAMAIS
être versionné. Le script **refuse** d'écrire dans le dépôt — passe un dossier de
sortie **hors du repo** (ex. ``/tmp/avp-acceptance``). La sortie stdout ne contient
QUE des compteurs, des booléens et un verdict (aucune donnée client).

Usage: python run_acceptance.py <out_dir_hors_repo> [phase]
"""

from __future__ import annotations

import json
import re
import sys
import unicodedata
import zipfile
from pathlib import Path

# Nom de projet « bouchon » injecté pour permettre la génération quand BIMData
# ne fournit aucun nom : il ne doit JAMAIS satisfaire le contrôle de métadonnées.
PLACEHOLDER_PROJECT_NAME = "ACCEPTANCE"

# Seuils du rapport Word (partagés runner ↔ tests via inspect_word_report).
WORD_MIN_PARAGRAPHS = 10
WORD_MIN_SIGNIFICANT_CELLS = 10

# Intitulés MÉTIER attendus des 9 sections : le contrôle vérifie le numéro **ET**
# l'intitulé (un « 1. Foo … 9. Bar » avec les bons numéros mais de faux titres est
# refusé). On matche par mot-clé (sous-chaîne, insensible à la casse/accents).
WORD_SECTION_TITLES = {
    1: "Données d'entrée",
    2: "Usages BIM 3F",
    3: "Synthèse",
    4: "Indicateurs",
    5: "Écarts",
    6: "Grille",
    7: "Points bloquants",
    8: "Recommandations",
    9: "Annexes",
}
_WORD_SECTION_RE = re.compile(r"^\s*([1-9])\.\s*(.+?)\s*$")


def _norm_title(s: str) -> str:
    """Normalise un intitulé pour comparaison : casse ET accents ignorés
    (« Écarts » matche « ecarts », « Donnees d'entree » matche la table)."""
    return "".join(
        c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c)
    ).casefold()


def _assert_outside_repo(out: Path) -> None:
    root = Path(__file__).resolve()
    while root != root.parent and not (root / ".git").exists():
        root = root.parent
    out = out.resolve()
    if out == root or root in out.parents:
        raise SystemExit(
            f"REFUS : {out} est dans le dépôt {root}. Le pack contient des données "
            f"client — écris-le HORS du repo (ex. /tmp/avp-acceptance)."
        )


def assert_catalog_usable(docs: dict[str, str | None], catalog) -> None:
    """Refuse un référentiel CCH inexploitable — **helper pur, testable**.

    ``build_catalog`` tolère des documents absents et rend un catalogue
    partiel/vide ; sans contrôle, l'acceptation pourrait rendre ``PASS`` sans
    aucun référentiel CCH réellement chargé. On refuse (``SystemExit``) si :

    - un des documents I3F (``docs`` = nom → chemin) est absent/introuvable ;
    - ``catalog.properties`` ou ``catalog.naming_rules`` est vide.
    """
    missing = [name for name, p in docs.items() if not p or not Path(p).exists()]
    if missing:
        raise SystemExit(f"REFUS : documents I3F absents {missing} — contrôle CCH impossible.")
    n_props = len(getattr(catalog, "properties", None) or [])
    n_rules = len(getattr(catalog, "naming_rules", None) or [])
    if n_props == 0 or n_rules == 0:
        raise SystemExit(
            f"REFUS : catalogue CCH vide (properties={n_props}, naming_rules={n_rules}) "
            f"— acceptation non fiable."
        )


def _charte_flags(path: Path, wordmark: str, primary: str, font: str) -> dict:
    with zipfile.ZipFile(path) as z:
        blob = b"".join(z.read(n) for n in z.namelist() if n.endswith((".xml", ".rels"))).upper()
    return {
        "wordmark": wordmark.encode().upper() in blob,
        "primary": primary.encode().upper() in blob,
        "font": font.encode().upper() in blob,
        "no_korhus": b"KORHUS" not in blob,
    }


def inspect_word_report(
    docx_path: Path,
    *,
    expected_project_name: str | None,
    phase: str | None,
    wordmark: str,
    primary: str,
    font: str,
    not_available: str,
    min_paragraphs: int = WORD_MIN_PARAGRAPHS,
    min_significant_cells: int = WORD_MIN_SIGNIFICANT_CELLS,
) -> dict:
    """Inspecte un rapport Word AVP — **helper pur, partagé runner ↔ tests**.

    Renvoie des constats (compteurs + booléens) et un ``ok`` global. Critères :

    - **contenu non vide** : ``>= min_paragraphs`` paragraphes ET
      ``>= min_significant_cells`` cellules **significatives** (non vides et
      distinctes de ``not_available``) ;
    - **sections 1 à 9** présentes (paragraphes commençant par ``"1."``…``"9."``) ;
    - **métadonnées** : ``expected_project_name`` (le **vrai** nom de projet
      BIMData, jamais le bouchon) présent dans le document, et ``phase`` présente ;
    - **charte** BIMData (wordmark / primaire / police) sans KORHUS.
    """
    from docx import Document

    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    significant_cells = 0
    cell_texts: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_texts.append(cell.text)
                txt = cell.text.strip()
                if txt and txt != not_available:
                    significant_cells += 1

    doc_txt = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(cell_texts)

    # Sections : on capture numéro + intitulé, et on vérifie que CHAQUE numéro
    # 1..9 porte le BON intitulé métier (pas seulement « N. quelque chose »).
    found_titles: dict[int, str] = {}
    for text in paragraphs:
        m = _WORD_SECTION_RE.match(text)
        if m:
            found_titles.setdefault(int(m.group(1)), m.group(2))
    sections_present = sorted(found_titles)
    sections_ok = all(
        n in found_titles and _norm_title(WORD_SECTION_TITLES[n]) in _norm_title(found_titles[n])
        for n in range(1, 10)
    )

    # Le bouchon ne satisfait JAMAIS le contrôle : on exige le vrai nom BIMData.
    real_name = expected_project_name or ""
    metadata_present = (
        bool(real_name)
        and real_name != PLACEHOLDER_PROJECT_NAME
        and real_name in doc_txt
        and bool(phase)
        and phase in doc_txt
    )

    charte = _charte_flags(docx_path, wordmark, primary, font)
    non_empty = len(paragraphs) >= min_paragraphs and significant_cells >= min_significant_cells
    ok = non_empty and sections_ok and metadata_present and all(charte.values())
    return {
        "n_paragraphs": len(paragraphs),
        "n_significant_cells": significant_cells,
        "sections_present": sections_present,
        "sections_ok": sections_ok,
        "non_empty": non_empty,
        "metadata_present": metadata_present,
        **charte,
        "ok": ok,
    }


def main(argv: list[str]) -> int:
    if len(argv) not in (2, 3):
        print("usage: python run_acceptance.py <out_dir_hors_repo> [phase]", file=sys.stderr)
        return 2

    from audit_bim import config
    from audit_bim.mcp.security import set_runtime_transport

    set_runtime_transport("script")  # PR3 §3a — entrypoint local (writes/token permis)
    from audit_bim.audit.engine import run_audit
    from audit_bim.extraction.client import BIMDataClient
    from audit_bim.extraction.model_data import extract_snapshot
    from audit_bim.reporting.avp_i3f import (
        _count_business_rows,
        _count_controle_rows,
        write_avp_i3f_report_pack,
    )
    from audit_bim.reporting.bimdata_brand import WORDMARK
    from audit_bim.reporting.theming import BIMDATA_FONT_PRIMARY, BIMDATA_PRIMARY
    from audit_bim.reporting.word_report import NOT_AVAILABLE
    from audit_bim.requirements.catalog import build_catalog
    from audit_bim.requirements.models import BIMPhase

    out = Path(argv[1])
    _assert_outside_repo(out)
    out.mkdir(parents=True, exist_ok=True)
    phase = BIMPhase(argv[2]) if len(argv) == 3 else BIMPhase.AVP

    catalog = build_catalog(
        cch_pdf=config.I3F_CCH_PDF,
        data_spec_xlsx=config.I3F_DATA_SPEC_XLSX,
        naming_spec_xlsx=config.I3F_NAMING_SPEC_XLSX,
    )
    # Garde CCH (helper pur, testé) : refuse un référentiel inexploitable.
    assert_catalog_usable(
        {
            "cch_pdf": config.I3F_CCH_PDF,
            "data_spec_xlsx": config.I3F_DATA_SPEC_XLSX,
            "naming_spec_xlsx": config.I3F_NAMING_SPEC_XLSX,
        },
        catalog,
    )

    client = BIMDataClient()  # cible + auth depuis l'environnement (read-only)
    snap = extract_snapshot(client)
    result = run_audit(snap, catalog, phase)

    # Vrai nom BIMData (peut être absent) vs nom de génération (bouchon si absent,
    # pour que le pack se génère quand même). Le contrôle de métadonnées n'accepte
    # QUE le vrai nom — jamais le bouchon.
    real_project_name = (snap.project or {}).get("name")
    pack = write_avp_i3f_report_pack(
        result,
        out,
        sources=None,  # chemin réel piloté par la maquette
        project_name=real_project_name or PLACEHOLDER_PROJECT_NAME,
        project_code="",
        phase=phase.value,
        export_pdf=False,
    )

    annexes = {
        "Contrôle": pack.controle_xlsx,
        "SHAB": pack.shab_xlsx,
        "Zones/Espaces": pack.zones_espaces_xlsx,
        "Enveloppe": pack.enveloppe_xlsx,
        "Menuiseries": pack.menuiseries_xlsx,
    }

    # Confidentialité : la sortie stdout ne porte AUCUNE donnée client (pas de
    # nom de modèle/projet). Uniquement phase, compteurs, booléens de charte,
    # verdict. Le pack complet (données client) reste hors repo.
    report: dict = {"phase": phase.value, "annexes": {}}
    ok = True
    for label, p in annexes.items():
        # Contrôle : compteur propre (points de contrôle sous la grille, hors
        # entête/légende/NOT_AVAILABLE) ; les 4 autres : compteur générique.
        rows = _count_controle_rows(p) if label == "Contrôle" else _count_business_rows(p)
        charte = _charte_flags(p, WORDMARK, BIMDATA_PRIMARY, BIMDATA_FONT_PRIMARY)
        entry = {"rows": rows, "non_empty": rows > 0, **charte}
        report["annexes"][label] = entry
        ok = ok and entry["non_empty"] and all(charte.values())

    # Rapport Word (analyse BIM AVP) : mêmes seuils que les tests (helper partagé).
    # Sortie = compteurs/booléens ; le vrai nom de projet est comparé au texte du
    # doc mais N'est PAS émis en clair.
    word = inspect_word_report(
        pack.analyse_docx,
        expected_project_name=real_project_name,
        phase=phase.value,
        wordmark=WORDMARK,
        primary=BIMDATA_PRIMARY,
        font=BIMDATA_FONT_PRIMARY,
        not_available=NOT_AVAILABLE,
    )
    report["word_report"] = word
    ok = ok and word["ok"]

    report["verdict"] = "PASS" if ok else "FAIL"
    # Sortie sûre : compteurs / booléens / verdict uniquement.
    print(json.dumps(report, ensure_ascii=False, indent=2))
    print("VERDICT:", report["verdict"])
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
