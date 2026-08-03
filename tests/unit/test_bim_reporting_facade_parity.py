"""Parité de la façade reporting : les primitives viennent bien du socle.

Cette suite est la conversion en test permanent de la preuve de parité menée
lors de l'extraction de `bim-reporting`. Elle tient deux rôles distincts :

1. **Identité** — les primitives bas niveau exposées par `audit_bim.reporting`
   *sont* celles du socle, pas des copies. Une réimplémentation locale qui
   réapparaîtrait ferait échouer ces tests.
2. **Comportement** — le rendu produit est inchangé, et surtout le garde-fou
   d'injection de formule couvre exactement le même jeu de caractères.

Le corpus du garde-fou est **dérivé** de la constante, jamais réécrit à la main.
C'est la leçon de `bim-reporting` v0.1.0 : une liste recopiée depuis une
docstring avait perdu la tabulation et le retour chariot, et la preuve de parité
d'alors, écrite depuis la même docstring, confirmait l'erreur au lieu de la
contredire. Un test qui ne peut pas contredire son auteur ne prouve rien.
"""

from __future__ import annotations

import io
from pathlib import Path

import bim_reporting.charts as bcharts
import bim_reporting.excel as bexcel
import bim_reporting.pdf as bpdf
import bim_reporting.theming as btheming
import bim_reporting.word as bword
import pytest
import xlsxwriter
from docx import Document

from audit_bim.reporting import bimdata_brand, pdf_export, theming, word_report, xlsx_annex

# ── 1. Identité : ré-exports directs, pas des copies ──────────────────


@pytest.mark.parametrize(
    ("local", "socle"),
    [
        (word_report._hex_to_rgb, bword.hex_to_rgb),
        (word_report._shade_cell, bword.shade_cell),
        (word_report._add_heading, bword.add_heading),
        (word_report._section_break, bword.section_break),
        (word_report._kpi_table, bword.kpi_table),
        (word_report._para_intro, bword.para_intro),
        (word_report._model_meta, bword.model_meta),
        (word_report._plt, bcharts.plt),
        (xlsx_annex._neutralize_formula, bexcel.neutralize_formula),
        (xlsx_annex._fmt_cell, bexcel.fmt_cell),
        (xlsx_annex.write_safe, bexcel.write_safe),
        (pdf_export.docx_to_pdf, bpdf.docx_to_pdf),
    ],
)
def test_primitive_is_the_socle_object(local, socle):
    assert local is socle


def test_brand_tokens_are_the_socle_objects():
    for name in ("BIMDATA_PRIMARY", "BIMDATA_SECONDARY", "BIMDATA_FONT_PRIMARY"):
        assert getattr(theming, name) is getattr(btheming, name)


def test_client_maps_stay_in_this_repo():
    """SEVERITY_COLORS et THEME_COLORS sont indexés par des énumérés métier.

    Les figer dans le socle obligerait un futur MCP à hériter du vocabulaire
    d'un autre — c'est la frontière que cette extraction pose.
    """
    assert not hasattr(btheming, "SEVERITY_COLORS")
    assert not hasattr(btheming, "THEME_COLORS")
    assert theming.SEVERITY_COLORS and theming.THEME_COLORS


# ── 2. Garde-fou d'injection : corpus DÉRIVÉ de la constante ──────────


def test_trigger_set_comes_from_the_socle():
    assert xlsx_annex._FORMULA_TRIGGERS is bexcel.FORMULA_TRIGGERS
    # Jeu historique d'audit-bim-i3f, tabulation et retour chariot compris.
    assert set(bexcel.FORMULA_TRIGGERS) == {"=", "+", "-", "@", "\t", "\r"}


def test_every_trigger_is_neutralized_including_invisible_ones():
    for char in xlsx_annex._FORMULA_TRIGGERS:  # dérivé, jamais recopié
        for payload in (char, f"{char}1+1", f"{char}cmd|' /C calc'!A0"):
            assert xlsx_annex._neutralize_formula(payload) == "'" + payload


def test_non_trigger_values_are_untouched():
    for value in ("Mur extérieur", "", 0, 42, -1.5, True, False, None):
        assert xlsx_annex._neutralize_formula(value) is value


def test_build_formats_exposes_one_format_per_severity():
    wb = xlsxwriter.Workbook(None)
    fmts = xlsx_annex._build_formats(wb)
    assert {"title", "header", "row", "row_alt", "kpi_key", "accent_filet"} <= set(fmts)
    for sev in theming.SEVERITY_COLORS:
        assert f"sev_{sev}" in fmts


# ── 3. Comportement : le rendu reste celui d'avant l'extraction ───────


def test_kv_or_na_keeps_the_repo_fallback_sentence():
    doc = Document()
    word_report._kv_or_na(doc, "Adresse", None)
    assert doc.paragraphs[-1].text == f"• Adresse : {word_report.NOT_AVAILABLE}"


def test_kv_or_na_keeps_the_repo_traceability_suffixes():
    doc = Document()
    word_report._kv_or_na(doc, "Adresse", "12 rue X", source="extracted")
    assert word_report.SOURCE_SUFFIX_EXTRACTED in doc.paragraphs[-1].text


def test_charts_keep_the_repo_labels():
    assert word_report._pie_chart({}, {}, "T").read(4) == b"\x89PNG"
    for build in (word_report._pie_chart, word_report._bar_chart):
        buf = build({"Nommage Pièce": 3}, theming.THEME_COLORS, "T")
        assert isinstance(buf, io.BytesIO) and buf.read(4) == b"\x89PNG"


def test_heading_level1_still_adds_the_accent_rule():
    doc = Document()
    before = len(doc.paragraphs)
    word_report._add_heading(doc, "Synthèse", 1)
    assert len(doc.paragraphs) == before + 2


# ── 4. Le logo doit rester résolu depuis CE dépôt ─────────────────────


def test_brand_search_origin_is_this_repo_not_site_packages():
    """Sans `search_from`, le logo disparaîtrait des livrables SANS erreur.

    L'absence de logo dégrade proprement vers le wordmark : la régression
    serait invisible en test comme en production. On vérifie donc que la façade
    ancre la recherche sur son propre fichier.
    """
    assert bimdata_brand._SEARCH_FROM.resolve() == Path(bimdata_brand.__file__).resolve()
    # La racine remontée doit être ce dépôt, pas l'installation du socle.
    assert "site-packages" not in str(bimdata_brand._SEARCH_FROM)


def test_brand_facade_resolves_kit_from_env(tmp_path, monkeypatch):
    kit = tmp_path / "kit"
    (kit / "assets").mkdir(parents=True)
    monkeypatch.setenv("BIMDATA_BRAND_KIT_DIR", str(kit))
    assert bimdata_brand.find_brand_kit_dir() == kit


def test_brand_facade_rejects_unknown_variant():
    with pytest.raises(ValueError, match="Variante logo inconnue"):
        bimdata_brand.find_logo("nope")
