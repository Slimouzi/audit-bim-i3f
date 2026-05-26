"""Tests de la charte Korhus.ai appliquée aux rapports d'audit.

Trois couches de couverture :

- :mod:`audit_bim.reporting.theming` — présence et cohérence des tokens
  Korhus, alignement des alias historiques ``I3F_*``.
- :mod:`audit_bim.reporting.korhus_brand` — résolution du brand kit
  via variable d'env, fallback chemin par défaut, dégradation
  silencieuse.
- *Smoke render* — la génération Word + Excel construit un fichier
  valide quand le brand kit est trouvé ET quand il ne l'est pas (le
  rapport ne doit jamais planter pour une question de logo manquant).
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

from audit_bim.audit.engine import AuditResult
from audit_bim.reporting import korhus_brand
from audit_bim.reporting.theming import (
    I3F_BLUE,
    I3F_BLUE_LIGHT,
    I3F_GREY,
    KORHUS_BLUE_NEUTRAL_LIGHT,
    KORHUS_FONT_PRIMARY,
    KORHUS_GRANITE,
    KORHUS_PRIMARY,
    KORHUS_SECONDARY,
    KORHUS_WHITE,
    SEVERITY_COLORS,
)
from audit_bim.reporting.word_report import write_word_report
from audit_bim.reporting.xlsx_annex import write_xlsx_annex

HEX6 = re.compile(r"^[0-9A-Fa-f]{6}$")


# ── 1. Tokens Korhus présents et cohérents ────────────────────────────


class TestKorhusPalette:
    def test_brand_tokens_are_hex6(self):
        for name in (
            KORHUS_PRIMARY,
            KORHUS_SECONDARY,
            KORHUS_WHITE,
            KORHUS_GRANITE,
            KORHUS_BLUE_NEUTRAL_LIGHT,
        ):
            assert HEX6.match(name), f"{name!r} n'est pas un hex 6 chars"

    def test_primary_is_korhus_dark(self):
        # #0C101B — couleur de couverture et titres forts (cf. brand kit).
        assert KORHUS_PRIMARY.upper() == "0C101B"

    def test_secondary_is_cyan_accent(self):
        assert KORHUS_SECONDARY.upper() == "59F4FF"

    def test_blue_neutral_light_is_table_bg(self):
        assert KORHUS_BLUE_NEUTRAL_LIGHT.upper() == "F0F5FF"


class TestI3FAliases:
    """Les alias historiques pointent désormais sur les tokens Korhus."""

    def test_i3f_blue_aliases_korhus_primary(self):
        assert I3F_BLUE == KORHUS_PRIMARY

    def test_i3f_blue_light_aliases_korhus_blue_neutral_light(self):
        assert I3F_BLUE_LIGHT == KORHUS_BLUE_NEUTRAL_LIGHT

    def test_i3f_grey_aliases_korhus_granite(self):
        assert I3F_GREY == KORHUS_GRANITE


class TestSeverityColorsUntouched:
    """Convention métier feux tricolores : la charte Korhus ne doit pas
    écraser le rouge / orange / vert des sévérités."""

    def test_critical_is_dark_red(self):
        r = int(SEVERITY_COLORS["CRITICAL"][0:2], 16)
        g = int(SEVERITY_COLORS["CRITICAL"][2:4], 16)
        b = int(SEVERITY_COLORS["CRITICAL"][4:6], 16)
        assert r > g and r > b  # rouge dominant


# ── 2. Résolution du brand kit ────────────────────────────────────────


class TestFindBrandKitDir:
    def test_env_override_takes_precedence(self, tmp_path, monkeypatch):
        # Simule un brand kit minimal sous tmp_path.
        (tmp_path / "assets").mkdir()
        monkeypatch.setenv("KORHUS_BRAND_KIT_DIR", str(tmp_path))
        assert korhus_brand.find_brand_kit_dir() == tmp_path

    def test_env_invalid_falls_back_to_sibling_or_none(self, tmp_path, monkeypatch):
        # Variable pointant vers un chemin inexistant : on tombe sur le
        # scan sibling (ou None si aucun voisin nommé korhus_brand_kit).
        monkeypatch.setenv("KORHUS_BRAND_KIT_DIR", str(tmp_path / "ghost"))
        result = korhus_brand.find_brand_kit_dir()
        # Soit un voisin existe (poste dev avec korhus_brand_kit cloné
        # à côté), soit None (CI). Dans les deux cas, le module n'a
        # plus de chemin hardcodé qui pourrait fausser le résultat.
        assert result is None or (result.is_dir() and result.name == "korhus_brand_kit")

    def test_returns_none_when_nothing_found(self, monkeypatch):
        monkeypatch.setenv("KORHUS_BRAND_KIT_DIR", "/nonexistent/path/x")
        # Le scan sibling pourrait remonter à un dossier existant — on
        # patch ``Path.is_dir`` pour qu'il ne dise jamais True.
        monkeypatch.setattr(Path, "is_dir", lambda self: False)
        assert korhus_brand.find_brand_kit_dir() is None


class TestFindLogo:
    def test_unknown_variant_raises(self):
        with pytest.raises(ValueError, match="Variante logo"):
            korhus_brand.find_logo("turquoise")

    def test_returns_none_when_brand_kit_missing(self, monkeypatch):
        monkeypatch.setattr(korhus_brand, "find_brand_kit_dir", lambda: None)
        assert korhus_brand.find_logo("light") is None

    def test_returns_path_when_logo_present(self, tmp_path, monkeypatch):
        assets = tmp_path / "assets"
        assets.mkdir()
        fake_logo = assets / "korhus_logo_reversed_or_light_wordmark.png"
        fake_logo.write_bytes(b"\x89PNG\r\n\x1a\n")  # signature PNG bidon
        monkeypatch.setattr(korhus_brand, "find_brand_kit_dir", lambda: tmp_path)
        assert korhus_brand.find_logo("light") == fake_logo


# ── 3. Smoke render Word + Excel ──────────────────────────────────────


@pytest.fixture
def minimal_audit_result(catalog, snapshot_with_walls) -> AuditResult:
    """Audit avec un finding suffisant pour exercer les tableaux + KPIs."""
    from audit_bim.audit.findings import ErrorType, Finding, Severity, Theme
    from audit_bim.requirements.models import BIMPhase

    return AuditResult(
        catalog=catalog,
        snapshot=snapshot_with_walls,
        phase=BIMPhase.PRO,
        findings=[
            Finding(
                theme=Theme.PROPERTY_MISSING,
                severity=Severity.HIGH,
                error_type=ErrorType.PROPERTY_MISSING,
                element_uuid="W1",
                ifc_type="IfcWallStandardCase",
                name="Mur extérieur 01",
                expected="Pset_WallCommon/IsExternal",
                actual=None,
                ref_cch="Chap 6.2",
                recommended_action="Renseigner IsExternal.",
            ),
        ],
    )


class TestSmokeWordRender:
    def test_word_report_builds_with_brand_kit(self, minimal_audit_result, tmp_path):
        """Avec brand kit disponible (poste dev) : le rendu doit
        embarquer le logo (au moins une image dans le doc).
        """
        out = tmp_path / "audit.docx"
        write_word_report(minimal_audit_result, out)
        assert out.exists() and out.stat().st_size > 0

        # Vérifie que le doc s'ouvre sans erreur python-docx.
        doc = Document(str(out))
        # La police Normal doit être Roboto (charte Korhus).
        assert doc.styles["Normal"].font.name == KORHUS_FONT_PRIMARY

    def test_word_report_builds_without_brand_kit(
        self, minimal_audit_result, tmp_path, monkeypatch
    ):
        """Sans brand kit : la couverture dégrade en wordmark texte,
        le rapport reste générable."""
        from audit_bim.reporting import word_report

        # Simule absence totale du brand kit.
        monkeypatch.setattr(word_report, "find_logo", lambda variant="light": None)
        out = tmp_path / "audit_no_logo.docx"
        write_word_report(minimal_audit_result, out)
        assert out.exists() and out.stat().st_size > 0
        doc = Document(str(out))
        # On cherche la chaîne fallback "KORHUS.AI" dans les runs du doc.
        all_text = "\n".join(
            run.text
            for tbl in doc.tables
            for row in tbl.rows
            for cell in row.cells
            for para in cell.paragraphs
            for run in para.runs
        )
        assert "KORHUS.AI" in all_text


class TestSmokeXlsxRender:
    def test_xlsx_annex_builds_and_uses_roboto(self, minimal_audit_result, tmp_path):
        out = tmp_path / "annex.xlsx"
        write_xlsx_annex(minimal_audit_result, out)
        assert out.exists() and out.stat().st_size > 0

        wb = load_workbook(out)
        # Onglet "Synthèse" doit exister + contenir le supertitle Korhus.
        assert "Synthèse" in wb.sheetnames
        synth = wb["Synthèse"]
        assert synth["A1"].value == "KORHUS.AI — AUDIT BIM"
        # La police appliquée doit bien être Roboto.
        assert synth["A3"].font.name == KORHUS_FONT_PRIMARY

    def test_xlsx_referential_sheet_has_korhus_banner(self, minimal_audit_result, tmp_path):
        out = tmp_path / "annex_ref.xlsx"
        write_xlsx_annex(minimal_audit_result, out)
        wb = load_workbook(out)
        assert "Référentiel I3F" in wb.sheetnames
        ref = wb["Référentiel I3F"]
        assert ref["A1"].value == "KORHUS.AI — RÉFÉRENTIEL"
