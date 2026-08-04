"""Le mode de filtrage d'enveloppe est pilotable depuis le tool, sans bricolage.

La recette Dieppe a d'abord été obtenue en filtrant **à la main** le contrat
``envelope.json`` produit par le backend : un résultat juste, mais non
reproductible par le produit — donc inutilisable pour un livrable client.

Ces tests verrouillent le chemin paramétrique de bout en bout :
``generate_avp_i3f_pack(envelope_filter_mode=…, envelope_type_pattern=…)``
atteint le backend, et un mode incohérent est refusé au lieu de se dégrader.

Ils verrouillent aussi la **note de lecture** : en façade Revit multicouche, les
baies sont portées par le mur porteur et non par la peau retenue comme façade.
La colonne « ouvertures » du livrable est alors nulle sur toutes les lignes face
à un total non nul — sans explication, cela se lit comme un défaut de calcul.
"""

from __future__ import annotations

import json

import pytest

from audit_bim.extraction.model_data import ModelSnapshot
from audit_bim.mcp.session import _Session, current_session
from audit_bim.profiles.i3f.tools_reporting import generate_avp_i3f_pack as tr_generate_avp_i3f_pack
from audit_bim.reporting import avp_autocompute
from audit_bim.reporting.avp.xlsx_enveloppe import _note_menuiseries, _note_methodologie
from audit_bim.reporting.avp_sources import read_envelope_json

PERIMETRE_AVANT_FILTRE = "murs_exterieurs_avant_filtre_type"


def _contrat(**diagnostics):
    return {
        "schema": "envelope_quantities/v1",
        "source": {"producer": "ifc-geometry", "ifc_file": "DIEPPE-7427L.ifc"},
        "created_at": "2026-08-03T08:00:00+00:00",
        "summary": {
            "superficie_facades_m2": 2206.19,
            "superficie_facades_nette_m2": 2206.19,
            "superficie_menuiseries_m2": 375.89,
            "shab_m2": 2392.64,
            "ratio_fac_shab": 0.9221,
            "methode_facade": "geometric_type_filter",
        },
        "par_type": [
            {
                "type": "Mur de base:MUR ENDUIT 20 mm",
                "etages": ["RDC"],
                "net_side_area_m2": 900.13,
                "n": 54,
                "menuiseries_m2": 0.0,
            }
        ],
        "hors_filtre_type": [
            {
                "type": "Mur de base:BETON 200mm",
                "etages": ["RDC"],
                "net_side_area_m2": 4001.30,
                "n": 279,
            }
        ],
        "diagnostics": diagnostics,
    }


@pytest.fixture
def session(tmp_path, monkeypatch):
    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    monkeypatch.setenv("AUDIT_INPUT_DIR", str(tmp_path))
    sess = _Session()
    sess.snapshot = ModelSnapshot(
        project={"name": "MCP_Audit"},
        model={"name": "DIEPPE-7427L.ifc"},
        elements=[{"uuid": "W1", "type": "IfcWall", "name": "Mur de base:MUR ENDUIT 20 mm"}],
    ).index()
    token = current_session.set(sess)
    try:
        yield sess, tmp_path
    finally:
        current_session.reset(token)


# ── le mode traverse bien tool -> autocompute -> backend ───────────────


def test_filter_mode_reaches_the_geometry_backend(session, tmp_path, monkeypatch):
    """Sans cette transmission, le mode resterait un paramètre décoratif."""
    recu = {}

    def _fake(ifc_path, **kw):
        recu.update(kw)
        return _contrat()

    monkeypatch.setattr(avp_autocompute, "compute_envelope_payload", _fake)
    ifc = tmp_path / "DIEPPE-7427L.ifc"
    ifc.write_text("ISO-10303-21;", encoding="utf-8")

    tr_generate_avp_i3f_pack(
        project_name="Dieppe",
        project_code="7427L",
        phase="APD",
        auditor_name="Stanislas Limouzi",
        envelope_filter_mode="geometric_type_filter",
        envelope_type_pattern=r"MUR ENDUIT|BARDAGE BOIS|ZINC|VERRE REGLIT",
        auto_compute_quantities=False,
        ifc_path=str(ifc),
        export_pdf=False,
    )

    assert recu["filter_mode"] == "geometric_type_filter"
    assert recu["type_pattern"] == r"MUR ENDUIT|BARDAGE BOIS|ZINC|VERRE REGLIT"
    assert recu["layer_pattern"] is None


def test_incoherent_filter_mode_is_refused_not_degraded(session, tmp_path, monkeypatch):
    """Le backend refuse ; le tool traduit en erreur d'appel, pas en exception."""

    def _fake(ifc_path, **kw):
        raise ValueError("``filter_mode='geometric_type_filter'`` exige ``type_pattern``")

    monkeypatch.setattr(avp_autocompute, "compute_envelope_payload", _fake)
    ifc = tmp_path / "DIEPPE-7427L.ifc"
    ifc.write_text("ISO-10303-21;", encoding="utf-8")

    res = tr_generate_avp_i3f_pack(
        project_name="Dieppe",
        project_code="7427L",
        phase="APD",
        auditor_name="Stanislas Limouzi",
        envelope_filter_mode="geometric_type_filter",
        auto_compute_quantities=False,
        ifc_path=str(ifc),
        export_pdf=False,
    )

    assert res["status"] == "error"
    assert res["error"] == "invalid_envelope_filter_mode"
    assert res["envelope_filter_mode"] == "geometric_type_filter"
    assert "type_pattern" in res["message"]


# ── la note de lecture des menuiseries ─────────────────────────────────


def test_envelope_source_carries_the_menuiserie_scope(tmp_path):
    chemin = tmp_path / "env.json"
    chemin.write_text(
        json.dumps(
            _contrat(
                menuiseries_perimetre=PERIMETRE_AVANT_FILTRE,
                menuiseries_m2_sur_types_rejetes=375.89,
            )
        ),
        encoding="utf-8",
    )

    src = read_envelope_json(chemin)

    assert src.menuiseries_perimetre == PERIMETRE_AVANT_FILTRE
    assert src.menuiseries_sur_types_rejetes == pytest.approx(375.89)


def test_note_explains_a_zero_openings_column(tmp_path):
    chemin = tmp_path / "env.json"
    chemin.write_text(
        json.dumps(
            _contrat(
                menuiseries_perimetre=PERIMETRE_AVANT_FILTRE,
                menuiseries_m2_sur_types_rejetes=375.89,
            )
        ),
        encoding="utf-8",
    )

    note = _note_menuiseries(read_envelope_json(chemin))

    assert note is not None
    assert "375.89" in note or "375,89" in note.replace(".", ",")
    assert "mur porteur" in note


def test_no_note_when_the_openings_are_attributed_to_kept_types(tmp_path):
    """Rien à expliquer sur une maquette ArchiCAD : pas de bruit inutile."""
    chemin = tmp_path / "env.json"
    chemin.write_text(
        json.dumps(
            _contrat(
                menuiseries_perimetre=PERIMETRE_AVANT_FILTRE,
                menuiseries_m2_sur_types_rejetes=0.0,
            )
        ),
        encoding="utf-8",
    )

    assert _note_menuiseries(read_envelope_json(chemin)) is None


def test_no_note_when_the_producer_says_nothing(tmp_path):
    """Contrat d'un producteur antérieur : aucune régression, aucune note."""
    chemin = tmp_path / "env.json"
    chemin.write_text(json.dumps(_contrat()), encoding="utf-8")

    src = read_envelope_json(chemin)

    assert src.menuiseries_perimetre is None
    assert _note_menuiseries(src) is None


# ── note de méthode : décrit le filtre RÉELLEMENT appliqué ─────────────


def _filtres(**kw):
    base = {
        "mode": "geometric_type_filter",
        "layer_pattern": None,
        "type_pattern": r"MUR ENDUIT|BARDAGE BOIS|ZINC|VERRE REGLIT",
        "types_retenus": ["Mur de base:MUR ENDUIT 20 mm"],
        "types_rejetes": ["Mur de base:BETON 200mm"],
    }
    base.update(kw)
    return base


def _source(tmp_path, **diagnostics):
    chemin = tmp_path / "env.json"
    chemin.write_text(json.dumps(_contrat(**diagnostics)), encoding="utf-8")
    return read_envelope_json(chemin)


def test_methodology_note_states_the_applied_filter(tmp_path):
    """Un texte générique affirmerait la même chose après un changement de mode.

    La note sort donc de ``diagnostics.filters`` : elle décrit ce que le calcul
    a fait, pas ce qu'on suppose qu'il a fait.
    """
    note = _note_methodologie(_source(tmp_path, filters=_filtres()))

    assert note is not None
    assert "IFC OpenShell" in note
    assert "double comptage" in note
    assert "MUR ENDUIT" in note


def test_methodology_note_follows_the_mode(tmp_path):
    """Mode ArchiCAD : le texte parle de calque, pas de peaux extérieures."""
    note = _note_methodologie(
        _source(
            tmp_path,
            filters=_filtres(
                mode="layer_type_filter", layer_pattern=r"221", type_pattern=r"^ME[ _]"
            ),
        )
    )

    assert "calque" in note
    assert "double comptage" not in note


def test_no_methodology_note_without_a_declared_filter(tmp_path):
    """Contrat d'un producteur antérieur : rien d'affirmé sur la méthode."""
    assert _note_methodologie(_source(tmp_path)) is None


# ── les notes atterrissent réellement dans les fichiers produits ───────


def test_notes_reach_the_generated_xlsx(session, tmp_path, monkeypatch):
    """Vérifier le helper ne dit rien de ce qui atterrit dans le classeur.

    C'est le même angle mort qui avait laissé passer la fuite « Solibri » : une
    logique juste et un livrable muet. On ouvre donc le fichier produit.
    """
    import openpyxl

    def _fake(ifc_path, **kw):
        return _contrat(
            filters=_filtres(),
            menuiseries_perimetre=PERIMETRE_AVANT_FILTRE,
            menuiseries_m2_sur_types_rejetes=375.89,
        )

    monkeypatch.setattr(avp_autocompute, "compute_envelope_payload", _fake)
    ifc = tmp_path / "DIEPPE-7427L.ifc"
    ifc.write_text("ISO-10303-21;", encoding="utf-8")

    res = tr_generate_avp_i3f_pack(
        project_name="Dieppe",
        project_code="7427L",
        phase="APD",
        auditor_name="Stanislas Limouzi",
        envelope_filter_mode="geometric_type_filter",
        envelope_type_pattern=r"MUR ENDUIT|BARDAGE BOIS|ZINC|VERRE REGLIT",
        auto_compute_quantities=False,
        ifc_path=str(ifc),
        export_pdf=False,
    )

    assert res.get("status") not in ("error", "needs_context"), res
    enveloppe = next(p for p in res["paths"] if "Extraction surface enveloppe" in p)
    wb = openpyxl.load_workbook(enveloppe)
    texte = "\n".join(
        str(c)
        for ws in wb.worksheets
        for row in ws.iter_rows(values_only=True)
        for c in row
        if isinstance(c, str)
    )
    wb.close()

    assert "IFC OpenShell" in texte
    assert "double comptage" in texte
    assert "mur porteur" in texte


# ── nature du dénominateur SHAB dans la note de méthode ────────────────


def test_shab_method_is_read_from_the_contract(tmp_path):
    """``summary.methode_shab`` est un champ ``extra`` du contrat V1."""
    chemin = tmp_path / "env.json"
    doc = _contrat(filters=_filtres())
    doc["summary"]["methode_shab"] = "pieces_zonees_hors_annexes"
    chemin.write_text(json.dumps(doc), encoding="utf-8")

    assert read_envelope_json(chemin).methode_shab == "pieces_zonees_hors_annexes"


def test_older_contract_declares_no_shab_method(tmp_path):
    """Antérieur à ifc-geometry-mcp v0.5.0 : l'absence reste silencieuse.

    Affirmer une méthode inconnue serait pire que de n'en affirmer aucune.
    """
    src = _source(tmp_path, filters=_filtres())

    assert src.methode_shab is None
    assert "SHAB :" not in (_note_methodologie(src) or "")


@pytest.mark.parametrize(
    ("methode", "attendu"),
    [
        ("pieces_zonees_hors_annexes", "pièces zonées hors annexes non habitables"),
        ("toutes_pieces_hors_annexes_sans_zonage", "repli sans zonage"),
    ],
)
def test_note_states_the_nature_of_the_shab(tmp_path, methode, attendu):
    """Un ratio ne se compare qu'à un ratio dont la SHAB est de même nature."""
    chemin = tmp_path / "env.json"
    doc = _contrat(filters=_filtres())
    doc["summary"]["methode_shab"] = methode
    chemin.write_text(json.dumps(doc), encoding="utf-8")

    note = _note_methodologie(read_envelope_json(chemin))

    assert "SHAB :" in note
    assert attendu in note


def _generer_pack_avec_shab(tmp_path, monkeypatch, methode="pieces_zonees_hors_annexes"):
    def _fake(ifc_path, **kw):
        doc = _contrat(
            filters=_filtres(),
            menuiseries_perimetre=PERIMETRE_AVANT_FILTRE,
            menuiseries_m2_sur_types_rejetes=375.89,
        )
        doc["summary"]["methode_shab"] = methode
        return doc

    monkeypatch.setattr(avp_autocompute, "compute_envelope_payload", _fake)
    ifc = tmp_path / "DIEPPE-7427L.ifc"
    ifc.write_text("ISO-10303-21;", encoding="utf-8")
    return tr_generate_avp_i3f_pack(
        project_name="Dieppe",
        project_code="7427L",
        phase="APD",
        auditor_name="Stanislas Limouzi",
        envelope_filter_mode="geometric_type_filter",
        envelope_type_pattern=r"MUR ENDUIT|BARDAGE BOIS|ZINC|VERRE REGLIT",
        auto_compute_quantities=False,
        ifc_path=str(ifc),
        export_pdf=False,
    )


def test_shab_method_reaches_the_generated_xlsx(session, tmp_path, monkeypatch):
    import openpyxl

    res = _generer_pack_avec_shab(tmp_path, monkeypatch)

    assert res.get("status") not in ("error", "needs_context"), res
    enveloppe = next(p for p in res["paths"] if "Extraction surface enveloppe" in p)
    wb = openpyxl.load_workbook(enveloppe)
    texte = "\n".join(
        str(c)
        for ws in wb.worksheets
        for row in ws.iter_rows(values_only=True)
        for c in row
        if isinstance(c, str)
    )
    wb.close()

    assert "SHAB : pièces zonées hors annexes non habitables." in texte


def test_shab_method_reaches_the_generated_docx(session, tmp_path, monkeypatch):
    """Le Word est le document que le client lit en premier."""
    from docx import Document

    res = _generer_pack_avec_shab(
        tmp_path, monkeypatch, methode="toutes_pieces_hors_annexes_sans_zonage"
    )

    assert res.get("status") not in ("error", "needs_context"), res
    d = Document(res["analyse_docx"])
    texte = "\n".join(p.text for p in d.paragraphs)
    texte += "\n" + "\n".join(c.text for t in d.tables for r in t.rows for c in r.cells)

    assert "SHAB : repli sans zonage, toutes pièces hors annexes." in texte
