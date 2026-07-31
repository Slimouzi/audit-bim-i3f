"""Tests de la validation de contexte avant audit/rapport.

Couvre :
- ``full_audit`` refuse si l'une des 3 infos obligatoires manque
  (adresse, phase, auditeur).
- ``generate_word_report`` refuse de même.
- Les deux tools acceptent ``confirm_context=True`` pour passer
  outre la validation.
- Les inputs utilisateur écrasent les valeurs déduites du snapshot
  et sont marqués ``source="user"``.
- Le rapport Word généré affiche les valeurs utilisateur **sans** la
  mention « déduit — à confirmer », et **avec** cette mention pour les
  valeurs extraites du snapshot.
"""

from __future__ import annotations

from types import SimpleNamespace

import pytest
from docx import Document

from audit_bim.audit.engine import AuditResult
from audit_bim.audit.findings import ErrorType, Finding, Severity, Theme
from audit_bim.extraction.model_data import ModelSnapshot
from audit_bim.mcp import phase
from audit_bim.mcp import server as mcp_server
from audit_bim.mcp import tools_audit as ta
from audit_bim.mcp.session import _Session, current_session
from audit_bim.reporting.context import (
    merge_user_context,
)
from audit_bim.requirements.models import BIMPhase, RequirementsCatalog

# ── Fixtures ────────────────────────────────────────────────────────────


@pytest.fixture
def _isolated(tmp_path, monkeypatch):
    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    sess = _Session()
    token = current_session.set(sess)
    try:
        yield sess, tmp_path
    finally:
        current_session.reset(token)


def _minimal_catalog() -> RequirementsCatalog:
    return RequirementsCatalog(
        cch_version="3.6",
        cch_source_pdf="/tmp/cch.pdf",
        data_spec_source="/tmp/data.xlsx",
        naming_spec_source="/tmp/naming.xlsx",
        properties=[],
        naming_rules=[],
        storey_names=[],
        zone_specs=[],
        room_specs=[],
    )


def _wire_audit(sess) -> None:
    """Pose un AuditResult minimal pour permettre generate_word_report."""
    snap = ModelSnapshot(
        project={
            "name": "Programme Test",
            "description": "Programme de 24 logements collectifs — phase test.",
        },
        model={"name": "TEST.ifc"},
        sites=[
            {
                "uuid": "S1",
                "name": "Site Liffré",
                "long_name": "12 rue de la Paix, 35340 LIFFRÉ",
                "type": "IfcSite",
                "SiteAddress": {
                    "AddressLines": ["12 rue de la Paix"],
                    "PostalCode": "35340",
                    "Town": "LIFFRÉ",
                },
            }
        ],
        buildings=[{"uuid": "B1", "name": "Bât A", "type": "IfcBuilding"}],
        storeys=[],
        spaces=[],
        zones=[],
        elements=[],
    ).index()
    sess.result = AuditResult(
        phase=BIMPhase.PRO,
        catalog=_minimal_catalog(),
        snapshot=snap,
        findings=[
            Finding(
                theme=Theme.NAMING_SPACE,
                severity=Severity.MEDIUM,
                error_type=ErrorType.NAMING_MISSING,
                element_uuid="X1",
                ifc_type="IfcSpace",
            )
        ],
    )
    sess.snapshot = snap


def _doc_text(path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ── generate_word_report : validation ────────────────────────────────────


class TestGenerateWordReportValidation:
    def test_refuses_when_address_missing(self, _isolated):
        sess, _ = _isolated
        _wire_audit(sess)
        res = mcp_server.generate_word_report(
            project_phase="PRO",
            auditor_name="Stanislas",
            # project_address manquant
        )
        assert res.get("status") == "needs_context"
        assert "project_address" in res["missing"]
        assert any(q["key"] == "project_address" for q in res["questions"])

    def test_refuses_when_phase_missing(self, _isolated):
        sess, _ = _isolated
        _wire_audit(sess)
        res = mcp_server.generate_word_report(
            project_address="12 rue de la Paix",
            auditor_name="Stanislas",
            # project_phase manquant
        )
        assert res.get("status") == "needs_context"
        assert "project_phase" in res["missing"]

    def test_refuses_when_auditor_missing(self, _isolated):
        sess, _ = _isolated
        _wire_audit(sess)
        res = mcp_server.generate_word_report(
            project_address="12 rue de la Paix",
            project_phase="PRO",
            # auditor_name manquant
        )
        assert res.get("status") == "needs_context"
        assert "auditor_name" in res["missing"]

    def test_refuses_when_phase_invalid(self, _isolated):
        sess, _ = _isolated
        _wire_audit(sess)
        res = mcp_server.generate_word_report(
            project_address="12 rue de la Paix",
            project_phase="NOPE",  # phase invalide
            auditor_name="Stanislas",
        )
        assert res.get("status") == "needs_context"
        assert "project_phase" in res["missing"]

    def test_accepts_when_all_fields_provided(self, _isolated, tmp_path):
        sess, _ = _isolated
        _wire_audit(sess)
        res = mcp_server.generate_word_report(
            output_path="rapport_complet.docx",
            project_address="12 rue de la Paix, 35340 LIFFRÉ",
            project_phase="PRO",
            auditor_name="Stanislas Limouzi",
            project_description="Programme de 24 logements collectifs.",
        )
        # Pas de needs_context — on a un path et size_bytes.
        assert "path" in res
        assert res.get("status") != "needs_context"
        assert (tmp_path / "rapport_complet.docx").exists()

    def test_confirm_context_bypasses_validation(self, _isolated):
        sess, _ = _isolated
        _wire_audit(sess)
        res = mcp_server.generate_word_report(
            output_path="rapport_minimal.docx",
            confirm_context=True,
            # Pas d'address ni phase ni auditor — mais confirm=True passe.
        )
        assert "path" in res
        assert res.get("status") != "needs_context"


# ── full_audit : validation (sans déclencher l'audit complet) ───────────


class TestFullAuditValidation:
    def test_default_push_mode_runs_full_audit_without_user_choice(self, _isolated, monkeypatch):
        """Sans ``push_mode`` explicite, le MCP lance le full audit nominal.

        Verrou produit : la classification et la publication ne doivent pas
        redevenir une question bloquante avant l'audit.
        """
        sess, tmp_path = _isolated
        calls: dict[str, str] = {}
        audit_result = SimpleNamespace(
            summary=lambda: {"n_findings": 0, "conformity_rate": 1.0},
            findings=[],
        )

        monkeypatch.setattr(ta, "_fa_prepare_catalog", lambda: None)
        monkeypatch.setattr(
            ta,
            "_fa_finalize_target",
            lambda *args, **kwargs: setattr(sess, "phase", BIMPhase.PRO),
        )
        monkeypatch.setattr(ta, "_fa_extract_snapshot", lambda *args, **kwargs: None)
        monkeypatch.setattr(ta, "run_audit", lambda *args, **kwargs: audit_result)
        monkeypatch.setattr(
            ta,
            "_fa_write_deliverables",
            lambda **kwargs: (
                tmp_path / "audit.docx",
                tmp_path / "audit.xlsx",
                tmp_path / "audit.docx",
            ),
        )
        monkeypatch.setattr(
            ta, "_fa_write_findings_json", lambda word_path: tmp_path / "audit_findings.json"
        )

        def _publication(mode: str) -> dict:
            calls["mode"] = mode
            return {"push_mode": mode}

        monkeypatch.setattr(ta, "_fa_prepare_publication", _publication)

        res = mcp_server.full_audit(
            phase="PRO",
            project_address="12 rue de la Paix",
            auditor_name="Stanislas Limouzi",
            project_description="Programme test.",
        )

        assert res.get("status") != "needs_user_choice"
        assert calls["mode"] == "none"
        assert res["publication"] == {"push_mode": "none"}
        assert "deliverables" in res

    def test_refuses_when_address_missing(self, _isolated):
        sess, _ = _isolated
        # Pas besoin de wire — la validation tombe AVANT toute exécution.
        res = mcp_server.full_audit(
            phase="PRO",
            auditor_name="Stanislas",
            push_mode="none",
        )
        assert res.get("status") == "needs_context"
        assert "project_address" in res["missing"]

    def test_refuses_when_auditor_missing(self, _isolated):
        sess, _ = _isolated
        res = mcp_server.full_audit(
            phase="PRO",
            project_address="12 rue de la Paix",
            push_mode="none",
        )
        assert res.get("status") == "needs_context"
        assert "auditor_name" in res["missing"]

    def test_refuses_when_phase_invalid(self, _isolated):
        sess, _ = _isolated
        res = mcp_server.full_audit(
            phase="WRONG",
            project_address="12 rue de la Paix",
            auditor_name="Stan",
            push_mode="none",
        )
        assert res.get("status") == "needs_context"
        assert "project_phase" in res["missing"]

    def test_validation_fires_before_push_mode_ask(self, _isolated):
        """La validation doit s'exécuter AVANT le check ``push_mode=ask``.
        Sinon, un agent obtiendrait la question push_mode avant de
        savoir qu'il lui manque adresse/auditeur."""
        sess, _ = _isolated
        # Même avec push_mode=ask explicite, on n'a pas fourni adresse/auditeur
        # → on doit obtenir needs_context, pas needs_user_choice.
        res = mcp_server.full_audit(phase="PRO", push_mode="ask")
        assert res.get("status") == "needs_context"
        assert res.get("status") != "needs_user_choice"


# ── Rapport Word : marquage source ──────────────────────────────────────


class TestWordReportSourceMarking:
    def test_user_provided_values_have_no_deduced_suffix(self, _isolated, tmp_path):
        sess, _ = _isolated
        _wire_audit(sess)
        mcp_server.generate_word_report(
            output_path="rapport.docx",
            project_address="42 boulevard Saint-Germain, 75005 PARIS",
            project_phase="DCE",
            auditor_name="Stanislas Limouzi",
            project_description="Réhabilitation.",
        )
        text = _doc_text(tmp_path / "rapport.docx")
        # Adresse user-fournie présente, SANS suffixe "à confirmer"
        assert "42 boulevard Saint-Germain" in text
        # Pas de suffixe "déduit — à confirmer" attaché à l'adresse user.
        # On cherche la ligne contenant l'adresse :
        for para_text in text.split("\n"):
            if "42 boulevard Saint-Germain" in para_text:
                assert "déduit" not in para_text, (
                    f"L'adresse user-fournie ne doit pas porter le suffixe « déduit » : "
                    f"{para_text!r}"
                )

    def test_extracted_values_carry_deduced_suffix(self, _isolated, tmp_path):
        """Quand l'adresse est extraite du snapshot (IfcSite.long_name)
        et PAS fournie par l'utilisateur, elle porte le suffixe
        « déduit — à confirmer »."""
        sess, _ = _isolated
        _wire_audit(sess)  # le snapshot contient une adresse via IfcSite
        res = mcp_server.generate_word_report(
            output_path="rapport_extracted.docx",
            # On ne fournit PAS project_address user, mais on bypass
            # la validation pour pouvoir générer le rapport.
            project_phase="PRO",
            auditor_name="Stan",
            confirm_context=True,
        )
        assert "path" in res
        text = _doc_text(tmp_path / "rapport_extracted.docx")
        # L'adresse extraite est présente
        assert "12 rue de la Paix" in text
        # Chercher la ligne adresse et vérifier qu'elle porte le suffixe
        found_line_with_address = False
        for para_text in text.split("\n"):
            if "12 rue de la Paix" in para_text and "Adresse" in para_text:
                found_line_with_address = True
                assert "à confirmer" in para_text, (
                    f"L'adresse extraite doit porter le suffixe « à confirmer » : {para_text!r}"
                )
        assert found_line_with_address, "Ligne « Adresse : ... » non trouvée"

    def test_auditor_name_appears_on_cover_page_and_in_context(self, _isolated, tmp_path):
        sess, _ = _isolated
        _wire_audit(sess)
        mcp_server.generate_word_report(
            output_path="rapport.docx",
            project_address="X",
            project_phase="PRO",
            auditor_name="Jean DUPONT (BET Acme)",
            project_description="Projet.",
        )
        text = _doc_text(tmp_path / "rapport.docx")
        # Page de garde : ligne « Auteur : ... »
        assert "Auteur : Jean DUPONT (BET Acme)" in text

    def test_phase_user_provided_displayed_correctly(self, _isolated, tmp_path):
        sess, _ = _isolated
        _wire_audit(sess)  # snapshot a phase PRO via _wire_audit
        mcp_server.generate_word_report(
            output_path="rapport.docx",
            project_address="X",
            project_phase="DCE",  # user fournit DCE
            auditor_name="Stan",
            project_description="Projet.",
        )
        text = _doc_text(tmp_path / "rapport.docx")
        # Phase DCE doit apparaître (user-fournie)
        assert "DCE" in text


# ── Non-régression : sections enrichies toujours présentes ──────────────


class TestEnrichedSectionsStillPresent:
    def test_all_enriched_sections_present_in_validated_report(self, _isolated, tmp_path):
        sess, _ = _isolated
        _wire_audit(sess)
        mcp_server.generate_word_report(
            output_path="rapport.docx",
            project_address="X",
            project_phase="PRO",
            auditor_name="Stan",
            project_description="Projet.",
        )
        text = _doc_text(tmp_path / "rapport.docx")
        # Sections du modèle de rapport de conformité (structure 0.3)
        for section in (
            "2. Synthèse exécutive",
            "3. Périmètre de l'audit",
            "Maquette auditée",
            "4. Méthodologie",
            "5. Résultats globaux",
            "6. Résultats détaillés",
            "7. Liste des non-conformités",
            "8. Recommandations",
            "9. Conclusion",
            "10. Annexes",
        ):
            assert section in text, f"section manquante dans le rapport : {section!r}"


# ── Phase : question unique, aide loi MOP, détection + mapping ──────────


class TestPhaseHelpers:
    def test_map_phase_valid_passthrough(self):
        assert phase._map_phase("AVP") == "AVP"
        assert phase._map_phase("avp") == "AVP"

    def test_map_phase_loi_mop_aliases(self):
        assert phase._map_phase("APD") == "AVP"
        assert phase._map_phase("ACT") == "EXE"
        assert phase._map_phase("VISA") == "EXE"
        assert phase._map_phase("DET") == "EXE"

    def test_map_phase_unknown_returns_none(self):
        assert phase._map_phase("ZZZ") is None
        assert phase._map_phase("") is None
        assert phase._map_phase(None) is None


class TestPhaseValidationDialogue:
    def test_single_phase_question_with_reading_aid(self):
        """Une seule question de phase, avec l'aide loi MOP embarquée
        (pas de second champ)."""
        res = phase._validate_audit_context(
            project_address="X",
            project_phase=None,
            auditor_name="Stan",
            confirm_context=False,
        )
        phase_qs = [q for q in res["questions"] if q["key"] == "project_phase"]
        assert len(phase_qs) == 1
        q = phase_qs[0]
        assert "phase du projet à auditer" in q["question"]
        # Aide de lecture loi MOP présente dans la même question.
        aid = q["aide_lecture_loi_mop"]
        assert aid["APS"] == "avant-projet sommaire"
        assert "APD" in aid["AVP"]
        assert "GESTION" in aid

    def test_recognized_detected_phase_asks_confirmation(self):
        res = phase._validate_audit_context(
            project_address="X",
            project_phase="AVP",
            auditor_name="Stan",
            confirm_context=False,
            suggested_phase="AVP",
            detected_phase_raw="AVP",
            require_phase_confirmation=True,
        )
        q = next(q for q in res["questions"] if q["key"] == "project_phase")
        assert q.get("suggested_value") == "AVP"
        assert "Phase détectée dans l'IFC" in q["question"]
        assert "Confirmez-vous" in q["question"]

    def test_unrecognized_phase_proposes_mapping(self):
        res = phase._validate_audit_context(
            project_address="X",
            project_phase="AVP",
            auditor_name="Stan",
            confirm_context=False,
            suggested_phase="AVP",
            detected_phase_raw="APD",
            require_phase_confirmation=True,
        )
        q = next(q for q in res["questions"] if q["key"] == "project_phase")
        assert "Phase détectée : « APD »" in q["question"]
        assert "Proposition d'audit : AVP" in q["question"]

    def test_explicit_valid_phase_no_confirmation(self):
        """Phase valide + pas de demande de confirmation → pas de question."""
        res = phase._validate_audit_context(
            project_address="X",
            project_phase="PRO",
            auditor_name="Stan",
            confirm_context=False,
            require_phase_confirmation=False,
        )
        assert res is None

    def test_no_duplicate_phase_field(self):
        """Une seule clé de phase dans les questions (pas de doublon
        loi MOP / phase BIM)."""
        res = phase._validate_audit_context(
            project_address=None,
            project_phase=None,
            auditor_name=None,
            confirm_context=False,
            require_phase_confirmation=True,
        )
        keys = [q["key"] for q in res["questions"]]
        assert keys.count("project_phase") == 1
        assert "mop_phase" not in keys


class TestFullAuditPhaseConfirmation:
    def test_no_explicit_phase_asks_confirmation(self, _isolated):
        """full_audit sans phase explicite → needs_context sur la phase
        (confirmation exigée), avant toute extraction."""
        sess, _ = _isolated
        res = mcp_server.full_audit(
            project_address="12 rue X",
            auditor_name="Stan",
            push_mode="none",
            # phase omise → confirmation exigée
        )
        assert res.get("status") == "needs_context"
        assert "project_phase" in res["missing"]

    def test_detected_phase_proposed_as_suggestion(self, _isolated):
        sess, _ = _isolated
        _wire_audit(sess)
        # Injecte une phase brute loi MOP dans le snapshot.
        sess.snapshot.project = {"name": "Programme Test", "phase": "APD"}
        res = mcp_server.full_audit(
            project_address="12 rue X",
            auditor_name="Stan",
            push_mode="none",
        )
        assert res.get("status") == "needs_context"
        q = next(q for q in res["questions"] if q["key"] == "project_phase")
        # APD (loi MOP) → proposition AVP.
        assert q.get("suggested_value") == "AVP"
        assert "APD" in q["question"]


# ── R4 : suggestions maquette + validation description ──────────────────


class TestValidateAuditContextSuggestions:
    """Tests unitaires directs de ``_validate_audit_context`` (R4)."""

    def test_address_question_carries_suggested_value(self):
        res = phase._validate_audit_context(
            project_address=None,
            project_phase="PRO",
            auditor_name="Stan",
            confirm_context=False,
            suggested_address="12 rue de la Paix 35340 LIFFRÉ",
        )
        assert res is not None
        q = next(q for q in res["questions"] if q["key"] == "project_address")
        assert q.get("suggested_value") == "12 rue de la Paix 35340 LIFFRÉ"
        assert "12 rue de la Paix" in q["question"]

    def test_address_question_without_suggestion_has_no_suggested_value(self):
        res = phase._validate_audit_context(
            project_address=None,
            project_phase="PRO",
            auditor_name="Stan",
            confirm_context=False,
        )
        q = next(q for q in res["questions"] if q["key"] == "project_address")
        assert "suggested_value" not in q

    def test_description_not_required_by_default(self):
        """Sans ``require_description`` (cold start / pas de snapshot),
        la description absente ne bloque pas."""
        res = phase._validate_audit_context(
            project_address="X",
            project_phase="PRO",
            auditor_name="Stan",
            confirm_context=False,
            project_description=None,
        )
        assert res is None

    def test_description_required_and_missing_is_asked(self):
        res = phase._validate_audit_context(
            project_address="X",
            project_phase="PRO",
            auditor_name="Stan",
            confirm_context=False,
            project_description=None,
            require_description=True,
        )
        assert res is not None
        assert "project_description" in res["missing"]
        q = next(q for q in res["questions"] if q["key"] == "project_description")
        assert "suggested_value" not in q

    def test_description_question_carries_suggestion(self):
        res = phase._validate_audit_context(
            project_address="X",
            project_phase="PRO",
            auditor_name="Stan",
            confirm_context=False,
            project_description=None,
            require_description=True,
            suggested_description="Résidence 24 logements",
        )
        q = next(q for q in res["questions"] if q["key"] == "project_description")
        assert q.get("suggested_value") == "Résidence 24 logements"

    def test_description_satisfied_when_provided(self):
        res = phase._validate_audit_context(
            project_address="X",
            project_phase="PRO",
            auditor_name="Stan",
            confirm_context=False,
            project_description="Une description fournie",
            require_description=True,
        )
        assert res is None


class TestWordReportSuggestionsAndDescription:
    def test_address_question_uses_snapshot_suggestion(self, _isolated):
        """generate_word_report sans adresse → la question propose
        l'adresse extraite de l'IfcSite.SiteAddress."""
        sess, _ = _isolated
        _wire_audit(sess)
        res = mcp_server.generate_word_report(
            project_phase="PRO",
            auditor_name="Stan",
            # project_address omis
        )
        assert res.get("status") == "needs_context"
        q = next(q for q in res["questions"] if q["key"] == "project_address")
        assert q.get("suggested_value") == "12 rue de la Paix 35340 LIFFRÉ"

    def test_snapshot_description_is_proposed_not_auto_accepted(self, _isolated):
        """La description du snapshot n'est PAS acceptée en silence : la
        question est posée avec la description maquette en suggestion, à
        valider/corriger par l'utilisateur (attendu CTO)."""
        sess, _ = _isolated
        _wire_audit(sess)  # snapshot avec project.description
        res = mcp_server.generate_word_report(
            project_address="X",
            project_phase="PRO",
            auditor_name="Stan",
            # project_description omis → doit être DEMANDÉE (pas reprise en silence)
        )
        assert res.get("status") == "needs_context"
        assert "project_description" in res["missing"]
        q = next(q for q in res["questions"] if q["key"] == "project_description")
        assert q.get("suggested_value") == "Programme de 24 logements collectifs — phase test."

    def test_user_description_flows_to_report(self, _isolated, tmp_path):
        sess, _ = _isolated
        _wire_audit(sess)
        mcp_server.generate_word_report(
            output_path="rapport_userdesc.docx",
            project_address="X",
            project_phase="PRO",
            auditor_name="Stan",
            project_description="Réhabilitation d'un immeuble de 12 logements sociaux.",
        )
        text = _doc_text(tmp_path / "rapport_userdesc.docx")
        assert "Réhabilitation d'un immeuble de 12 logements sociaux." in text

    def test_missing_description_is_asked_when_snapshot_silent(self, _isolated):
        """Snapshot sans description ET utilisateur silencieux →
        needs_context sur project_description (P1b)."""
        sess, _ = _isolated
        _wire_audit(sess)
        # Retire la description du snapshot pour simuler une maquette muette.
        sess.snapshot.project = {"name": "Programme Test"}
        res = mcp_server.generate_word_report(
            project_address="X",
            project_phase="PRO",
            auditor_name="Stan",
            # project_description omis, snapshot muet
        )
        assert res.get("status") == "needs_context"
        assert "project_description" in res["missing"]


# ── merge_user_context : tests unitaires de la primitive ────────────────


class TestMergeUserContext:
    def test_overwrites_address_and_marks_source_user(self):
        # Build un contexte minimal avec adresse extraite
        from audit_bim.reporting.context import ReportProjectContext

        base = ReportProjectContext(
            address="Adresse déduite snapshot",
            field_sources={"address": "extracted"},
        )
        new = merge_user_context(base, project_address="Adresse user explicite")
        assert new.address == "Adresse user explicite"
        assert new.source_of("address") == "user"

    def test_no_input_returns_same_instance(self):
        from audit_bim.reporting.context import ReportProjectContext

        base = ReportProjectContext(address="X", field_sources={"address": "extracted"})
        new = merge_user_context(base)
        assert new is base
        # Source inchangée
        assert new.source_of("address") == "extracted"

    def test_empty_string_ignored(self):
        from audit_bim.reporting.context import ReportProjectContext

        base = ReportProjectContext(address="X", field_sources={"address": "extracted"})
        new = merge_user_context(base, project_address="   ")
        # Strings blanches sont ignorées
        assert new is base

    def test_no_hallucination_via_merge(self):
        """``merge_user_context`` ne doit injecter une valeur que si
        l'utilisateur la fournit explicitement. ``None`` n'écrase pas."""
        from audit_bim.reporting.context import ReportProjectContext

        base = ReportProjectContext(
            address="Adresse fiable",
            field_sources={"address": "user"},
        )
        new = merge_user_context(base, project_address=None)
        # Adresse user-fournie inchangée
        assert new.address == "Adresse fiable"
        assert new.source_of("address") == "user"

    def test_missing_information_cleaned_when_field_provided(self):
        """Si missing_information contenait une entrée pour un champ
        que l'utilisateur a comblé, elle doit être retirée."""
        from audit_bim.reporting.context import ReportProjectContext

        base = ReportProjectContext(
            address=None,
            field_sources={"address": "missing"},
            missing_information=[
                "Adresse du projet : non renseignée sur l'IfcSite ni dans les métadonnées BIMData.",
                "Maîtrise d'ouvrage : non identifiée formellement.",
            ],
        )
        new = merge_user_context(base, project_address="Nouvelle adresse")
        assert new.address == "Nouvelle adresse"
        # L'entrée "Adresse du projet" doit avoir disparu
        assert not any("Adresse du projet" in m for m in new.missing_information)
        # Mais l'autre entrée reste
        assert any("Maîtrise d'ouvrage" in m for m in new.missing_information)
