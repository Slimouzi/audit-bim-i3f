"""Tests du pack de livrables AVP I3F (``avp_i3f``).

Sources synthétiques (openpyxl) → génération → relecture. Couvre :
structure d'onglets, ordre des en-têtes, charte BIMData, absence de
l'ancienne charte, principe « ne jamais inventer », sections du consolidé.
"""

from __future__ import annotations

import zipfile

import openpyxl
import pytest
from docx import Document

from audit_bim.reporting.avp_i3f import write_avp_i3f_report_pack
from audit_bim.reporting.avp_sources import AvpSourcePaths
from audit_bim.reporting.pdf_export import docx_to_pdf


def _wb(path, sheet_rows):
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for name, rows in sheet_rows.items():
        ws = wb.create_sheet(title=name[:31])
        for r in rows:
            ws.append(r)
    wb.save(str(path))
    return path


@pytest.fixture
def sources(tmp_path):
    src_dir = tmp_path / "src"
    src_dir.mkdir()
    enveloppe = _wb(
        src_dir / "env.xlsx",
        {
            "TDB 2022 04.2": [
                ["Composant", "Type", "Étages", "Archicad BQ NetSideArea", "Surface Solibri"],
                ["Mur", "ME_36", "RDC", 313.14, 325.33],
                [],
                [None, None, "ratio FAC/SHAB : ", 0.9567],
                [None, None, "Seuil 3F 2026 : ", 0.9],
                [None, None, "SHAB : ", 2164.98],
            ]
        },
    )
    shab = _wb(
        src_dir / "260201 shab.xlsx",
        {
            # Onglet pivot (doit être préservé) + onglet détail TDB.
            "Feuil1": [
                ["SHAB (Qté de Base)", "Pièces"],
                ["Logement", "SEJOUR", "CHAMBRE 01", "Total général"],
                ["Zone Logement T3", "", 12.98, 12.98],
            ],
            "TDB 2022 01.3 - Export Zones": [
                ["Composant", "Nom Zone", "Pièce", "Surface Nette (Qté de Base)", "Étage"],
                ["Zone", "0546L-1101", "CHAMBRE 01", 12.98, "R+1"],
            ],
        },
    )
    zones = _wb(
        src_dir / "zones.xlsx",
        {
            "TDB 2022 01.3 - Export Zones": [
                ["Composant", "Nom Zone", "Pièce (Nombre)", "Surface Nette (Qté de Base)"],
                ["Zone", "0546L-1101", "CHAMBRE 01", 12.98],
            ]
        },
    )
    menuiseries = _wb(
        src_dir / "men.xlsx",
        {
            "TDB 2022 05.1 - Fenêtres": [
                ["Composant", "Type", "Matériau", "Largeur", "Hauteur"],
                ["Fenêtre", "F25", None, 0.6, 1.3],
                [None, "Nombre de types de menuiseries", 1],
            ]
        },
    )
    controle = _wb(
        src_dir / "ctrl.xlsx",
        {
            "Grille de contrôle": [
                [None, "Projet", "Tarare"],
                [None, "ESI", "0546L"],
                [None, "Phase", "AVP"],
                [None, None, None, None, 0, "Non fourni / non trouvé"],
                [None, None, None, None, 2, "Satisfaisant"],
                [
                    "CODE 3F",
                    "POINTS DE CONTROLE",
                    "EXIGENCE CCH BIM 3F",
                    "Outil utilisé",
                    "EVALUATION",
                    "Commentaires CdP Bim",
                ],
                ["1.1", "Conformité plans", "les plans…", "", 0, "non testé"],
                ["4.1", "Présence de zones", "6.1.2", "", 2, ""],
            ],
            "Pièces Nommage": [
                [None, "Nombre de Noms"],
                [None, None, "MN", None, "Conforme", None, "Non Conforme"],
                [None, "Nombre de Noms", 316, 247, 0.7816, 16, 0.0506],
            ],
            "Zones Nommage": [
                [None, None, "MN", None, "Conforme", None, "Non Conforme"],
                [None, "Noms (nbre)", 24, 24, 1, 0, 0],
            ],
            "ARC bsence de matériau": [
                [None, None, "MN"],
                [
                    None,
                    "Nombre d'élements sans matériaux",
                    617,
                    0.0586,
                    None,
                    "Nombre d'élements :",
                    10530,
                ],
            ],
        },
    )
    return AvpSourcePaths(
        controle=controle,
        shab=shab,
        zones_espaces=zones,
        enveloppe=enveloppe,
        menuiseries=menuiseries,
    )


def _find_row(ws, anchor):
    for r in range(1, ws.max_row + 1):
        for c in range(1, ws.max_column + 1):
            if ws.cell(r, c).value == anchor:
                return r
    return None


# ── Structure ──────────────────────────────────────────────────────────


def test_pack_generates_six_deliverables(tmp_path, sources):
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    assert len(pack.paths()) == 6
    for p in pack.paths():
        assert p.exists() and p.stat().st_size > 0


def test_controle_has_expected_sheets(tmp_path, sources):
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    wb = openpyxl.load_workbook(pack.controle_xlsx)
    assert wb.sheetnames == [
        "Grille de contrôle",
        "Zones Nommage",
        "Pièces Nommage",
        "ARC absence de matériau",
        "Zones ObjectType",
    ]
    wb.close()


def test_export_headers_order_preserved(tmp_path, sources):
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    wb = openpyxl.load_workbook(pack.enveloppe_xlsx)
    ws = wb.active
    hr = _find_row(ws, "Composant")
    assert hr is not None
    headers = [ws.cell(hr, c).value for c in range(1, 6)]
    assert headers == ["Composant", "Type", "Étages", "Archicad BQ NetSideArea", "Surface Solibri"]
    wb.close()


def test_enveloppe_summary_block(tmp_path, sources):
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    wb = openpyxl.load_workbook(pack.enveloppe_xlsx)
    ws = wb.active
    r = _find_row(ws, "ratio FAC/SHAB")
    assert r is not None and ws.cell(r, 2).value == pytest.approx(0.9567)
    r2 = _find_row(ws, "Seuil 3F 2026")
    assert r2 is not None and ws.cell(r2, 2).value == pytest.approx(0.9)
    wb.close()


# ── Charte BIMData ───────────────────────────────────────────────────────


def test_bimdata_branding(tmp_path, sources):
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    wb = openpyxl.load_workbook(pack.enveloppe_xlsx)
    ws = wb.active
    # Bannière BIMDATA.
    assert str(ws["A1"].value).startswith("BIMDATA —")
    # En-tête de table : fond primaire 2F374A, police Roboto.
    hr = _find_row(ws, "Composant")
    cell = ws.cell(hr, 1)
    assert (cell.fill.fgColor.rgb or "").upper().endswith("2F374A")
    assert cell.font.name == "Roboto"
    wb.close()


def test_no_old_charter_in_outputs(tmp_path, sources):
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    for p in pack.paths():
        with zipfile.ZipFile(p) as z:
            blob = b"".join(z.read(n) for n in z.namelist() if n.endswith((".xml", ".rels")))
        assert b"KORHUS" not in blob.upper(), f"ancienne charte trouvée dans {p.name}"
        assert b"BIMDATA" in blob.upper()


# ── Ne jamais inventer ───────────────────────────────────────────────────


def test_never_invent_without_sources(tmp_path):
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=None, export_pdf=False)
    wb = openpyxl.load_workbook(pack.enveloppe_xlsx)
    ws = wb.active
    seen = {str(c.value) for row in ws.iter_rows() for c in row if isinstance(c.value, str)}
    assert any("Information non disponible dans les documents fournis" in v for v in seen)
    wb.close()


# ── Consolidé ────────────────────────────────────────────────────────────


def test_consolidated_docx_sections(tmp_path, sources):
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    doc = Document(str(pack.analyse_docx))
    txt = "\n".join(p.text for p in doc.paragraphs)
    for section in (
        "Analyse BIM",
        "1. Données d'entrée",
        "2. Usages BIM 3F",
        "3. Synthèse",
        "4. Indicateurs de conformité",
        "5. Écarts",
        "6. Grille de contrôle",
        "7. Points bloquants",
        "8. Recommandations AMO BIM",
        "9. Annexes",
    ):
        assert section in txt, f"section manquante : {section}"


# ── Conformité aux fichiers I3F réels (bugs de revue) ────────────────────


def test_exports_preserve_source_sheets(tmp_path, sources):
    # Les onglets pivot + détail des exports SHAB/Zones sont préservés.
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    wb = openpyxl.load_workbook(pack.shab_xlsx)
    assert wb.sheetnames == ["Feuil1", "TDB 2022 01.3 - Export Zones"]
    wb.close()


def test_filenames_follow_i3f_convention(tmp_path, sources):
    # Convention documentaire I3F générée depuis l'identité projet confirmée :
    # YYMMDD Nom Code Phase - TypeLivrable.ext (date de génération imposée ici).
    pack = write_avp_i3f_report_pack(
        None,
        tmp_path / "out",
        sources=sources,
        project_name="Tarare",
        project_code="0546L",
        phase="AVP",
        date="260702",
        export_pdf=False,
    )
    assert pack.controle_xlsx.name == "260702 Tarare 0546L AVP - Contrôle Maquettes.xlsx"
    assert pack.shab_xlsx.name == "260702 Tarare 0546L AVP - export SHAB maquette.xlsx"
    assert pack.zones_espaces_xlsx.name == "260702 Tarare 0546L AVP - Export Zones et Espaces.xlsx"
    assert pack.enveloppe_xlsx.name == "260702 Tarare 0546L AVP - Extraction surface enveloppe.xlsx"
    assert pack.menuiseries_xlsx.name == "260702 Tarare 0546L AVP - export Menuiseries.xlsx"
    assert pack.analyse_docx.name == "260702 Tarare 0546L AVP - Rapport analyse BIM.docx"


def test_filenames_default_date_is_generation_date(tmp_path, sources):
    from datetime import datetime

    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    today = datetime.now().strftime("%y%m%d")
    # Date de génération (YYMMDD) en préfixe, phase avant le tiret.
    assert pack.shab_xlsx.name.startswith(f"{today} ")
    assert pack.shab_xlsx.name.endswith(" AVP - export SHAB maquette.xlsx")


def test_filenames_omit_missing_code(tmp_path, sources):
    # Code absent → fragment simplement omis (jamais inventé).
    pack = write_avp_i3f_report_pack(
        None,
        tmp_path / "out",
        sources=sources,
        project_name="Tarare",
        project_code="",
        phase="AVP",
        date="260702",
        export_pdf=False,
    )
    assert pack.shab_xlsx.name == "260702 Tarare AVP - export SHAB maquette.xlsx"


def test_writer_defaults_no_client_identity(tmp_path, sources):
    # P3 : un appel direct sans identité ne doit PAS inventer « Tarare » /
    # « 0546L » (défauts génériques neutres).
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    for p in pack.paths():
        assert "Tarare" not in p.name
        assert "0546L" not in p.name


def test_filename_sanitizes_path_separators(tmp_path, sources):
    pack = write_avp_i3f_report_pack(
        None,
        tmp_path / "out",
        sources=sources,
        project_name="Rue A/B",
        project_code="05/46",
        phase="AVP",
        date="260702",
        export_pdf=False,
    )
    # Aucun séparateur de chemin ne doit subsister dans le nom de fichier.
    assert "/" not in pack.shab_xlsx.name
    assert pack.shab_xlsx.name == "260702 Rue A B 05 46 AVP - export SHAB maquette.xlsx"


def test_materiau_ratio_not_exploded(tmp_path, sources):
    # ARC matériau : le taux consolidé n'explose plus (bug 1053000 %).
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    doc = Document(str(pack.analyse_docx))
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text for c in row.cells]
            if cells and cells[0].startswith("Éléments sans matériau"):
                # source ARC absente dans cette fixture -> NOT_AVAILABLE (jamais 1053000 %)
                assert "%" not in cells[1] or float(cells[1].rstrip(" %")) < 100


def test_seuil_not_invented(tmp_path):
    from audit_bim.reporting.avp_i3f import NOT_AVAILABLE

    src_dir = tmp_path / "src"
    src_dir.mkdir()
    env = _wb(
        src_dir / "env.xlsx",
        {
            "TDB": [
                ["Composant", "Type", "Surface Solibri"],
                ["Mur", "ME", 100.0],
                [],
                [None, None, "ratio FAC/SHAB : ", 0.95],  # pas de « Seuil 3F 2026 »
            ]
        },
    )
    pack = write_avp_i3f_report_pack(
        None, tmp_path / "out", sources=AvpSourcePaths(enveloppe=env), export_pdf=False
    )
    doc = Document(str(pack.analyse_docx))
    seuil_value = None
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text for c in row.cells]
            if cells and cells[0].startswith("Seuil 3F 2026"):
                assert cells[0] == "Seuil 3F 2026"  # pas de « (≥ 0.9) » inventé
                seuil_value = cells[1]
    assert seuil_value == NOT_AVAILABLE  # verdict indisponible faute de seuil source


# ── Reproduction du détail + audit + mise en page (2e revue) ─────────────


def test_controle_detail_grid_reproduced(tmp_path, sources):
    # Les onglets de contrôle ne sont plus réduits à une ligne KPI :
    # la grille détaillée source est reproduite (ligne « 316 » présente).
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    wb = openpyxl.load_workbook(pack.controle_xlsx)
    ws = wb["Pièces Nommage"]
    vals = {c.value for row in ws.iter_rows() for c in row}
    assert 316 in vals  # détail reproduit, pas seulement l'agrégat
    assert ws.max_row > 3
    wb.close()


def test_export_sheet_names_preserved(tmp_path, sources):
    # P3 : proximité I3F — les onglets Enveloppe/Menuiseries gardent le
    # nom source (« TDB … »).
    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    wb = openpyxl.load_workbook(pack.enveloppe_xlsx)
    assert wb.sheetnames == ["TDB 2022 04.2"]
    wb.close()
    wb = openpyxl.load_workbook(pack.menuiseries_xlsx)
    assert wb.sheetnames == ["TDB 2022 05.1 - Fenêtres"]
    wb.close()


def _mini_audit_result():
    from audit_bim.audit.engine import AuditResult
    from audit_bim.audit.findings import ErrorType, Finding, Severity, Theme
    from audit_bim.extraction.model_data import ModelSnapshot
    from audit_bim.requirements.models import BIMPhase, RequirementsCatalog

    cat = RequirementsCatalog(
        cch_version="3.6",
        properties=[],
        naming_rules=[],
        storey_names=[],
        zone_specs=[],
        room_specs=[],
    )
    snap = ModelSnapshot(project={"name": "Tarare"}, model={"name": "MN.ifc"}).index()
    findings = [
        Finding(
            theme=Theme.QUANTITY,
            severity=Severity.MEDIUM,
            error_type=ErrorType.SPATIAL_MISSING_QUANTITY,
            element_uuid=f"q{i}",
            ifc_type="IfcSpace",
        )
        for i in range(3)
    ]
    return AuditResult(phase=BIMPhase.AVP, catalog=cat, snapshot=snap, findings=findings)


def test_consolidated_uses_real_audit_result(tmp_path, sources):
    # P2a : le consolidé restitue la synthèse d'audit (anomalies, quantités
    # manquantes) et ne l'ignore plus.
    pack = write_avp_i3f_report_pack(
        _mini_audit_result(), tmp_path / "out", sources=sources, export_pdf=False
    )
    doc = Document(str(pack.analyse_docx))
    txt = "\n".join(p.text for p in doc.paragraphs)
    assert "Audit BIMData automatisé de la maquette active" in txt
    found = any(
        c.text.startswith("Quantités manquantes") and r.cells[1].text == "3"
        for tbl in doc.tables
        for r in tbl.rows
        for c in r.cells[:1]
    )
    assert found


def test_consolidated_grille_is_landscape(tmp_path, sources):
    from docx.enum.section import WD_ORIENT

    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    doc = Document(str(pack.analyse_docx))
    assert any(s.orientation == WD_ORIENT.LANDSCAPE for s in doc.sections)


# ── 3e revue : LOW/INFO, métadonnées opérationnelles, onglet vide ────────


def _low_audit_result(n_low=3362):
    from audit_bim.audit.engine import AuditResult
    from audit_bim.audit.findings import ErrorType, Finding, Severity, Theme
    from audit_bim.extraction.model_data import ModelSnapshot
    from audit_bim.requirements.models import BIMPhase, RequirementsCatalog

    cat = RequirementsCatalog(
        cch_version="3.6",
        properties=[],
        naming_rules=[],
        storey_names=[],
        zone_specs=[],
        room_specs=[],
    )
    snap = ModelSnapshot(project={"name": "T"}, model={"name": "MN.ifc"}).index()
    findings = [
        Finding(
            theme=Theme.NAMING_SPACE,
            severity=Severity.LOW,
            error_type=ErrorType.NAMING_TOO_LONG,
            element_uuid=f"l{i}",
            ifc_type="IfcSpace",
        )
        for i in range(n_low)
    ]
    return AuditResult(phase=BIMPhase.AVP, catalog=cat, snapshot=snap, findings=findings)


def test_audit_synthese_includes_low_and_info(tmp_path, sources):
    pack = write_avp_i3f_report_pack(
        _low_audit_result(3362), tmp_path / "out", sources=sources, export_pdf=False
    )
    doc = Document(str(pack.analyse_docx))
    rows = {
        r.cells[0].text: r.cells[1].text
        for tbl in doc.tables
        for r in tbl.rows
        if len(r.cells) >= 2
    }
    assert rows.get("LOW") == "3362"  # la répartition se réconcilie (plus de LOW masqué)
    assert "INFO" in rows


def test_metadata_fill_usages_and_donnees(tmp_path, sources):
    pack = write_avp_i3f_report_pack(
        None,
        tmp_path / "out",
        sources=sources,
        usages_bim=["Usage 3F 1", "Usage 3F 2"],
        nombre_logements="24 logements",
        temoin_virtuel="Absent",
        date_controle="2026-01-30",
        auteur_controle="CdP BIM 3F",
        export_pdf=False,
    )
    doc = Document(str(pack.analyse_docx))
    txt = "\n".join(p.text for p in doc.paragraphs)
    labels = {c.text for tbl in doc.tables for r in tbl.rows for c in r.cells}
    assert "Usage 3F 1" in txt and "Usage 3F 2" in txt
    assert "24 logements" in labels and "Absent" in labels and "CdP BIM 3F" in labels


def test_metadata_absent_not_invented(tmp_path, sources):
    from audit_bim.reporting.avp_i3f import NOT_AVAILABLE

    pack = write_avp_i3f_report_pack(None, tmp_path / "out", sources=sources, export_pdf=False)
    doc = Document(str(pack.analyse_docx))
    txt = "\n".join(p.text for p in doc.paragraphs)
    # Usages absents -> NOT_AVAILABLE, jamais inventés.
    assert f"Usages BIM 3F : {NOT_AVAILABLE}" in txt


def test_auteur_controle_defaults_to_auditor(tmp_path, sources):
    """R4 P2 : sans ``auteur_controle`` explicite, on reprend ``auditor``
    (donnée fournie) plutôt que ``NOT_AVAILABLE``."""
    from audit_bim.reporting.avp_i3f import NOT_AVAILABLE

    pack = write_avp_i3f_report_pack(
        None,
        tmp_path / "out",
        sources=sources,
        auditor="AMO BIM BIMData",
        export_pdf=False,
    )
    doc = Document(str(pack.analyse_docx))
    labels = {c.text for tbl in doc.tables for r in tbl.rows for c in r.cells}
    assert "AMO BIM BIMData" in labels
    # « Auteur du contrôle » ne doit pas rester NOT_AVAILABLE.
    for tbl in doc.tables:
        for r in tbl.rows:
            cells = [c.text for c in r.cells]
            if cells and cells[0] == "Auteur du contrôle":
                assert cells[1] != NOT_AVAILABLE
                assert cells[1] == "AMO BIM BIMData"


def test_auteur_controle_explicit_wins(tmp_path, sources):
    pack = write_avp_i3f_report_pack(
        None,
        tmp_path / "out",
        sources=sources,
        auditor="AMO BIM BIMData",
        auteur_controle="CdP BIM 3F",
        export_pdf=False,
    )
    doc = Document(str(pack.analyse_docx))
    for tbl in doc.tables:
        for r in tbl.rows:
            cells = [c.text for c in r.cells]
            if cells and cells[0] == "Auteur du contrôle":
                assert cells[1] == "CdP BIM 3F"


def test_empty_source_sheet_preserved(tmp_path):
    src_dir = tmp_path / "src"
    src_dir.mkdir()
    zones = _wb(
        src_dir / "zones.xlsx",
        {
            "Feuil1": [],  # onglet vide (structure I3F stricte)
            "TDB 2022 01.3 - Export Zones": [
                ["Composant", "Nom Zone"],
                ["Zone", "0546L-1101"],
            ],
        },
    )
    pack = write_avp_i3f_report_pack(
        None, tmp_path / "out", sources=AvpSourcePaths(zones_espaces=zones), export_pdf=False
    )
    wb = openpyxl.load_workbook(pack.zones_espaces_xlsx)
    assert "Feuil1" in wb.sheetnames  # onglet vide préservé
    wb.close()


# ── PDF best-effort ──────────────────────────────────────────────────────


def test_pdf_export_none_when_engine_missing(tmp_path, monkeypatch):
    monkeypatch.setenv("AUDIT_BIM_SOFFICE", str(tmp_path / "no-such-soffice"))
    monkeypatch.setattr("shutil.which", lambda name: None)
    docx = tmp_path / "x.docx"
    Document().save(str(docx))
    assert docx_to_pdf(docx) is None
