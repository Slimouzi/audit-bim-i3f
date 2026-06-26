"""Tests de génération du rapport Word avec contexte enrichi.

On utilise ``docx.Document`` pour relire le fichier généré et chercher
les titres et textes des nouvelles sections.
"""

from __future__ import annotations

import pytest
from docx import Document

from audit_bim.audit.engine import AuditResult
from audit_bim.audit.findings import ErrorType, Finding, Severity, Theme
from audit_bim.extraction.model_data import ModelSnapshot
from audit_bim.reporting.context import build_report_context
from audit_bim.reporting.word_report import NOT_AVAILABLE, write_word_report
from audit_bim.requirements.models import BIMPhase, RequirementsCatalog

# ── Fixtures ────────────────────────────────────────────────────────────


def _minimal_catalog() -> RequirementsCatalog:
    return RequirementsCatalog(
        cch_version="3.6",
        cch_source_pdf="/tmp/cch.pdf",
        data_spec_source="/tmp/data.xlsx",
        naming_spec_source="/tmp/naming.xlsx",
        properties=[],
        naming_rules=[],
        storey_names=[],
        zone_specs=[],
        room_specs=[],
    )


def _result(
    *,
    project: dict | None = None,
    findings: list[Finding] | None = None,
    catalog: RequirementsCatalog | None = None,
) -> AuditResult:
    snap = ModelSnapshot(
        project=project or {"name": "Programme Test"},
        model={"name": "TEST.ifc"},
        sites=[{"uuid": "S1", "name": "Site Test", "type": "IfcSite"}],
        buildings=[{"uuid": "B1", "name": "Bât A", "type": "IfcBuilding"}],
        storeys=[],
        spaces=[],
        zones=[],
        elements=[],
    ).index()
    return AuditResult(
        phase=BIMPhase.PRO,
        catalog=catalog or _minimal_catalog(),
        snapshot=snap,
        findings=findings
        or [
            Finding(
                theme=Theme.NAMING_SPACE,
                severity=Severity.MEDIUM,
                error_type=ErrorType.NAMING_MISSING,
                element_uuid="X1",
                ifc_type="IfcSpace",
            )
        ],
    )


def _doc_text(path: str) -> str:
    """Concatène tous les paragraphes + textes de cellules d'un docx."""
    doc = Document(path)
    parts: list[str] = []
    for para in doc.paragraphs:
        parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ── Tests ───────────────────────────────────────────────────────────────


class TestWordReportContextSections:
    def test_all_new_sections_present(self, tmp_path):
        """Le rapport Word doit contenir les 9 sections numérotées du
        modèle de rapport de conformité + la mention "Informations non
        disponibles" (déclenchée par les données manquantes typiques
        d'un AuditResult minimal)."""
        out = tmp_path / "report.docx"
        write_word_report(_result(), output_path=out, auditor="Test AMO")
        assert out.exists()
        text = _doc_text(str(out))

        # Titres obligatoires (structure refondue 0.3)
        for heading in (
            "2. Synthèse exécutive",
            "3. Périmètre de l'audit",
            "4. Méthodologie",
            "5. Résultats globaux",
            "6. Résultats détaillés",
            "7. Liste des non-conformités",
            "8. Recommandations",
            "9. Conclusion",
            "10. Annexes",
        ):
            assert heading in text, f"Section manquante : {heading}"
        # Sous-sections 6.x du détail
        for sub in ("6.1 Structure", "6.3 Classification", "6.7 Détection des conflits"):
            assert sub in text, f"Sous-section manquante : {sub}"
        # Page de garde : titre du rapport + référence CCBIM
        assert "Rapport d'audit de conformité de la maquette numérique" in text
        assert "Référence du CCBIM utilisé" in text
        # Décision d'acceptation présente
        assert "Décision finale" in text
        # Le contexte d'un projet minimal déclenche au moins une mention
        # manquante (description / MOA / adresse / objectifs).
        assert (
            "Informations non disponibles" in text
            or "non disponible dans les documents fournis" in text
        )

    def test_no_hallucination_when_minimal_data(self, tmp_path):
        """Quand les données sont absentes, la mention
        ``NOT_AVAILABLE`` doit apparaître plutôt qu'une valeur inventée.
        """
        out = tmp_path / "report.docx"
        write_word_report(_result(project={"name": "X"}), output_path=out)
        text = _doc_text(str(out))
        # Mention de fallback explicite
        assert NOT_AVAILABLE in text

    def test_works_with_explicit_context(self, tmp_path):
        """Le caller peut pré-builder un contexte et le passer."""
        result = _result()
        ctx = build_report_context(result)
        out = tmp_path / "report.docx"
        write_word_report(result, output_path=out, context=ctx)
        assert out.exists()
        text = _doc_text(str(out))
        assert "3. Périmètre de l'audit" in text

    def test_does_not_raise_when_no_findings(self, tmp_path):
        """Génération sans findings ne doit pas planter (la section
        Détail saute)."""
        out = tmp_path / "report.docx"
        write_word_report(_result(findings=[]), output_path=out)
        assert out.exists()

    def test_does_not_raise_when_no_project_description(self, tmp_path):
        """Description projet absente → pas d'exception, fallback affiché."""
        out = tmp_path / "report.docx"
        write_word_report(
            _result(project={"name": "Projet anonyme"}),
            output_path=out,
        )
        assert out.exists()

    def test_controls_table_has_expected_themes(self, tmp_path):
        """La table des contrôles (Méthodologie) doit contenir au moins
        Classification, Nommage, Propriétés, Spatial."""
        out = tmp_path / "report.docx"
        write_word_report(_result(), output_path=out)
        text = _doc_text(str(out))
        assert "Classification IFC" in text
        # Au moins un nommage présent
        assert "Nommage" in text
        assert "Propriétés attendues" in text
        assert "Hiérarchie spatiale" in text

    def test_out_of_scope_controls_flagged_not_conforming(self, tmp_path):
        """Les familles de contrôles non couvertes (cohérence métier,
        détection de conflits) doivent être explicitement signalées comme
        hors périmètre — jamais présentées comme conformes."""
        out = tmp_path / "report.docx"
        write_word_report(_result(), output_path=out)
        text = _doc_text(str(out))
        assert "6.6 Cohérence métier" in text
        assert "6.7 Détection des conflits" in text
        assert "Contrôle non réalisé dans le périmètre" in text

    def test_decision_reflects_severity(self, tmp_path):
        """Une non-conformité critique force une décision « Refusée »."""
        result = _result(
            findings=[
                Finding(
                    theme=Theme.CLASSIFICATION,
                    severity=Severity.CRITICAL,
                    error_type=ErrorType.CLASSIFICATION_MISSING,
                    element_uuid="C1",
                    ifc_type="IfcWall",
                )
            ]
        )
        out = tmp_path / "report.docx"
        write_word_report(result, output_path=out)
        text = _doc_text(str(out))
        assert "Refusée" in text


class TestWordReportBackwardsCompat:
    """L'extension ne doit pas casser l'appel sans ``context``."""

    def test_legacy_call_without_context(self, tmp_path):
        out = tmp_path / "report.docx"
        # Signature historique : pas de paramètre context
        write_word_report(
            _result(),
            output_path=out,
            auditor="Legacy auditor",
            xlsx_annex_path=tmp_path / "annex.xlsx",
        )
        assert out.exists()

    def test_legacy_returns_path(self, tmp_path):
        out = tmp_path / "report.docx"
        returned = write_word_report(_result(), output_path=out)
        assert returned == out


@pytest.mark.parametrize(
    "phase",
    [BIMPhase.APS, BIMPhase.AVP, BIMPhase.PRO, BIMPhase.DCE, BIMPhase.EXE, BIMPhase.DOE],
)
def test_word_report_all_phases(tmp_path, phase):
    """Le rapport doit se générer pour toutes les phases sans exception."""
    snap = ModelSnapshot(
        project={"name": f"Test {phase.value}"},
        model={"name": "T.ifc"},
        sites=[],
        buildings=[],
        storeys=[],
        spaces=[],
        zones=[],
        elements=[],
    ).index()
    result = AuditResult(
        phase=phase,
        catalog=_minimal_catalog(),
        snapshot=snap,
        findings=[],
    )
    out = tmp_path / f"report_{phase.value}.docx"
    write_word_report(result, output_path=out)
    assert out.exists()
