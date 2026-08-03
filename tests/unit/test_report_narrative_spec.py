"""Narratif Word piloté par le profil : texte I3F figé, profil tiers propre.

PR C1 sort de `word_report.py` les phrases qui citent un référentiel, un maître
d'ouvrage ou un système de classification propriétaire. Le risque est double et
opposé : casser le livrable I3F d'un côté, laisser fuiter du I3F chez un autre
AMO de l'autre. Ces tests couvrent les deux sens.
"""

from __future__ import annotations

import ast
import subprocess
from dataclasses import replace
from pathlib import Path

import pytest
from docx import Document

import audit_bim.profiles.registry as reg
from audit_bim.audit.engine import AuditResult
from audit_bim.audit.findings import ErrorType, Finding, Severity, Theme
from audit_bim.extraction.model_data import ModelSnapshot
from audit_bim.profiles import get_profile
from audit_bim.profiles.models import (
    ClassificationNarrativeSpec,
    ReferenceFrameworkSpec,
    ReportNarrativeSpec,
)
from audit_bim.reporting import word_report
from audit_bim.requirements.models import BIMPhase

WORD_REPORT = Path(word_report.__file__)


class _Catalog:
    cch_version = "3.6"
    cch_source_pdf = "cch.pdf"
    data_spec_source = "spec.xlsx"
    naming_spec_source = "nom.xlsx"
    properties: list = []
    naming_rules: list = []
    storey_names: list = []
    zone_specs: list = []
    room_specs: list = []


def _result() -> AuditResult:
    snap = ModelSnapshot()
    snap.project = {"name": "P"}
    snap.model = {"name": "M"}
    themes = [
        Theme.CLASSIFICATION,
        Theme.NAMING_ZONE,
        Theme.NAMING_SPACE,
        Theme.SPATIAL_HIERARCHY,
        Theme.NAMING_SITE_BAT_ETAGE,
        Theme.QUANTITY,
        Theme.PROPERTY_MISSING,
    ]
    findings = [
        Finding(
            element_uuid=f"u{i}",
            ifc_type="IfcWall",
            name=f"W{i}",
            theme=th,
            error_type=ErrorType.CLASSIFICATION_MISSING,
            severity=Severity.HIGH,
            expected="x",
            actual=None,
        )
        for i, th in enumerate(themes)
    ]
    return AuditResult(snapshot=snap, catalog=_Catalog(), phase=BIMPhase.PRO, findings=findings)


def _render(tmp_path, monkeypatch, *, profile_id=None) -> list[str]:
    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    out = tmp_path / "r.docx"
    word_report.write_word_report(_result(), out, profile_id=profile_id)
    return [p.text for p in Document(str(out)).paragraphs]


# ── 1. Le narratif ne vit plus dans le module de rendu ────────────────


def test_theme_hints_is_no_longer_a_local_constant():
    """`_THEME_HINTS` doit avoir disparu de word_report.py."""
    tree = ast.parse(WORD_REPORT.read_text(encoding="utf-8"))
    names = {
        t.id
        for node in tree.body
        if isinstance(node, (ast.Assign, ast.AnnAssign))
        for t in (node.targets if isinstance(node, ast.Assign) else [node.target])
        if isinstance(t, ast.Name)
    }
    assert "_THEME_HINTS" not in names
    assert not hasattr(word_report, "_THEME_HINTS")


def test_no_client_string_printed_from_word_report():
    """Aucune chaîne client dans un littéral vif de word_report.py.

    Contrôle par AST : on ignore docstrings et commentaires, on ne regarde que
    les chaînes réellement évaluées.
    """
    src = WORD_REPORT.read_text(encoding="utf-8")
    tree = ast.parse(src)
    docstrings = set()
    for node in ast.walk(tree):
        body = getattr(node, "body", None)
        if isinstance(node, (ast.Module, ast.ClassDef, ast.FunctionDef)) and body:
            first = body[0]
            if isinstance(first, ast.Expr) and isinstance(first.value, ast.Constant):
                docstrings.add(id(first.value))

    offenders = [
        (node.lineno, node.value[:70])
        for node in ast.walk(tree)
        if isinstance(node, ast.Constant)
        and isinstance(node.value, str)
        and id(node) not in docstrings
        and any(
            term in node.value
            for term in ("I3F", "CCH", "CCBIM", "table 3F", "UniFormat", "Omniclass")
        )
    ]
    assert not offenders, f"chaînes client encore imprimées depuis word_report.py : {offenders}"


def test_profile_carries_the_full_theme_hints_table():
    hints = get_profile("i3f").report_narrative.theme_hints
    assert len(hints) == 9
    assert hints["Nommage Zone"] == (
        "reprendre le nommage des zones (codification I3F, CCH chap. 6.3)"
    )


# ── 2. I3F : rendu byte-identique ─────────────────────────────────────

#: Paragraphes I3F qui devaient survivre au découplage, à l'octet près.
I3F_EXPECTED_PARAGRAPHS = [
    "Présence et cohérence de la classification IFC (UniFormat II par défaut ; "
    "Omniclass / CCI / table interne 3F selon le référentiel).",
    "Contrôle du nommage des objets, niveaux, zones et espaces selon les listes "
    "fermées et la codification I3F (CCH chap. 6.3).",
    "• Référentiel CCH I3F : documents transmis par la maîtrise d'ouvrage "
    "(Cahier des annexes, annexe Spécifications, annexe Nommage).",
]


def test_i3f_narrative_paragraphs_are_byte_identical(tmp_path, monkeypatch):
    rendered = _render(tmp_path, monkeypatch)
    for expected in I3F_EXPECTED_PARAGRAPHS:
        assert expected in rendered, f"paragraphe I3F perdu ou modifié : {expected[:60]!r}"


def test_i3f_theme_hints_still_reach_the_recommendations(tmp_path, monkeypatch):
    joined = "\n".join(_render(tmp_path, monkeypatch))
    assert "codification I3F, CCH chap. 6.3" in joined
    assert "UniFormat / Omniclass / table 3F" in joined


def test_i3f_cover_reference_label_unchanged(tmp_path, monkeypatch):
    tables_text = []
    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    out = tmp_path / "r.docx"
    word_report.write_word_report(_result(), out)
    for table in Document(str(out)).tables:
        for row in table.rows:
            tables_text.extend(c.text for c in row.cells)
    assert any("Référence du CCBIM utilisé" in t for t in tables_text)


def test_low_conformity_recommendation_is_byte_identical(tmp_path, monkeypatch):
    joined = "\n".join(_render(tmp_path, monkeypatch))
    assert (
        "Ré-itérer un audit après reprise : l'écart au CCH est important — "
        "prévoir une revue conjointe MOA / MOE avant la phase suivante."
    ) in joined


# ── 3. Profil tiers : aucune fuite I3F dans le Word ───────────────────

FORBIDDEN_IN_THIRD_PARTY = (
    "I3F",
    "CCH BIM I3F",
    "table 3F",
    "table interne 3F",
    "codification I3F",
)


@pytest.fixture
def third_party_profile(monkeypatch):
    """Profil AMO complet et distinct, branché sur le registre."""
    tiers = replace(
        reg._BIM_IN_MOTION_PROFILE,
        reference_framework=ReferenceFrameworkSpec(
            name="Référentiel AMO BIM in Motion",
            short_name="Référentiel",
            long_name="Référentiel AMO BIM in Motion",
        ),
        report_narrative=ReportNarrativeSpec(
            theme_hints={"Nommage Zone": "reprendre le nommage des zones"},
            classification_intro="Présence et cohérence de la classification IFC.",
            naming_intro="Contrôle du nommage des objets, niveaux, zones et espaces.",
            reference_documents_line="• Référentiel client : documents transmis.",
            cover_reference_label="Référence contractuelle",
            applied_reference_label="Référentiel appliqué",
            low_conformity_recommendation="Ré-itérer un audit après reprise.",
        ),
        classification_narrative=ClassificationNarrativeSpec(
            default_system="UniFormat II",
            known_systems=("UniFormat",),
            proprietary_systems=(),
            proprietary_label="",
        ),
    )
    monkeypatch.setattr(reg, "_PROFILES", (reg._I3F_PROFILE, tiers))
    return "bim_in_motion"


def test_third_party_word_has_no_i3f_trace(tmp_path, monkeypatch, third_party_profile):
    joined = "\n".join(_render(tmp_path, monkeypatch, profile_id=third_party_profile))
    for term in FORBIDDEN_IN_THIRD_PARTY:
        assert term not in joined, f"{term!r} imprimé dans le Word d'un AMO tiers"


def test_third_party_prints_its_own_narrative(tmp_path, monkeypatch, third_party_profile):
    joined = "\n".join(_render(tmp_path, monkeypatch, profile_id=third_party_profile))
    assert "Référentiel AMO BIM in Motion" in joined
    assert "• Référentiel client : documents transmis." in joined


def test_profile_without_narrative_degrades_without_leaking(tmp_path, monkeypatch):
    """`bim_in_motion` nu : replis neutres, jamais les phrases d'I3F."""
    joined = "\n".join(_render(tmp_path, monkeypatch, profile_id="bim_in_motion"))
    for term in FORBIDDEN_IN_THIRD_PARTY:
        assert term not in joined, f"{term!r} hérité par un profil sans narratif"


# ── 4. Composition de l'intro classification ──────────────────────────


def test_classification_intro_declared_by_a_profile_is_actually_printed(tmp_path, monkeypatch):
    """Un champ de profil déclaré mais jamais lu est une fausse commande.

    `classification_intro` était renseigné dans le registre et ignoré par le
    rendu, qui recomposait la phrase depuis `ClassificationNarrativeSpec`. Un
    profil tiers aurait pu définir sa phrase sans qu'elle s'imprime jamais —
    le pire cas, parce que rien n'échoue.
    """
    phrase = "Phrase classification BIM in Motion"
    tiers = replace(
        reg._BIM_IN_MOTION_PROFILE,
        reference_framework=ReferenceFrameworkSpec(
            name="Référentiel AMO BIM in Motion",
            short_name="Référentiel",
            long_name="Référentiel AMO BIM in Motion",
        ),
        report_narrative=ReportNarrativeSpec(
            theme_hints={"Nommage Zone": "reprendre le nommage des zones"},
            classification_intro=phrase,
            naming_intro="Contrôle du nommage des objets, niveaux, zones et espaces.",
            reference_documents_line="• Référentiel client : documents transmis.",
            cover_reference_label="Référence contractuelle",
            applied_reference_label="Référentiel appliqué",
            low_conformity_recommendation="Ré-itérer un audit après reprise.",
        ),
        classification_narrative=ClassificationNarrativeSpec(
            default_system="UniFormat II",
            known_systems=("UniFormat",),
            proprietary_systems=(),
            proprietary_label="",
        ),
    )
    monkeypatch.setattr(reg, "_PROFILES", (reg._I3F_PROFILE, tiers))

    rendered = _render(tmp_path, monkeypatch, profile_id="bim_in_motion")
    assert phrase in rendered, "la phrase du profil tiers n'est pas imprimée"

    joined = "\n".join(rendered)
    for term in (*FORBIDDEN_IN_THIRD_PARTY, "Omniclass", "CCI"):
        assert term not in joined, f"{term!r} a fuité dans le Word d'un AMO tiers"


def test_i3f_classification_intro_comes_from_the_narrative_spec(tmp_path, monkeypatch):
    """Côté I3F, la phrase imprimée est bien celle du profil, à l'octet près."""
    declared = get_profile("i3f").report_narrative.classification_intro
    assert declared == I3F_EXPECTED_PARAGRAPHS[0]
    assert declared in _render(tmp_path, monkeypatch)


def test_profile_without_narrative_uses_the_neutral_classification_fallback(tmp_path, monkeypatch):
    """Sans narratif, repli neutre — jamais la phrase d'un autre profil."""
    rendered = _render(tmp_path, monkeypatch, profile_id="bim_in_motion")
    assert "Présence et cohérence de la classification IFC." in rendered
    assert I3F_EXPECTED_PARAGRAPHS[0] not in rendered


# ── 5. Quels champs de profil sont RÉELLEMENT lus ─────────────────────


def test_classification_spec_consumed_fields_are_pinned():
    """Fige ce qui est branché, pour que le déclaratif reste visible comme tel.

    Après C1, seul `default_system` est consommé (`_build_controls_performed`).
    `known_systems`, `proprietary_systems` et `proprietary_label` attendent C2 et
    le futur `bim-classifier`. Ce test échouera le jour où on les branchera —
    c'est voulu : il faudra alors ajouter le test de rendu correspondant, au lieu
    de croire la commande active parce que le champ existe.
    """
    import audit_bim.reporting.context as ctx_mod
    import audit_bim.reporting.word_report as wr_mod

    sources = Path(ctx_mod.__file__).read_text(encoding="utf-8") + Path(wr_mod.__file__).read_text(
        encoding="utf-8"
    )
    assert "classification.default_system" in sources or "default_system" in sources
    for dormant in ("known_systems", "proprietary_systems", "proprietary_label"):
        assert dormant not in sources, (
            f"{dormant!r} est désormais lu par le rendu : ajouter un test de "
            "rendu tiers et retirer ce champ de la liste dormante"
        )


def _narrative_fields_read_by(module) -> set[str]:
    """Champs de narratif réellement LUS par un module, détectés par AST.

    Une simple recherche de sous-chaîne est piégeuse ici : avant l'amend,
    `word_report.py` contenait bien « classification_intro » — c'était le nom
    d'une fonction locale `_classification_intro`, pas une lecture du champ. Le
    contrôle serait passé sur le bug qu'il visait.

    On ne retient donc que deux formes de lecture réelle :
    `_narrative_text(profile_id, "<champ>")` et `<qqch>.<champ>`.
    """
    tree = ast.parse(Path(module.__file__).read_text(encoding="utf-8"))
    read: set[str] = set()
    for node in ast.walk(tree):
        if (
            isinstance(node, ast.Call)
            and isinstance(node.func, ast.Name)
            and node.func.id == "_narrative_text"
            and len(node.args) >= 2
            and isinstance(node.args[1], ast.Constant)
        ):
            read.add(node.args[1].value)
        elif isinstance(node, ast.Attribute):
            read.add(node.attr)
    return read


def test_narrative_spec_fields_are_all_consumed():
    """TOUS les champs de ReportNarrativeSpec doivent être réellement lus.

    C'est le contrôle qui attrape un champ de profil mort — le cas
    `classification_intro`, déclaré, renseigné, et ignoré par le rendu. Rien
    n'échouait : un profil tiers aurait défini sa phrase pour rien.
    """
    import audit_bim.reporting.word_report as wr_mod

    declared = set(get_profile("i3f").report_narrative.to_dict())
    unused = sorted(declared - _narrative_fields_read_by(wr_mod))
    assert not unused, f"champs de ReportNarrativeSpec déclarés mais jamais lus : {unused}"


def test_the_consumption_guard_would_have_caught_the_dead_field():
    """Preuve de non-vacuité, contre la version d'avant l'amend.

    On rejoue la détection sur le `word_report.py` du commit précédent : le
    garde-fou doit y voir `classification_intro` non lu.
    """
    previous = subprocess.run(
        ["git", "show", "HEAD:audit_bim/reporting/word_report.py"],
        capture_output=True,
        text=True,
        cwd=Path(__file__).resolve().parents[2],
    ).stdout
    if not previous.strip():
        pytest.skip("historique Git indisponible")

    tree = ast.parse(previous)
    read: set[str] = set()
    for node in ast.walk(tree):
        if (
            isinstance(node, ast.Call)
            and isinstance(node.func, ast.Name)
            and node.func.id == "_narrative_text"
            and len(node.args) >= 2
            and isinstance(node.args[1], ast.Constant)
        ):
            read.add(node.args[1].value)
        elif isinstance(node, ast.Attribute):
            read.add(node.attr)
    assert "classification_intro" not in read, (
        "le garde-fou ne détecte pas le champ mort d'avant l'amend : il est vacueux"
    )
