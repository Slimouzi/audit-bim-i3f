"""Adoption des primitives `bim_reporting.sections` — et absence de squelette.

Deux exigences opposées :

1. le Word I3F ne change pas — les primitives rendent ce que rendaient les
   helpers locaux ;
2. un profil tiers assemble son document **sans** l'orchestrateur I3F. C'est le
   seul test qui prouve que le socle fournit des briques et non un rapport.
"""

from __future__ import annotations

import ast
import subprocess
import sys
import tempfile
import textwrap
from pathlib import Path

from docx import Document

from audit_bim.audit.engine import AuditResult
from audit_bim.audit.findings import ErrorType, Finding, Severity, Theme
from audit_bim.extraction.model_data import ModelSnapshot
from audit_bim.reporting import word_report
from audit_bim.requirements.models import BIMPhase

WORD_REPORT = Path(word_report.__file__)


class _Catalog:
    cch_version = "3.6"
    cch_source_pdf = "cch.pdf"
    data_spec_source = "spec.xlsx"
    naming_spec_source = "nom.xlsx"
    properties: list = []
    naming_rules: list = []
    storey_names: list = []
    zone_specs: list = []
    room_specs: list = []


def _result(themes=None) -> AuditResult:
    snap = ModelSnapshot()
    snap.project = {"name": "P"}
    snap.model = {"name": "M"}
    themes = themes or [Theme.CLASSIFICATION, Theme.QUANTITY, Theme.NAMING_ZONE]
    findings = [
        Finding(
            element_uuid=f"u{i}",
            ifc_type="IfcWall",
            name=f"W{i}",
            theme=th,
            error_type=ErrorType.CLASSIFICATION_MISSING,
            severity=Severity.HIGH,
            expected=["a", "b"],
            actual=None,
            ref_cch="6.3",
            recommended_action="corriger",
        )
        for i, th in enumerate(themes)
    ]
    return AuditResult(snapshot=snap, catalog=_Catalog(), phase=BIMPhase.PRO, findings=findings)


def _render(tmp_path, monkeypatch) -> Document:
    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    # La résolution du logo dépend de l'emplacement du dépôt : on la fige pour
    # que le test ne dépende pas de la présence d'un brand kit sur la machine.
    monkeypatch.setattr(word_report, "find_logo", lambda variant="light": None)
    out = tmp_path / "r.docx"
    word_report.write_word_report(_result(), out)
    return Document(str(out))


# ── 1. Le socle rend, l'orchestration reste ici ───────────────────────


def test_local_render_helpers_are_gone():
    """Les helpers remplacés par le socle ne doivent plus exister ici."""
    for name in ("_header_row", "_findings_table", "_write_cover_page"):
        assert not hasattr(word_report, name), f"{name} aurait dû partir dans le socle"


def test_orchestration_stays_in_this_module():
    """L'ordre des dix sections et les décisions éditoriales restent I3F."""
    for name in (
        "write_word_report",
        "_decision",
        "_domain_status",
        "DOMAINS",
        "GRAVITY_FR",
        "_STATUS_LABEL",
        "MAX_FINDINGS_PER_THEME",
        "MAX_NONCONFORMITIES",
    ):
        assert hasattr(word_report, name), f"{name} ne doit PAS partir dans le socle"


def test_decision_thresholds_are_unchanged():
    """Seuils contractuels : 90 % sans bloquant, 70 % sans critique."""
    src = ast.parse(WORD_REPORT.read_text(encoding="utf-8"))
    fn = next(n for n in src.body if isinstance(n, ast.FunctionDef) and n.name == "_decision")
    numbers = {
        n.value
        for n in ast.walk(fn)
        if isinstance(n, ast.Constant) and isinstance(n.value, (int, float))
    }
    assert {90, 70} <= numbers


# ── 2. Le Word I3F ne change pas ──────────────────────────────────────


def test_report_still_renders_expected_shape(tmp_path, monkeypatch):
    doc = _render(tmp_path, monkeypatch)
    texts = [p.text for p in doc.paragraphs]
    assert any(t.startswith("1.") or "Synthèse" in t for t in texts)
    assert any("10. Annexes" in t for t in texts)
    assert len(doc.tables) >= 10


def test_cover_labels_are_assembled_client_side(tmp_path, monkeypatch):
    """Les intitulés de couverture viennent d'ici, pas du socle."""
    doc = _render(tmp_path, monkeypatch)
    blob = doc.element.xml
    for label in ("Projet", "Maquette auditée", "Version", "Date", "Auteur"):
        assert label in blob
    assert "Rapport d'audit de conformité de la maquette numérique" in blob


def test_severity_column_is_still_colored(tmp_path, monkeypatch):
    doc = _render(tmp_path, monkeypatch)
    assert any("DC3545" in t._tbl.xml or "8B0000" in t._tbl.xml for t in doc.tables)


def test_empty_quantity_note_is_really_italic(tmp_path, monkeypatch):
    """Correction assumée d'un no-op : `Paragraph.italic` n'écrivait rien."""
    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    monkeypatch.setattr(word_report, "find_logo", lambda variant="light": None)
    out = tmp_path / "r.docx"
    # Aucun finding de quantité -> la mention « aucune anomalie » est rendue.
    word_report.write_word_report(_result(themes=[Theme.CLASSIFICATION]), out)
    doc = Document(str(out))
    note = next(p for p in doc.paragraphs if "Aucune anomalie de quantité détectée." in p.text)
    assert [r.italic for r in note.runs] == [True]
    assert "<w:i/>" in note._p.xml


# ── 3. Non-squelette : un tiers assemble sans l'orchestrateur I3F ─────

THIRD_PARTY_SCRIPT = textwrap.dedent(
    """
    import sys
    from docx import Document
    from bim_reporting.sections import bullet_list, cover_page, data_table, document_base

    doc = Document()
    document_base(doc)
    cover_page(doc, title="Rapport AMO tiers", meta_rows=[("Auteur", "X")])
    data_table(doc, ["Lot", "État"], [["A", "OK"]])
    bullet_list(doc, ["Point 1"])
    doc.save(sys.argv[1])

    leaked = [m for m in sys.modules if m.startswith("audit_bim")]
    if leaked:
        raise SystemExit("ORCHESTRATEUR I3F IMPORTÉ : " + ", ".join(sorted(leaked)))
    print("OK")
    """
)


def test_third_party_assembles_a_document_without_the_i3f_orchestrator(tmp_path):
    """Le test décisif : trois blocs, aucun import d'`audit_bim`.

    Exécuté dans un interpréteur SÉPARÉ : dans le processus de test, `audit_bim`
    est déjà importé par les autres cas, donc la vérification y serait vide de
    sens.
    """
    script = tmp_path / "tiers.py"
    script.write_text(THIRD_PARTY_SCRIPT, encoding="utf-8")
    out = tmp_path / "tiers.docx"
    proc = subprocess.run(
        [sys.executable, str(script), str(out)],
        capture_output=True,
        text=True,
        cwd=tempfile.gettempdir(),  # hors du dépôt : `audit_bim` non importable par accident
    )
    assert proc.returncode == 0, f"{proc.stdout}\n{proc.stderr}"
    assert "OK" in proc.stdout
    assert out.is_file()

    doc = Document(str(out))
    assert "Rapport AMO tiers" in doc.element.xml
    assert any("Point 1" in p.text for p in doc.paragraphs)


def test_the_leak_check_of_that_script_is_not_vacuous(tmp_path):
    """Le script tiers doit savoir échouer si l'orchestrateur est importé."""
    script = tmp_path / "fuite.py"
    script.write_text(
        "import audit_bim.reporting.word_report  # noqa: F401\n" + THIRD_PARTY_SCRIPT,
        encoding="utf-8",
    )
    proc = subprocess.run(
        [sys.executable, str(script), str(tmp_path / "x.docx")],
        capture_output=True,
        text=True,
        cwd=str(Path(word_report.__file__).resolve().parents[3]),
    )
    assert proc.returncode != 0
    assert "ORCHESTRATEUR I3F IMPORTÉ" in (proc.stdout + proc.stderr)


def test_word_report_does_not_import_a_report_builder():
    """Aucune API du socle en « report » n'est importée ici."""
    tree = ast.parse(WORD_REPORT.read_text(encoding="utf-8"))
    for node in ast.walk(tree):
        if isinstance(node, ast.ImportFrom) and (node.module or "").startswith("bim_reporting"):
            for alias in node.names:
                assert "report" not in alias.name.lower(), (
                    f"{alias.name} : le socle ne doit exposer aucun assembleur"
                )
