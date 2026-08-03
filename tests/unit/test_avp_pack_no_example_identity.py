"""Aucun livrable client ne peut porter l'identité d'un projet d'exemple.

« Tarare 0546L » est le projet de **référence MOA** : il a sa place dans les
fixtures, la documentation et les comparaisons de mise en forme — jamais dans
le nom d'un fichier remis au client. Le classeur de contrôle MOA étant
auto-découvert dans les documents maître d'ouvrage, son entête nommait les
livrables d'après ce projet-là, quelle que soit la maquette auditée.

Ces tests verrouillent la règle : l'identité vient **exclusivement des paramètres
explicites**, sinon **on demande sans générer**. Le repli sur le contexte du
modèle actif a été retiré — il a livré un pack « MCP_Audit » — et la maquette ne
fournit plus qu'une *suggestion* dans la question posée.
"""

from __future__ import annotations

from pathlib import Path

import openpyxl
import pytest

from audit_bim.extraction.model_data import ModelSnapshot
from audit_bim.mcp import server as mcp_server
from audit_bim.mcp.session import _Session, current_session

EXEMPLE_MOA = ("Tarare", "0546L")


@pytest.fixture
def session(tmp_path, monkeypatch):
    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    monkeypatch.delenv("AUDIT_INPUT_DIR", raising=False)
    sess = _Session()
    token = current_session.set(sess)
    try:
        yield sess, tmp_path
    finally:
        current_session.reset(token)


def _moa_template(path, *, projet="Tarare", esi="0546L", phase="AVP"):
    """Classeur de contrôle du projet de RÉFÉRENCE MOA (Tarare 0546L)."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Grille de contrôle"
    r = 1
    for label, val in (("Projet", projet), ("ESI", esi), ("Phase", phase)):
        ws.cell(r, 2, label)
        ws.cell(r, 3, val)
        r += 1
    ws.cell(r, 1, "CODE 3F")
    ws.cell(r, 2, "POINTS DE CONTROLE")
    wb.save(path)
    return str(path)


def _snapshot(sess, project_name="I3F"):
    sess.snapshot = ModelSnapshot(
        project={"name": project_name}, model={"name": "DIEPPE-7427L.ifc"}
    ).index()


# ── sans identité : on demande, on ne génère pas ───────────────────────


def test_missing_identity_asks_and_generates_nothing(session):
    sess, tmp_path = session
    sess.snapshot = ModelSnapshot(model={"name": "M.ifc"}).index()  # aucun nom projet

    res = mcp_server.generate_avp_i3f_pack(auditor="AMO BIM", export_pdf=False)

    assert res["status"] == "needs_context"
    assert {"project_name", "project_code"} <= set(res["missing"])
    keys = {q["key"] for q in res["questions"]}
    assert {"project_name", "project_code"} <= keys
    # Rien n'a été écrit.
    assert not list(Path(tmp_path).glob("**/*.xlsx"))
    assert not list(Path(tmp_path).glob("**/*.docx"))


def test_identity_cannot_be_bypassed_by_confirm_context(session):
    """``confirm_context`` couvre le contexte documentaire, pas l'identité."""
    sess, tmp_path = session
    sess.snapshot = ModelSnapshot(model={"name": "M.ifc"}).index()

    res = mcp_server.generate_avp_i3f_pack(confirm_context=True, export_pdf=False)

    assert res["status"] == "needs_context"
    assert "project_code" in res["missing"]
    assert "OBLIGATOIRE" in res["next_step"]
    assert not list(Path(tmp_path).glob("**/*.xlsx"))


def test_moa_template_identity_never_reaches_the_deliverables(session):
    """Le template MOA fournit la MISE EN FORME, jamais l'identité projet."""
    sess, tmp_path = session
    _snapshot(sess)
    ctrl = _moa_template(tmp_path / "controle maquettes.xlsx")

    res = mcp_server.generate_avp_i3f_pack(controle_xlsx=ctrl, auditor="AMO BIM", export_pdf=False)

    assert res["status"] == "needs_context"  # le code ESI manque toujours
    assert not list(Path(tmp_path).glob("**/*Tarare*"))
    assert not list(Path(tmp_path).glob("**/*0546L*"))


# ── avec identité explicite : le vrai projet, partout ──────────────────


def test_explicit_identity_names_every_deliverable(session):
    sess, tmp_path = session
    _snapshot(sess)
    ctrl = _moa_template(tmp_path / "controle maquettes.xlsx")

    res = mcp_server.generate_avp_i3f_pack(
        controle_xlsx=ctrl,
        project_name="Dieppe",
        project_code="7427L",
        phase="AVP",
        auditor="CdP BIM 3F",
        export_pdf=False,
    )

    assert res.get("status") != "needs_context", res
    assert res["project_name"] == "Dieppe"
    assert res["project_code"] == "7427L"

    # Le préfixe est la date de GÉNÉRATION : la calculer, ne pas la coder en
    # dur — un littéral ne passerait que le jour où le test a été écrit.
    from datetime import datetime

    attendu = f"{datetime.now().strftime('%y%m%d')} Dieppe 7427L AVP - "
    noms = [Path(p).name for p in res["paths"]]
    assert noms, "le pack doit produire des livrables"
    for nom in noms:
        assert nom.startswith(attendu), nom


def test_no_example_identity_anywhere_in_the_generated_pack(session):
    """Ni dans les noms de fichiers, ni dans le contenu des livrables."""
    sess, tmp_path = session
    _snapshot(sess)
    ctrl = _moa_template(tmp_path / "controle maquettes.xlsx")

    res = mcp_server.generate_avp_i3f_pack(
        controle_xlsx=ctrl,
        project_name="Dieppe",
        project_code="7427L",
        phase="AVP",
        auditor="CdP BIM 3F",
        export_pdf=False,
    )
    assert res.get("status") != "needs_context", res

    for chemin in res["paths"]:
        p = Path(chemin)
        for exemple in EXEMPLE_MOA:
            assert exemple not in p.name, f"{exemple} dans le nom : {p.name}"
        if p.suffix == ".xlsx":
            wb = openpyxl.load_workbook(p, data_only=True)
            texte = "\n".join(
                str(c)
                for ws in wb.worksheets
                for row in ws.iter_rows(values_only=True)
                for c in row
                if c is not None
            )
            wb.close()
            for exemple in EXEMPLE_MOA:
                assert exemple not in texte, f"{exemple} dans le contenu de {p.name}"


def test_bimdata_project_name_never_supplies_the_identity(session):
    """Le nom du projet BIMData n'est PLUS une source d'identité — même crédible.

    Il l'a été, et un pack est sorti sous « MCP_Audit ». La règle ne peut pas être
    « on accepte le nom BIMData s'il a l'air d'un chantier » : ce qui a l'air d'un
    nom de chantier reste le nom d'un espace de travail, choisi par celui qui l'a
    créé, et personne ne le relit avant qu'il n'atterrisse sur un livrable client.
    On demande, y compris quand le repli aurait donné la bonne réponse.
    """
    sess, tmp_path = session
    _snapshot(sess, project_name="Dieppe")

    res = mcp_server.generate_avp_i3f_pack(
        project_code="7427L", phase="AVP", auditor="CdP BIM 3F", export_pdf=False
    )

    assert res["status"] == "needs_context"
    assert res["missing"] == ["project_name"]
    assert not list(tmp_path.glob("avp_pack_*"))


# ── auteur du contrôle : demandé, jamais inventé ───────────────────────


def test_missing_auditor_is_asked_not_defaulted(session):
    sess, tmp_path = session
    _snapshot(sess, project_name="Dieppe")

    res = mcp_server.generate_avp_i3f_pack(project_code="7427L", phase="AVP", export_pdf=False)

    assert res["status"] == "needs_context"
    # Clé alignée sur le PARAMÈTRE du tool : une clé sans paramètre
    # correspondant guiderait vers un appel MCP invalide.
    assert "auditor_name" in res["missing"]
    q = next(q for q in res["questions"] if q["key"] == "auditor_name")
    assert set(q["accepted_aliases"]) == {"auteur_controle", "auditor"}
    assert not list(Path(tmp_path).glob("**/*.xlsx"))


def test_word_report_has_no_hardcoded_auditor_name():
    """``generate_word_report`` ne porte plus de nom d'auteur codé en dur."""
    import inspect

    from audit_bim.mcp import tools_reporting

    sig = inspect.signature(tools_reporting.generate_word_report)
    assert sig.parameters["auditor"].default is None
    assert sig.parameters["auditor_name"].default is None


def test_auditor_name_is_accepted_and_appears_in_the_pack(session):
    """Le paramètre que le prompt demande d'employer doit exister — et servir.

    Le prompt guide Claude vers ``auditor_name`` : si le tool ne l'expose pas,
    l'appel MCP échoue avec un ``TypeError`` côté serveur. Ce test appelle donc
    le tool exactement comme le prompt l'indique.
    """
    sess, tmp_path = session
    _snapshot(sess, project_name="Dieppe")

    res = mcp_server.generate_avp_i3f_pack(
        project_name="Dieppe",
        project_code="7427L",
        phase="AVP",
        auditor_name="Stanislas Limouzi",
        export_pdf=False,
    )

    assert res.get("status") != "needs_context", res
    docx = next(p for p in res["paths"] if p.endswith(".docx"))
    from docx import Document

    texte = "\n".join(p.text for p in Document(docx).paragraphs)
    texte += "\n" + "\n".join(
        c.text for t in Document(docx).tables for r in t.rows for c in r.cells
    )
    assert "Stanislas Limouzi" in texte


@pytest.mark.parametrize("param", ["auditor_name", "auteur_controle", "auditor"])
def test_every_documented_auditor_param_is_accepted(session, param):
    """Les trois noms coexistent : le nouveau, le métier I3F et l'historique."""
    sess, tmp_path = session
    _snapshot(sess, project_name="Dieppe")

    res = mcp_server.generate_avp_i3f_pack(
        project_name="Dieppe",
        project_code="7427L",
        phase="AVP",
        export_pdf=False,
        **{param: "CdP BIM 3F"},
    )

    assert res.get("status") != "needs_context", res


def test_auditor_name_wins_over_the_compat_aliases(session):
    sess, tmp_path = session
    _snapshot(sess, project_name="Dieppe")

    res = mcp_server.generate_avp_i3f_pack(
        project_name="Dieppe",
        project_code="7427L",
        phase="AVP",
        auditor_name="Nom retenu",
        auteur_controle="Nom metier",
        auditor="Nom legacy",
        export_pdf=False,
    )
    assert res.get("status") != "needs_context", res

    from docx import Document

    docx = next(p for p in res["paths"] if p.endswith(".docx"))
    texte = "\n".join(p.text for p in Document(docx).paragraphs)
    texte += "\n" + "\n".join(
        c.text for t in Document(docx).tables for r in t.rows for c in r.cells
    )
    assert "Nom retenu" in texte
    assert "Nom legacy" not in texte


def test_prompt_only_names_parameters_the_tool_exposes():
    """Garde-fou : le prompt ne doit citer que des paramètres réels.

    Un prompt qui nomme un paramètre inexistant guide vers un appel MCP
    invalide (``TypeError`` côté serveur) — l'UX devient pire que son absence.
    On vérifie la section qui instruit la génération du pack, où les noms cités
    sont des arguments à passer.
    """
    import inspect
    import re

    from audit_bim.mcp import prompts, tools_reporting

    texte = prompts.AMO_BIM_I3F_PROMPT
    debut = texte.index("## Identité projet et auteur du contrôle")
    section = texte[debut:]
    fin = section.find("\n## ", 1)
    if fin != -1:
        section = section[:fin]

    exposes = set(inspect.signature(tools_reporting.generate_avp_i3f_pack).parameters)
    cites = set(re.findall(r"`([a-z_]+)`", section)) | set(re.findall(r"`([a-z_]+)=", section))
    # Noms de tools et statuts cités dans la section : hors périmètre.
    cites -= {"generate_avp_i3f_pack", "questions", "suggestion", "status"}

    assert "auditor_name" in cites, "la section doit nommer le paramètre à employer"
    assert cites <= exposes, f"cités mais inexistants : {sorted(cites - exposes)}"
