"""Gardes du runner d'acceptation AVP (`scripts/avp_acceptance/run_acceptance.py`).

Teste, hors réseau, les refus purs du runner : garde catalogue CCH (document
absent / catalogue vide / catalogue valide) et garde d'écriture hors dépôt
(chemin dans le repo refusé, /tmp accepté). Le module est chargé par chemin (il
vit sous ``scripts/``) ; ses imports ``audit_bim`` sont tardifs (dans ``main``),
donc l'import du module ne déclenche aucune dépendance lourde.
"""

from __future__ import annotations

import importlib.util
import types
from pathlib import Path

import pytest
from docx import Document as _Docx
from docx.shared import RGBColor

from audit_bim.reporting.bimdata_brand import WORDMARK
from audit_bim.reporting.theming import BIMDATA_FONT_PRIMARY, BIMDATA_PRIMARY
from audit_bim.reporting.word_report import NOT_AVAILABLE

_RUNNER_PATH = (
    Path(__file__).resolve().parents[2] / "scripts" / "avp_acceptance" / "run_acceptance.py"
)
_spec = importlib.util.spec_from_file_location("avp_acceptance_runner", _RUNNER_PATH)
runner = importlib.util.module_from_spec(_spec)
_spec.loader.exec_module(runner)


def _fake_catalog(n_props: int, n_rules: int):
    return types.SimpleNamespace(properties=[0] * n_props, naming_rules=[0] * n_rules)


# ── Garde catalogue CCH ───────────────────────────────────────────────────


def test_catalog_guard_refuses_missing_document(tmp_path):
    docs = {"cch_pdf": str(tmp_path / "absent.pdf")}  # fichier inexistant
    with pytest.raises(SystemExit):
        runner.assert_catalog_usable(docs, _fake_catalog(10, 10))


def test_catalog_guard_refuses_empty_catalog(tmp_path):
    doc = tmp_path / "cch.pdf"
    doc.write_text("x")  # document présent...
    # ...mais catalogue sans règle de nommage → refus.
    with pytest.raises(SystemExit):
        runner.assert_catalog_usable({"cch_pdf": str(doc)}, _fake_catalog(10, 0))


def test_catalog_guard_refuses_empty_properties(tmp_path):
    doc = tmp_path / "cch.pdf"
    doc.write_text("x")
    with pytest.raises(SystemExit):
        runner.assert_catalog_usable({"cch_pdf": str(doc)}, _fake_catalog(0, 10))


def test_catalog_guard_accepts_valid(tmp_path):
    doc = tmp_path / "cch.pdf"
    doc.write_text("x")
    # Ne doit pas lever.
    runner.assert_catalog_usable({"cch_pdf": str(doc)}, _fake_catalog(5, 5))


# ── Garde d'écriture hors dépôt ───────────────────────────────────────────


def test_outside_repo_refuses_path_in_repo():
    # Un chemin sous le dépôt (à côté du runner) doit être refusé.
    inside = Path(runner.__file__).resolve().parent / "out"
    with pytest.raises(SystemExit):
        runner._assert_outside_repo(inside)


def test_outside_repo_accepts_tmp(tmp_path):
    # tmp_path (hors du dépôt) est accepté : ne doit pas lever.
    runner._assert_outside_repo(tmp_path)


# ── inspect_word_report : cas négatifs (docx synthétiques) ────────────────


def _brand(run):
    run.font.name = BIMDATA_FONT_PRIMARY  # Roboto
    run.font.color.rgb = RGBColor(0x2F, 0x37, 0x4A)  # primaire #2F374A


def _valid_docx(
    path, *, project_text="PROG", phase_text="AVP", sections=range(1, 10), significant_cells=12
):
    """Fabrique un docx qui satisfait TOUS les critères (base des cas négatifs)."""
    d = _Docx()
    for n in sections:
        title = runner.WORD_SECTION_TITLES[n]  # vrai intitulé métier attendu
        _brand(d.add_paragraph().add_run(f"{n}. {title}"))
    for i in range(5):
        d.add_paragraph().add_run(f"Contenu réel {i}")
    d.add_paragraph().add_run(WORDMARK)  # wordmark BIMDATA
    d.add_paragraph().add_run(f"Projet {project_text} — Phase {phase_text}")  # métadonnées
    t = d.add_table(rows=max(1, significant_cells), cols=1)
    for i in range(significant_cells):
        t.rows[i].cells[0].text = f"valeur {i}"
    d.save(str(path))
    return path


def _inspect(path, **over):
    kwargs = {
        "expected_project_name": "PROG",
        "phase": "AVP",
        "wordmark": WORDMARK,
        "primary": BIMDATA_PRIMARY,
        "font": BIMDATA_FONT_PRIMARY,
        "not_available": NOT_AVAILABLE,
    }
    kwargs.update(over)
    return runner.inspect_word_report(path, **kwargs)


def test_word_valid_synthetic_passes(tmp_path):
    # Garde-fou : le docx synthétique « valide » passe bien (sinon les cas
    # négatifs prouveraient n'importe quoi).
    r = _inspect(_valid_docx(tmp_path / "ok.docx"))
    assert r["ok"] is True, r


def test_word_reject_thin_branded_docx(tmp_path):
    # DOCX 1×1 brandé : charte OK mais contenu maigre → rejeté.
    p = tmp_path / "thin.docx"
    d = _Docx()
    _brand(d.add_paragraph().add_run(f"1. {WORDMARK}"))  # brandé + wordmark, 1 paragraphe
    t = d.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "donnée"  # 1 cellule
    d.save(str(p))
    r = _inspect(p)
    assert r["wordmark"] and r["primary"] and r["font"]  # bien brandé
    assert r["non_empty"] is False
    assert r["ok"] is False


def test_word_reject_cells_only_not_available(tmp_path):
    # Assez de paragraphes/sections, mais toutes les cellules = NOT_AVAILABLE.
    p = _valid_docx(tmp_path / "na.docx", significant_cells=0)
    d = _Docx(str(p))
    t = d.add_table(rows=12, cols=1)
    for i in range(12):
        t.rows[i].cells[0].text = NOT_AVAILABLE
    d.save(str(p))
    r = _inspect(p)
    assert r["n_significant_cells"] == 0
    assert r["non_empty"] is False
    assert r["ok"] is False


def test_word_reject_missing_required_section(tmp_path):
    # Sections 1..8 seulement (9 manquante) → sections_ok False.
    p = _valid_docx(tmp_path / "missing.docx", sections=range(1, 9))
    r = _inspect(p)
    assert 9 not in r["sections_present"]
    assert r["sections_ok"] is False
    assert r["ok"] is False


def test_word_reject_numbers_present_wrong_titles(tmp_path):
    # Les 9 numéros sont présents mais les INTITULÉS sont faux (« 1. Foo »… « 9.
    # Bar ») → le contrôle par intitulé métier doit refuser.
    p = tmp_path / "wrongtitles.docx"
    d = _Docx()
    for n in range(1, 10):
        _brand(d.add_paragraph().add_run(f"{n}. Titre bidon {n}"))
    for i in range(5):
        d.add_paragraph().add_run(f"Contenu réel {i}")
    d.add_paragraph().add_run(WORDMARK)
    d.add_paragraph().add_run("Projet PROG — Phase AVP")
    t = d.add_table(rows=12, cols=1)
    for i in range(12):
        t.rows[i].cells[0].text = f"valeur {i}"
    d.save(str(p))
    r = _inspect(p)
    assert r["sections_present"] == list(range(1, 10))  # numéros bien présents…
    assert r["sections_ok"] is False  # …mais intitulés faux → refus
    assert r["ok"] is False


def test_word_reject_absent_project_and_placeholder(tmp_path):
    # Doc valide par ailleurs, mais vrai nom de projet ABSENT → metadata False.
    p = _valid_docx(tmp_path / "noproj.docx", project_text=runner.PLACEHOLDER_PROJECT_NAME)
    r_none = _inspect(p, expected_project_name=None)
    assert r_none["metadata_present"] is False
    assert r_none["ok"] is False
    # Le bouchon ne satisfait JAMAIS le gate, même présent dans le doc.
    r_ph = _inspect(p, expected_project_name=runner.PLACEHOLDER_PROJECT_NAME)
    assert r_ph["metadata_present"] is False
    assert r_ph["ok"] is False
