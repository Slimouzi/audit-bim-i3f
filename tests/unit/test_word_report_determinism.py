"""Le rapport Word doit être reproductible : mêmes entrées, même document.

`_theme_block` recevait un `set[Theme]`. L'itération d'un ensemble dépend du
hash de ses membres, donc du processus : deux rendus du même audit sortaient
les lignes d'un bloc de thème dans un ordre différent. Sans erreur, sans que
rien ne le signale — un client comparant deux audits successifs y aurait vu un
changement qui n'existait pas.

Ces tests tournent dans des **interpréteurs séparés avec des `PYTHONHASHSEED`
différents** : dans un seul processus, le hash est fixe et le bug invisible.
C'est précisément ce qui l'avait laissé passer.
"""

from __future__ import annotations

import ast
import json
import subprocess
import sys
import textwrap
from pathlib import Path

import pytest

from audit_bim.audit.findings import Theme
from audit_bim.reporting import word_report

REPO = Path(word_report.__file__).resolve().parents[3]

RENDER_SCRIPT = textwrap.dedent(
    """
    import json, os, sys, tempfile, pathlib
    from docx import Document
    from audit_bim.audit.engine import AuditResult
    from audit_bim.audit.findings import ErrorType, Finding, Severity, Theme
    from audit_bim.extraction.model_data import ModelSnapshot
    from audit_bim.requirements.models import BIMPhase
    import audit_bim.reporting.word_report as wr

    wr.find_logo = lambda variant="light": None

    class Cat:
        cch_version = "3.6"; cch_source_pdf = "cch.pdf"; data_spec_source = "spec.xlsx"
        naming_spec_source = "nom.xlsx"; properties = []; naming_rules = []
        storey_names = []; zone_specs = []; room_specs = []

    snap = ModelSnapshot(); snap.project = {"name": "P"}; snap.model = {"name": "M"}
    themes = [Theme.NAMING_SITE_BAT_ETAGE, Theme.NAMING_ZONE, Theme.NAMING_SPACE,
              Theme.PROPERTY_MISSING, Theme.PROPERTY_INVALID, Theme.CLASSIFICATION]
    findings = [
        Finding(element_uuid=f"u{i}", ifc_type="IfcWall", name=f"W{i}", theme=th,
                error_type=ErrorType.CLASSIFICATION_MISSING, severity=Severity.HIGH,
                expected="x", actual=None, ref_cch="6.3")
        for i, th in enumerate(themes)
    ]
    result = AuditResult(snapshot=snap, catalog=Cat(), phase=BIMPhase.PRO, findings=findings)

    d = pathlib.Path(tempfile.mkdtemp()); os.environ["AUDIT_OUTPUT_DIR"] = str(d)
    wr.write_word_report(result, d / "r.docx")
    doc = Document(str(d / "r.docx"))
    dump = {
        "paragraphs": [p.text for p in doc.paragraphs],
        "tables": [[[c.text for c in row.cells] for row in t.rows] for t in doc.tables],
    }
    pathlib.Path(sys.argv[1]).write_text(json.dumps(dump, ensure_ascii=False))
    """
)


def _render_with_seed(tmp_path: Path, seed: str) -> dict:
    script = tmp_path / f"render_{seed}.py"
    script.write_text(RENDER_SCRIPT, encoding="utf-8")
    out = tmp_path / f"dump_{seed}.json"
    proc = subprocess.run(
        [sys.executable, str(script), str(out)],
        capture_output=True,
        text=True,
        cwd=str(REPO),
        env={**dict(__import__("os").environ), "PYTHONHASHSEED": seed},
    )
    assert proc.returncode == 0, f"rendu seed={seed} échoué :\n{proc.stderr}"
    return json.loads(out.read_text(encoding="utf-8"))


@pytest.mark.parametrize("seed", ["2", "7", "12345", "99", "31337"])
def test_report_is_identical_across_hash_seeds(tmp_path, seed):
    """Même document quel que soit le hash seed — sans PYTHONHASHSEED=0.

    Détecteur **probabiliste** : pour un petit ensemble, deux seeds peuvent
    donner le même ordre par chance. Il ne remplace donc pas
    :func:`test_theme_block_preserves_declared_order`, qui prouve la propriété
    directement. Les deux sont gardés : celui-ci attrape une régression
    ailleurs dans la chaîne, l'autre verrouille le contrat.
    """
    reference = _render_with_seed(tmp_path, "1")
    other = _render_with_seed(tmp_path, seed)
    assert reference["paragraphs"] == other["paragraphs"]
    assert reference["tables"] == other["tables"], (
        f"ordre des lignes dépendant du hash (seed={seed}) : le rendu n'est pas reproductible"
    )


def test_theme_block_preserves_declared_order(tmp_path, monkeypatch):
    """Preuve DIRECTE : les lignes suivent l'ordre déclaré du bloc nommage.

    Trois findings de même sévérité, un par thème de nommage. Le tri par
    sévérité étant stable, l'ordre des lignes ne peut venir que de l'itération
    des thèmes.

    Mesuré contre l'implémentation à ``set`` : échoue sur 3 hash seeds sur 4.
    Pas 4/4, parce qu'une permutation peut coïncider avec l'ordre déclaré — un
    test comportemental ne peut pas faire mieux face à un ordre aléatoire. Ce
    sont les garde-fous **statiques** ci-dessous qui constituent la vraie
    barrière : eux échouent systématiquement.
    """
    from docx import Document

    from audit_bim.audit.engine import AuditResult
    from audit_bim.audit.findings import ErrorType, Finding, Severity
    from audit_bim.extraction.model_data import ModelSnapshot
    from audit_bim.requirements.models import BIMPhase

    class Cat:
        cch_version = "3.6"
        cch_source_pdf = "cch.pdf"
        data_spec_source = "s"
        naming_spec_source = "n"
        properties: list = []
        naming_rules: list = []
        storey_names: list = []
        zone_specs: list = []
        room_specs: list = []

    snap = ModelSnapshot()
    snap.project = {"name": "P"}
    snap.model = {"name": "M"}
    declared = [
        (Theme.NAMING_SITE_BAT_ETAGE, "AAA-SITE"),
        (Theme.NAMING_ZONE, "BBB-ZONE"),
        (Theme.NAMING_SPACE, "CCC-SPACE"),
    ]
    findings = [
        Finding(
            element_uuid=f"u{i}",
            ifc_type="IfcWall",
            name=name,
            theme=theme,
            error_type=ErrorType.NAMING_MISSING,
            severity=Severity.HIGH,  # sévérité IDENTIQUE : le tri ne départage pas
            expected="x",
            actual=None,
        )
        for i, (theme, name) in enumerate(declared)
    ]
    result = AuditResult(snapshot=snap, catalog=Cat(), phase=BIMPhase.PRO, findings=findings)

    monkeypatch.setenv("AUDIT_OUTPUT_DIR", str(tmp_path))
    monkeypatch.setattr(word_report, "find_logo", lambda variant="light": None)
    out = tmp_path / "r.docx"
    word_report.write_word_report(result, out)

    names = [n for _, n in declared]
    seen: list[str] = []
    for table in Document(str(out)).tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text in names and cell.text not in seen:
                    seen.append(cell.text)
    assert seen == names, f"ordre rendu {seen}, ordre déclaré {names}"


# ── Garde-fou statique : plus aucun set de thèmes ─────────────────────


def test_theme_block_takes_a_sequence_not_a_set():
    src = ast.parse(Path(word_report.__file__).read_text(encoding="utf-8"))
    for node in ast.walk(src):
        if isinstance(node, ast.FunctionDef) and node.name == "_theme_block":
            annotation = ast.unparse(node.args.args[0].annotation)
            assert "set" not in annotation.lower(), f"_theme_block({annotation}) : ordre instable"
            return
    pytest.fail("_theme_block introuvable")


def test_no_theme_collection_is_declared_as_a_set():
    """Un set de Theme qui traîne finit par être itéré pour produire de la sortie."""
    src = Path(word_report.__file__).read_text(encoding="utf-8")
    offenders = [
        node.lineno
        for node in ast.walk(ast.parse(src))
        if isinstance(node, ast.Set)
        and any(
            isinstance(e, ast.Attribute) and getattr(e.value, "id", None) == "Theme"
            for e in node.elts
        )
    ]
    assert not offenders, f"littéraux set[Theme] aux lignes {offenders} : ordre non garanti"


def test_domains_themes_are_ordered_collections():
    for label, themes in word_report.DOMAINS:
        assert isinstance(themes, tuple), f"DOMAINS[{label!r}] n'est pas ordonné"
        assert all(isinstance(t, Theme) for t in themes)


def test_naming_block_order_follows_declaration():
    """L'ordre déclaré est la référence : site → zone → pièce."""
    src = Path(word_report.__file__).read_text(encoding="utf-8")
    call = "(Theme.NAMING_SITE_BAT_ETAGE, Theme.NAMING_ZONE, Theme.NAMING_SPACE)"
    assert call in src, "l'ordre de lecture du bloc nommage doit rester explicite"
