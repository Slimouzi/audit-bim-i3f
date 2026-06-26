"""Tests de la charte BIMData appliquée aux rapports d'audit.

Trois couches de couverture :

- :mod:`audit_bim.reporting.theming` — présence et cohérence des tokens
  BIMData, alignement des alias historiques ``KORHUS_*`` / ``I3F_*``.
- :mod:`audit_bim.reporting.bimdata_brand` — résolution du brand kit
  via variable d'env, fallback chemin par défaut, dégradation
  silencieuse.
- *Smoke render* — la génération Word + Excel construit un fichier
  valide quand le brand kit est trouvé ET quand il ne l'est pas (le
  rapport ne doit jamais planter pour une question de logo manquant).
"""

from __future__ import annotations

import re
from pathlib import Path

import pytest
from docx import Document
from openpyxl import load_workbook

from audit_bim.audit.engine import AuditResult
from audit_bim.reporting import bimdata_brand, theming
from audit_bim.reporting.theming import (
    BIMDATA_BLACK,
    BIMDATA_BLUE_NEUTRAL_LIGHT,
    BIMDATA_FONT_FALLBACK,
    BIMDATA_FONT_PRIMARY,
    BIMDATA_GRANITE,
    BIMDATA_GRANITE_LIGHT,
    BIMDATA_PRIMARY,
    BIMDATA_ROYAL_BLUE,
    BIMDATA_SECONDARY,
    BIMDATA_SILVER_DARK,
    BIMDATA_SILVER_LIGHT,
    BIMDATA_WHITE,
    I3F_BLUE,
    I3F_BLUE_LIGHT,
    I3F_GREY,
    KORHUS_FONT_PRIMARY,
    KORHUS_PRIMARY,
    KORHUS_SECONDARY,
    SEVERITY_COLORS,
)
from audit_bim.reporting.word_report import write_word_report
from audit_bim.reporting.xlsx_annex import write_xlsx_annex

HEX6 = re.compile(r"^[0-9A-Fa-f]{6}$")

# Valeurs canoniques attendues (cf. BRAND_GUIDELINES.md). Toute
# divergence ici = dérive par rapport à la charte BIMData 2022 v1.0.
_EXPECTED_TOKENS = {
    "BIMDATA_PRIMARY": "2F374A",
    "BIMDATA_SECONDARY": "F9C72C",
    "BIMDATA_ROYAL_BLUE": "3375DD",
    "BIMDATA_WHITE": "FFFFFF",
    "BIMDATA_GRANITE": "606060",
    "BIMDATA_BLACK": "000000",
    "BIMDATA_GRANITE_LIGHT": "7A7A7A",
    "BIMDATA_SILVER_DARK": "BDBDBD",
    "BIMDATA_SILVER_LIGHT": "F7F7F7",
    "BIMDATA_BLUE_NEUTRAL_LIGHT": "F0F5FF",
    "BIMDATA_HIGH": "FF3D1E",
    "BIMDATA_WARNING": "FF9100",
    "BIMDATA_SUCCESS": "00AF50",
    "BIMDATA_FONT_PRIMARY": "Roboto",
    "BIMDATA_FONT_FALLBACK": "Arial",
}


# ── 1. Tokens BIMData présents et cohérents ───────────────────────────


class TestBIMDataPalette:
    def test_brand_tokens_are_hex6(self):
        for name in (
            BIMDATA_PRIMARY,
            BIMDATA_SECONDARY,
            BIMDATA_WHITE,
            BIMDATA_GRANITE,
            BIMDATA_BLUE_NEUTRAL_LIGHT,
        ):
            assert HEX6.match(name), f"{name!r} n'est pas un hex 6 chars"

    def test_primary_is_bimdata_slate(self):
        # #2F374A — bleu ardoise, couleur de couverture et titres forts.
        assert BIMDATA_PRIMARY.upper() == "2F374A"

    def test_secondary_is_yellow_accent(self):
        assert BIMDATA_SECONDARY.upper() == "F9C72C"

    def test_blue_neutral_light_is_table_bg(self):
        assert BIMDATA_BLUE_NEUTRAL_LIGHT.upper() == "F0F5FF"


class TestDeprecatedAliases:
    """Les alias historiques pointent désormais sur les tokens BIMData."""

    def test_korhus_primary_aliases_bimdata(self):
        assert KORHUS_PRIMARY == BIMDATA_PRIMARY

    def test_korhus_secondary_aliases_bimdata(self):
        assert KORHUS_SECONDARY == BIMDATA_SECONDARY

    def test_korhus_font_aliases_bimdata(self):
        assert KORHUS_FONT_PRIMARY == BIMDATA_FONT_PRIMARY

    def test_i3f_blue_aliases_bimdata_primary(self):
        assert I3F_BLUE == BIMDATA_PRIMARY

    def test_i3f_blue_light_aliases_blue_neutral_light(self):
        assert I3F_BLUE_LIGHT == BIMDATA_BLUE_NEUTRAL_LIGHT

    def test_i3f_grey_aliases_bimdata_granite(self):
        assert I3F_GREY == BIMDATA_GRANITE


class TestBrandTokensExactValues:
    """Les tokens exécutables doivent matcher la charte (BRAND_GUIDELINES.md)."""

    def test_all_tokens_have_expected_values(self):
        actual = {name: getattr(theming, name) for name in _EXPECTED_TOKENS}
        assert actual == _EXPECTED_TOKENS

    def test_fonts_are_roboto_arial(self):
        assert BIMDATA_FONT_PRIMARY == "Roboto"
        assert BIMDATA_FONT_FALLBACK == "Arial"


class TestSeverityColorsUntouched:
    """Convention métier feux tricolores : la charte BIMData ne doit pas
    écraser le rouge / orange / vert des sévérités."""

    def test_critical_is_dark_red(self):
        r = int(SEVERITY_COLORS["CRITICAL"][0:2], 16)
        g = int(SEVERITY_COLORS["CRITICAL"][2:4], 16)
        b = int(SEVERITY_COLORS["CRITICAL"][4:6], 16)
        assert r > g and r > b  # rouge dominant

    def test_severity_colors_independent_from_brand_palette(self):
        """Les couleurs de sévérité ne doivent PAS être tirées de la
        palette de marque : c'est une convention métier autonome."""
        brand_palette = {
            BIMDATA_PRIMARY,
            BIMDATA_SECONDARY,
            BIMDATA_ROYAL_BLUE,
            BIMDATA_WHITE,
            BIMDATA_GRANITE,
            BIMDATA_BLACK,
            BIMDATA_GRANITE_LIGHT,
            BIMDATA_SILVER_DARK,
            BIMDATA_SILVER_LIGHT,
            BIMDATA_BLUE_NEUTRAL_LIGHT,
        }
        for sev, color in SEVERITY_COLORS.items():
            assert color.upper() not in {c.upper() for c in brand_palette}, (
                f"La sévérité {sev} ({color}) ne doit pas réutiliser un token de marque."
            )


class TestBrandGuidelinesDoc:
    """La charte éditoriale BRAND_GUIDELINES.md doit exister et rester la
    source de vérité référencée par le code."""

    def _doc_path(self) -> Path:
        return Path(theming.__file__).parent / "BRAND_GUIDELINES.md"

    def test_brand_guidelines_file_exists(self):
        assert self._doc_path().is_file(), "BRAND_GUIDELINES.md manquant"

    def test_brand_guidelines_mentions_key_tokens(self):
        text = self._doc_path().read_text(encoding="utf-8")
        # Hex primaires/accents présents dans la charte.
        for hexval in ("#2F374A", "#F9C72C", "#3375DD", "#F0F5FF"):
            assert hexval in text, f"{hexval} absent de BRAND_GUIDELINES.md"
        # Typographie + attributs de marque.
        assert "Roboto" in text and "Arial" in text
        for attr in ("Simplicity", "Modernity", "Technology", "Scalable"):
            assert attr in text

    def test_theming_module_references_the_doc(self):
        src = Path(theming.__file__).read_text(encoding="utf-8")
        assert "BRAND_GUIDELINES.md" in src


# ── 2. Résolution du brand kit ────────────────────────────────────────


class TestFindBrandKitDir:
    def test_env_override_takes_precedence(self, tmp_path, monkeypatch):
        # Simule un brand kit minimal sous tmp_path.
        (tmp_path / "assets").mkdir()
        monkeypatch.setenv("BIMDATA_BRAND_KIT_DIR", str(tmp_path))
        assert bimdata_brand.find_brand_kit_dir() == tmp_path

    def test_legacy_env_var_still_supported(self, tmp_path, monkeypatch):
        # Variable héritée KORHUS_BRAND_KIT_DIR conservée en compat.
        (tmp_path / "assets").mkdir()
        monkeypatch.delenv("BIMDATA_BRAND_KIT_DIR", raising=False)
        monkeypatch.setenv("KORHUS_BRAND_KIT_DIR", str(tmp_path))
        assert bimdata_brand.find_brand_kit_dir() == tmp_path

    def test_env_invalid_falls_back_to_sibling_or_none(self, tmp_path, monkeypatch):
        # Variable pointant vers un chemin inexistant : on tombe sur le
        # scan sibling (ou None si aucun voisin brand kit).
        monkeypatch.setenv("BIMDATA_BRAND_KIT_DIR", str(tmp_path / "ghost"))
        monkeypatch.delenv("KORHUS_BRAND_KIT_DIR", raising=False)
        result = bimdata_brand.find_brand_kit_dir()
        assert result is None or (
            result.is_dir() and result.name in {"bimdata_brand_kit", "korhus_brand_kit"}
        )

    def test_returns_none_when_nothing_found(self, monkeypatch):
        monkeypatch.setenv("BIMDATA_BRAND_KIT_DIR", "/nonexistent/path/x")
        monkeypatch.delenv("KORHUS_BRAND_KIT_DIR", raising=False)
        # Le scan sibling pourrait remonter à un dossier existant — on
        # patch ``Path.is_dir`` pour qu'il ne dise jamais True.
        monkeypatch.setattr(Path, "is_dir", lambda self: False)
        assert bimdata_brand.find_brand_kit_dir() is None


class TestFindLogo:
    def test_unknown_variant_raises(self):
        with pytest.raises(ValueError, match="Variante logo"):
            bimdata_brand.find_logo("turquoise")

    def test_returns_none_when_brand_kit_missing(self, monkeypatch):
        monkeypatch.setattr(bimdata_brand, "find_brand_kit_dir", lambda: None)
        # Pas de logo « vrac » non plus → repli wordmark côté rapport.
        monkeypatch.setattr(bimdata_brand, "_find_loose_logo", lambda: None)
        assert bimdata_brand.find_logo("light") is None

    def test_loose_logo_used_when_brand_kit_missing(self, tmp_path, monkeypatch):
        """Sans brand kit structuré, un logo « vrac » valide est retenu."""
        valid = tmp_path / "Logo_Bimdata.png"
        valid.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 16)
        monkeypatch.setattr(bimdata_brand, "find_brand_kit_dir", lambda: None)
        monkeypatch.setattr(bimdata_brand, "_find_loose_logo", lambda: valid)
        assert bimdata_brand.find_logo("light") == valid

    def test_returns_path_when_logo_present(self, tmp_path, monkeypatch):
        assets = tmp_path / "assets"
        assets.mkdir()
        fake_logo = assets / "bimdata_logo_white.png"
        fake_logo.write_bytes(b"\x89PNG\r\n\x1a\n")  # signature PNG bidon
        monkeypatch.setattr(bimdata_brand, "find_brand_kit_dir", lambda: tmp_path)
        # Pas de logo vrac à scanner pour ce test.
        monkeypatch.setattr(bimdata_brand, "_find_loose_logo", lambda: None)
        assert bimdata_brand.find_logo("light") == fake_logo


class TestRasterImageValidation:
    def test_real_png_accepted(self, tmp_path):
        p = tmp_path / "logo.png"
        p.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 16)
        assert bimdata_brand._is_raster_image(p) is True

    def test_real_jpeg_accepted(self, tmp_path):
        p = tmp_path / "logo.jpg"
        p.write_bytes(b"\xff\xd8\xff\xe0" + b"\x00" * 16)
        assert bimdata_brand._is_raster_image(p) is True

    def test_appledouble_metadata_file_rejected(self, tmp_path):
        # Fichier AppleDouble macOS : que des métadonnées, pas d'image.
        p = tmp_path / "Logo-bimdata.jpg"
        p.write_bytes(b"\x00\x05\x16\x07\x00\x02\x00\x00Mac OS X")
        assert bimdata_brand._is_raster_image(p) is False

    def test_first_valid_raster_skips_appledouble_and_picks_valid(self, tmp_path):
        folder = tmp_path / "Logo_BIMData"
        folder.mkdir()
        # Fichier AppleDouble (alphabétiquement avant 'real') doit être ignoré.
        (folder / "junk.jpg").write_bytes(b"\x00\x05\x16\x07Mac OS X")
        valid = folder / "real.png"
        valid.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 16)
        assert bimdata_brand._first_valid_raster(folder) == valid

    def test_first_valid_raster_none_when_only_invalid(self, tmp_path):
        folder = tmp_path / "Logo_BIMData"
        folder.mkdir()
        (folder / "Logo-bimdata.jpg").write_bytes(b"\x00\x05\x16\x07Mac OS X")
        assert bimdata_brand._first_valid_raster(folder) is None

    def test_first_valid_raster_none_when_missing_folder(self, tmp_path):
        assert bimdata_brand._first_valid_raster(tmp_path / "nope") is None


# ── 3. Smoke render Word + Excel ──────────────────────────────────────


@pytest.fixture
def minimal_audit_result(catalog, snapshot_with_walls) -> AuditResult:
    """Audit avec un finding suffisant pour exercer les tableaux + KPIs."""
    from audit_bim.audit.findings import ErrorType, Finding, Severity, Theme
    from audit_bim.requirements.models import BIMPhase

    return AuditResult(
        catalog=catalog,
        snapshot=snapshot_with_walls,
        phase=BIMPhase.PRO,
        findings=[
            Finding(
                theme=Theme.PROPERTY_MISSING,
                severity=Severity.HIGH,
                error_type=ErrorType.PROPERTY_MISSING,
                element_uuid="W1",
                ifc_type="IfcWallStandardCase",
                name="Mur extérieur 01",
                expected="Pset_WallCommon/IsExternal",
                actual=None,
                ref_cch="Chap 6.2",
                recommended_action="Renseigner IsExternal.",
            ),
        ],
    )


class TestSmokeWordRender:
    def test_word_report_builds_with_brand_kit(self, minimal_audit_result, tmp_path):
        """Avec brand kit disponible (poste dev) : le rendu doit se
        construire sans erreur python-docx."""
        out = tmp_path / "audit.docx"
        write_word_report(minimal_audit_result, out)
        assert out.exists() and out.stat().st_size > 0

        doc = Document(str(out))
        # La police Normal doit être Roboto (charte BIMData).
        assert doc.styles["Normal"].font.name == BIMDATA_FONT_PRIMARY

    def test_word_report_builds_without_brand_kit(
        self, minimal_audit_result, tmp_path, monkeypatch
    ):
        """Sans brand kit : la couverture dégrade en wordmark texte
        « BIMDATA », le rapport reste générable."""
        from audit_bim.reporting import word_report

        # Simule absence totale du brand kit.
        monkeypatch.setattr(word_report, "find_logo", lambda variant="light": None)
        out = tmp_path / "audit_no_logo.docx"
        write_word_report(minimal_audit_result, out)
        assert out.exists() and out.stat().st_size > 0
        doc = Document(str(out))
        # On cherche la chaîne fallback "BIMDATA" dans les runs du doc.
        all_text = "\n".join(
            run.text
            for tbl in doc.tables
            for row in tbl.rows
            for cell in row.cells
            for para in cell.paragraphs
            for run in para.runs
        )
        assert "BIMDATA" in all_text
        assert "KORHUS" not in all_text


class TestSmokeXlsxRender:
    def test_xlsx_annex_builds_and_uses_roboto(self, minimal_audit_result, tmp_path):
        out = tmp_path / "annex.xlsx"
        write_xlsx_annex(minimal_audit_result, out)
        assert out.exists() and out.stat().st_size > 0

        wb = load_workbook(out)
        # Onglet "Synthèse" doit exister + contenir le supertitle BIMData.
        assert "Synthèse" in wb.sheetnames
        synth = wb["Synthèse"]
        assert synth["A1"].value == "BIMDATA — AUDIT BIM"
        # La police appliquée doit bien être Roboto.
        assert synth["A3"].font.name == BIMDATA_FONT_PRIMARY

    def test_xlsx_referential_sheet_has_bimdata_banner(self, minimal_audit_result, tmp_path):
        out = tmp_path / "annex_ref.xlsx"
        write_xlsx_annex(minimal_audit_result, out)
        wb = load_workbook(out)
        assert "Référentiel I3F" in wb.sheetnames
        ref = wb["Référentiel I3F"]
        assert ref["A1"].value == "BIMDATA — RÉFÉRENTIEL"
