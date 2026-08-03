"""Rapport consolidé « Analyse BIM AVP » (.docx) et l'ensemble de ses helpers."""

from __future__ import annotations

import re
from datetime import date, datetime
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..avp_snapshot import snapshot_shab_total
from ..theming import BIMDATA_FONT_FALLBACK, BIMDATA_PRIMARY
from ..word_report import NOT_AVAILABLE, _kpi_table, _shade_cell
from .models import _CONTROLE_STATS_SHEETS, AvpMeta
from .xlsx_common import _cell, _fmt_meta, _norm, _pct, _stat_lookup
from .xlsx_controle import _controle_grid, _controle_rows_for_moa, _find_control_header_row
from .xlsx_enveloppe import _note_menuiseries, _note_methodologie

_MOA_FONT = "Calibri"
_MOA_TABLE_HEADER = "D9EAD3"
_MOA_TABLE_SECTION = "EAF2F8"


def _setup_docx() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _MOA_FONT
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _MOA_FONT)
    rfonts.set(qn("w:hAnsi"), _MOA_FONT)
    rfonts.set(qn("w:cs"), BIMDATA_FONT_FALLBACK)
    _configure_section(doc.sections[0], WD_ORIENT.PORTRAIT)
    return doc


def _build_analyse_bim_avp_docx(
    path, result, sources, meta, snap=None, controle_xlsx: Path | None = None
) -> Path:
    doc = _setup_docx()
    ctrl = sources.controle if sources else None
    env = sources.enveloppe if sources else None
    pieces = _stat_lookup(ctrl, "Pièces Nommage")
    zones = _stat_lookup(ctrl, "Zones Nommage")
    materiau = _stat_lookup(ctrl, "ARC absence de matériau")

    _write_moa_cover_page(doc, ctrl, meta)
    doc.add_page_break()
    _write_moa_contents_and_inputs(doc, ctrl, meta, result, sources, snap)
    _write_audit_synthese(doc, result)
    _write_computed_coverage(doc, snap)
    _write_ecarts(doc, result, sources, snap)
    _write_moa_indicateurs(doc, pieces, zones, materiau, env)

    _set_orientation(doc, WD_ORIENT.LANDSCAPE)
    _write_moa_control_grid_pages(doc, ctrl, meta, controle_xlsx)
    _write_moa_control_annex_pages(
        doc,
        ctrl,
        meta,
        controle_xlsx,
        use_source_cached_values=(snap is None and result is None),
    )

    doc.save(str(path))
    return path


def _configure_section(section, orient) -> None:
    section.orientation = orient
    if orient == WD_ORIENT.LANDSCAPE:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.left_margin = Cm(0.9)
        section.right_margin = Cm(0.9)
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
    else:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)


def _set_docx_run_font(run, *, size: float = 10, bold: bool | None = None) -> None:
    run.font.name = _MOA_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _MOA_FONT)
    rfonts.set(qn("w:hAnsi"), _MOA_FONT)
    rfonts.set(qn("w:cs"), BIMDATA_FONT_FALLBACK)


def _add_moa_paragraph(
    doc: Document,
    text: str = "",
    *,
    size: float = 10,
    bold: bool = False,
    align=None,
    space_after: float = 3,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_docx_run_font(run, size=size, bold=bold)
    return p


def _write_moa_title(doc: Document, text: str, *, size: float = 14) -> None:
    p = _add_moa_paragraph(doc, text, size=size, bold=True, space_after=6)
    p.paragraph_format.keep_with_next = True


def _write_moa_cover_page(doc: Document, ctrl, meta: AvpMeta) -> None:
    _add_moa_paragraph(
        doc,
        "Rapport d’analyse des maquettes numériques",
        size=12,
        bold=True,
        space_after=38,
    )
    _add_moa_paragraph(
        doc,
        "Rapport d’analyse\ndes maquettes numériques – 3F",
        size=22,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=22,
    )
    operation = _operation_label(meta)
    _add_moa_paragraph(
        doc,
        operation,
        size=13,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    _add_moa_paragraph(
        doc,
        f"Programme {meta.project_code}" if meta.project_code else "Programme",
        size=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=28,
    )
    _add_moa_paragraph(
        doc,
        f"Les maquettes IFC remises pour le rendu {meta.phase} ont permis une analyse "
        f"partielle de la phase {meta.phase} :",
        size=10,
        space_after=6,
    )
    _write_moa_usages(doc, meta)
    _write_moa_phase_table(doc, ctrl, meta)


def _operation_label(meta: AvpMeta) -> str:
    if meta.nombre_logements and meta.project_name:
        return f"Opération de construction de {meta.nombre_logements} à {meta.project_name}"
    if meta.project_name and meta.project_code:
        return f"Opération {meta.project_name} – Programme {meta.project_code}"
    if meta.project_name:
        return f"Opération {meta.project_name}"
    return "Opération"


def _write_moa_usages(doc: Document, meta: AvpMeta) -> None:
    if meta.usages_bim:
        for usage in meta.usages_bim:
            _add_moa_paragraph(doc, str(usage), size=10, space_after=2)
    else:
        _add_moa_paragraph(doc, "Usages BIM 3F : " + NOT_AVAILABLE + ".", size=10, space_after=2)
    if meta.temoin_virtuel:
        _add_moa_paragraph(doc, str(meta.temoin_virtuel), size=10, space_after=12)


def _write_moa_phase_table(doc: Document, ctrl, meta: AvpMeta) -> None:
    h = (ctrl.header if ctrl else {}) or {}
    date_controle = meta.date_controle or h.get("date d'analyse")
    auteur = meta.auteur_controle or meta.auditor
    rows = [
        ["Phase", "Date du contrôle", "Contrôle effectué par :"],
        [meta.phase, _fmt_meta(date_controle), _fmt_meta(auteur)],
    ]
    _write_docx_table(doc, rows, widths=(3.0, 5.0, 8.0), font_size=10, header_rows=1)


def _write_moa_contents_and_inputs(
    doc: Document, ctrl, meta: AvpMeta, result, sources, snap=None
) -> None:
    _add_moa_paragraph(doc, "Rapport d’analyse des maquettes numériques", size=11, bold=True)
    _write_moa_title(doc, "Table des matières", size=16)
    for label, page in (
        ("1. Données d’entrées", "2"),
        ("2. Grille de contrôle", "3"),
        ("3. Annexes de contrôle", "4"),
    ):
        _add_moa_paragraph(doc, f"{label} {' .' * 42} {page}", size=10, space_after=1)
    _write_moa_title(doc, "1. Données d’entrées", size=14)
    h = (ctrl.header if ctrl else {}) or {}
    n_models = _snapshot_model_count(
        snap if snap is not None else (result.snapshot if result else None)
    )
    _add_moa_paragraph(doc, _models_transmitted_label(n_models), size=10, space_after=3)
    _add_moa_paragraph(doc, "Exports IFC réalisés le :", size=10, space_after=2)
    export_date = h.get("maquettes ifc transmises le") or h.get("date d'analyse")
    model_name = _snapshot_model_name(
        snap if snap is not None else (result.snapshot if result else None)
    )
    if export_date not in (None, ""):
        _add_moa_paragraph(doc, f"{_fmt_meta(export_date)} : {model_name}", size=10, space_after=6)
    else:
        _add_moa_paragraph(doc, NOT_AVAILABLE, size=10, space_after=6)
    _write_donnees_entree(doc, ctrl, meta)
    _write_moa_title(doc, "Synthèse technique MCP", size=12)


def _snapshot_model_count(snap) -> int | None:
    if snap is None:
        return None
    models = getattr(snap, "models", None)
    if isinstance(models, list) and models:
        return len(models)
    model = getattr(snap, "model", None)
    return 1 if model else None


def _models_transmitted_label(n_models: int | None) -> str:
    if not isinstance(n_models, int):
        return "Maquettes transmises : " + NOT_AVAILABLE
    return (
        f"{n_models} maquette{'s' if n_models > 1 else ''} transmise{'s' if n_models > 1 else ''}"
    )


def _snapshot_model_name(snap) -> str:
    model = getattr(snap, "model", None) if snap is not None else None
    if isinstance(model, dict):
        name = model.get("name")
        if name:
            return str(name)
    return "maquette IFC"


def _write_moa_indicateurs(doc: Document, pieces, zones, materiau, env) -> None:
    ratio = env.ratio_fac_shab if env else None
    seuil = env.seuil_3f if env else None
    if isinstance(ratio, (int, float)) and isinstance(seuil, (int, float)):
        ratio_ok = "Conforme" if ratio >= seuil else "Non conforme"
    else:
        ratio_ok = NOT_AVAILABLE
    rows = [
        ["Indicateur", "Valeur"],
        [
            "Taux de conformité nommage pièces",
            _pct(pieces.get("conforme_ratio")) if pieces else NOT_AVAILABLE,
        ],
        [
            "Taux de conformité nommage zones",
            _pct(zones.get("conforme_ratio")) if zones else NOT_AVAILABLE,
        ],
        [
            "Éléments sans matériau (taux)",
            _pct(materiau.get("non_conforme_ratio")) if materiau else NOT_AVAILABLE,
        ],
        ["Ratio FAC/SHAB", f"{ratio:.3f}" if isinstance(ratio, (int, float)) else NOT_AVAILABLE],
        [
            f"Seuil 3F 2026 (≥ {seuil})" if isinstance(seuil, (int, float)) else "Seuil 3F 2026",
            ratio_ok,
        ],
    ]
    _write_docx_table(doc, rows, widths=(7.5, 4.0), font_size=8.5, header_rows=1)


def _write_moa_control_grid_pages(
    doc: Document, ctrl, meta: AvpMeta, controle_xlsx: Path | None
) -> None:
    _write_moa_title(doc, "3. Grille de contrôle des exigences du CCH BIM 3F", size=13)
    _add_moa_paragraph(
        doc,
        "Les maquettes numériques ont été analysées suivant la liste des critères suivants :",
        size=8,
        space_after=4,
    )
    _write_moa_grid_metadata(doc, ctrl, meta)
    rows = _control_grid_rows_from_xlsx(controle_xlsx) if controle_xlsx else []
    if not rows and ctrl and ctrl.grille:
        rows = [ctrl.grille.headers, *_controle_rows_for_moa(ctrl.grille)]
    if not rows:
        _add_moa_paragraph(doc, NOT_AVAILABLE, size=9)
        return
    _write_docx_table(
        doc,
        rows,
        widths=_GRILLE_COL_WIDTHS,
        font_size=6.2,
        header_rows=1,
        max_rows=85,
        repeat_header=True,
    )


def _write_moa_grid_metadata(doc: Document, ctrl, meta: AvpMeta) -> None:
    h = (ctrl.header if ctrl else {}) or {}
    left = [
        ["Projet", meta.project_name or _fmt_meta(h.get("projet"))],
        ["ESI", meta.project_code or _fmt_meta(h.get("esi"))],
        ["Phase", meta.phase or _fmt_meta(h.get("phase"))],
        ["Maquettes IFC transmises le", _fmt_meta(h.get("maquettes ifc transmises le"))],
        ["Date d'analyse", _fmt_meta(h.get("date d'analyse"))],
        ["Version d'analyse", _fmt_meta(h.get("version d'analyse"))],
    ]
    legend = (ctrl.legend if ctrl else {}) or {
        0: "Non fourni / non trouvé",
        1: "Insuffisant : à reprendre ou compléter",
        2: "Satisfaisant",
    }
    rows = []
    for idx in range(max(len(left), len(legend))):
        left_row = left[idx] if idx < len(left) else ["", ""]
        code = sorted(legend)[idx] if idx < len(legend) else ""
        rows.append(
            [left_row[0], left_row[1], "" if idx else "Légende :", code, legend.get(code, "")]
        )
    _write_docx_table(doc, rows, widths=(4.0, 4.0, 2.0, 1.2, 6.0), font_size=7.2)


def _write_moa_control_annex_pages(
    doc: Document,
    ctrl,
    meta: AvpMeta,
    controle_xlsx: Path | None,
    *,
    use_source_cached_values: bool = False,
) -> None:
    for sheet_name in _CONTROLE_STATS_SHEETS:
        doc.add_page_break()
        _write_moa_title(doc, f"onglet {sheet_name}", size=12)
        rows = _control_annex_rows_from_xlsx(controle_xlsx, sheet_name) if controle_xlsx else []
        grid = _controle_grid(sheet_name, ctrl)
        if use_source_cached_values and grid and grid.rows:
            rows = grid.rows
        elif rows:
            rows = _inject_annex_summary_values(rows, _stat_lookup(ctrl, sheet_name), sheet_name)
        if not rows:
            rows = grid.rows if grid else []
        if not rows:
            rows = _fallback_control_annex_rows(sheet_name, _stat_lookup(ctrl, sheet_name), meta)
        _write_docx_table(
            doc,
            rows,
            widths=_annex_widths(sheet_name),
            font_size=5.8,
            header_rows=0,
            max_rows=_annex_row_cap(sheet_name),
        )


#: Lignes de **pilotage** du classeur MOA : consignes de saisie (« zone de copie
#: de la liste… », « coller ici ») et marques d'outils tiers. Elles ont un sens
#: dans un Excel qu'on remplit à la main ; reprises dans le rapport remis au
#: client, elles produisent un document qui explique comment se remplir et cite
#: l'outillage d'un autre chantier. Le .xlsx généré est déjà nettoyé de ses
#: marques — ce filtre traite le cas distinct de la ligne d'instruction, et
#: couvre aussi la relecture d'un classeur non régénéré.
_PLACEHOLDER_ROW_RE = re.compile(r"zone de copie|coller ici|bimcollab|solibri", re.IGNORECASE)


def _is_placeholder_row(values: list) -> bool:
    return any(isinstance(v, str) and _PLACEHOLDER_ROW_RE.search(v) for v in values)


def _control_grid_rows_from_xlsx(controle_xlsx: Path | None) -> list[list]:
    if not controle_xlsx or not Path(controle_xlsx).exists():
        return []
    wb = openpyxl.load_workbook(controle_xlsx, data_only=False, read_only=True)
    try:
        if "Grille de contrôle" not in wb.sheetnames:
            return []
        ws = wb["Grille de contrôle"]
        header_row = _find_control_header_row(ws)
        if header_row is None:
            return []
        rows: list[list] = []
        for row in ws.iter_rows(min_row=header_row, max_col=6, values_only=True):
            values = [_docx_excel_value(v) for v in row]
            if not any(values) or _is_placeholder_row(values):
                continue
            rows.append(values)
        return rows
    finally:
        wb.close()


def _control_annex_rows_from_xlsx(controle_xlsx: Path | None, sheet_name: str) -> list[list]:
    if not controle_xlsx or not Path(controle_xlsx).exists():
        return []
    wb = openpyxl.load_workbook(controle_xlsx, data_only=True, read_only=True)
    try:
        if sheet_name not in wb.sheetnames:
            return []
        ws = wb[sheet_name]
        max_cols = min(ws.max_column, 11)
        max_rows = min(ws.max_row, _annex_row_cap(sheet_name))
        rows: list[list] = []
        for row in ws.iter_rows(min_row=1, max_row=max_rows, max_col=max_cols, values_only=True):
            values = [_docx_excel_value(v) for v in row]
            if any(values) and not _is_placeholder_row(values):
                rows.append(values)
        return rows
    finally:
        wb.close()


def _docx_excel_value(value) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, float):
        return f"{value:.2f}".rstrip("0").rstrip(".")
    if isinstance(value, str) and value.startswith("="):
        return ""
    text = str(value)
    if "openpyxl.worksheet.formula" in text:
        return ""
    return text


def _inject_annex_summary_values(
    rows: list[list], stats: dict | None, sheet_name: str
) -> list[list]:
    if not stats:
        return rows
    out = [list(row) for row in rows]
    for row in out:
        label = str(row[1] if len(row) > 1 else "").strip().lower()
        if "matériau" in sheet_name.lower() or "materiau" in _norm(sheet_name):
            if "sans" in label and "mat" in label:
                _set_row_value(row, 2, stats.get("non_conforme"))
                _set_row_value(row, 3, _pct(stats.get("non_conforme_ratio")))
                _set_row_value(row, 5, "Nombre d'élements :")
                _set_row_value(row, 6, stats.get("total"))
            elif "type" in label:
                _set_row_value(row, 2, _count_annex_types(out))
            continue
        if "noms" in label and "type" not in label:
            _set_row_value(row, 2, stats.get("total"))
            _set_row_value(row, 3, stats.get("conforme"))
            _set_row_value(row, 4, _pct(stats.get("conforme_ratio")))
            _set_row_value(row, 5, stats.get("non_conforme"))
            _set_row_value(row, 6, _pct(stats.get("non_conforme_ratio")))
        elif "type" in label:
            _set_row_value(row, 2, _count_annex_types(out))
    return out


def _set_row_value(row: list, idx: int, value) -> None:
    while len(row) <= idx:
        row.append("")
    if value not in (None, ""):
        row[idx] = str(value)


def _count_annex_types(rows: list[list]) -> int:
    values = {
        str(row[1]).strip()
        for row in rows
        if len(row) > 2
        and str(row[1]).strip()
        and str(row[2]).strip()
        and not str(row[1]).strip().lower().startswith(("name", "object", "ifc"))
    }
    return len(values)


def _fallback_control_annex_rows(sheet_name: str, stats: dict | None, meta: AvpMeta) -> list[list]:
    rows = [[meta.project_code], [meta.project_code], [f"onglet {sheet_name}"]]
    if not stats:
        rows.append([NOT_AVAILABLE])
        return rows
    if "matériau" in sheet_name.lower():
        rows.extend(
            [
                ["", "MN"],
                [
                    stats.get("label"),
                    stats.get("non_conforme"),
                    _pct(stats.get("non_conforme_ratio")),
                ],
                ["Total éléments", stats.get("total")],
            ]
        )
        return rows
    rows.extend(
        [
            ["", "", "MN", "", "Conforme", "", "Non Conforme"],
            [
                "",
                stats.get("label"),
                stats.get("total"),
                stats.get("conforme"),
                _pct(stats.get("conforme_ratio")),
                stats.get("non_conforme"),
                _pct(stats.get("non_conforme_ratio")),
            ],
        ]
    )
    return rows


def _annex_row_cap(sheet_name: str) -> int:
    if sheet_name == "Zones Nommage":
        return 48
    if sheet_name == "Pièces Nommage":
        return 62
    if "matériau" in sheet_name.lower():
        return 42
    return 58


def _annex_widths(sheet_name: str) -> tuple[float, ...]:
    if "matériau" in sheet_name.lower():
        return (2.4, 5.2, 1.4, 4.0, 1.0, 5.2, 1.8, 1.0)
    if sheet_name == "Pièces Nommage":
        return (2.0, 3.6, 1.1, 2.3, 1.0, 3.0, 1.0, 1.0, 3.0, 4.0, 1.0)
    return (2.0, 4.2, 1.1, 2.3, 1.0, 3.0, 1.0, 1.0, 4.2, 4.4)


def _write_docx_table(
    doc: Document,
    rows: list[list],
    *,
    widths: tuple[float, ...] | None = None,
    font_size: float = 8,
    header_rows: int = 0,
    max_rows: int | None = None,
    repeat_header: bool = False,
):
    if max_rows is not None:
        rows = rows[:max_rows]
    if not rows:
        return None
    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if repeat_header and row_idx == 0:
            _set_repeat_table_header(table.rows[0])
        for col_idx, cell in enumerate(cells):
            value = values[col_idx] if col_idx < len(values) else ""
            cell.text = "" if value is None else str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths and col_idx < len(widths):
                cell.width = Cm(widths[col_idx])
            if row_idx < header_rows or _looks_like_moa_section_row(values):
                _shade_cell(
                    cell, _MOA_TABLE_HEADER if row_idx < header_rows else _MOA_TABLE_SECTION
                )
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_docx_run_font(run, size=font_size, bold=row_idx < header_rows)
    return table


def _looks_like_moa_section_row(values: list) -> bool:
    non_empty = [v for v in values if v not in (None, "")]
    if len(non_empty) > 2:
        return False
    first = str(values[0] if values else "").strip()
    return first.isdigit() or first in {"Zones", "Pièces", "ObjT Pièces"}


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_orientation(doc, orient) -> None:
    """Nouvelle section avec orientation portrait/paysage (lisibilité des
    tableaux larges en livrable client)."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(sec, orient)


# Largeurs (cm) des 6 colonnes de la grille de contrôle en paysage.
_GRILLE_COL_WIDTHS = (2.0, 6.5, 5.0, 3.0, 3.0, 7.0)


def _docx_header_table(doc, headers, *, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    tbl.allow_autofit = False
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = str(h)
        if col_widths and i < len(col_widths):
            cell.width = Cm(col_widths[i])
        _shade_cell(cell, BIMDATA_PRIMARY)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
    return tbl


def _write_computed_coverage(doc, snap) -> None:
    """Couverture des quantités **calculées** (fusion IfcOpenShell, Lot 3), si
    présente sur le snapshot — sinon rien. Ne masque aucun gap : affiche les
    valeurs conservées / ignorées / uuid inconnus."""
    cov = getattr(snap, "computed_coverage", None) if snap is not None else None
    if not cov:
        return
    p = doc.add_paragraph()
    p.add_run("Quantités calculées (IfcOpenShell) — couverture").bold = True
    doc.add_paragraph(
        "Valeurs géométriques calculées par IfcOpenShell et fusionnées en gap-only "
        "(les BaseQuantities natives BIMData ne sont jamais écrasées). Valeurs "
        "**non contractuelles** — un ré-export maquette avec BaseQuantities natives "
        "reste recommandé.",
        style="Intense Quote",
    )
    _kpi_table(
        doc,
        [
            ("Quantités fusionnées (comblées)", str(cov.get("n_merged", 0))),
            ("BaseQuantities natives conservées", str(cov.get("n_gap_kept", 0))),
            ("Entrées ignorées (skipped/failed)", str(cov.get("n_skipped_status", 0))),
            ("GlobalId inconnus du snapshot", str(cov.get("n_unknown_uuid", 0))),
        ],
    )


def _write_audit_synthese(doc, result) -> None:
    """Synthèse de l'audit BIMData réel (sévérité, thèmes, quantités
    manquantes) — le consolidé ne doit pas ignorer l'``AuditResult``."""
    if result is None:
        doc.add_paragraph("Audit automatisé : " + NOT_AVAILABLE + " (aucun audit chargé).")
        return
    by_sev = result.count_by_severity()
    by_theme = result.count_by_theme()
    by_type = result.count_by_error_type()
    p = doc.add_paragraph()
    p.add_run("Audit automatisé de la maquette active").bold = True
    # Répartition par sévérité COMPLÈTE (CRITICAL→INFO) pour que le total
    # se réconcilie côté client.
    _kpi_table(
        doc,
        [
            ("Anomalies détectées", str(len(result.findings))),
            ("Taux de conformité (pondéré)", f"{result.conformity_rate() * 100:.0f} %"),
            ("CRITICAL", str(by_sev.get("CRITICAL", 0))),
            ("HIGH", str(by_sev.get("HIGH", 0))),
            ("MEDIUM", str(by_sev.get("MEDIUM", 0))),
            ("LOW", str(by_sev.get("LOW", 0))),
            ("INFO", str(by_sev.get("INFO", 0))),
            ("Quantités manquantes", str(by_type.get("spatial_missing_quantity", 0))),
        ],
    )
    top = sorted(by_theme.items(), key=lambda kv: -kv[1])[:5]
    if top:
        doc.add_paragraph("Principaux thèmes en écart :")
        for theme, count in top:
            doc.add_paragraph(f"• {theme} : {count}", style="List Bullet")


def _write_donnees_entree(doc, ctrl, meta) -> None:
    h = (ctrl.header if ctrl else {}) or {}
    _kpi_table(
        doc,
        [
            ("Opération", _fmt_meta(meta.nombre_logements)),
            ("Maquettes IFC transmises le", _fmt_meta(h.get("maquettes ifc transmises le"))),
            ("Date d'analyse", _fmt_meta(h.get("date d'analyse"))),
            ("Version d'analyse", _fmt_meta(h.get("version d'analyse"))),
            ("Date du contrôle", _fmt_meta(meta.date_controle)),
            ("Auteur du contrôle", _fmt_meta(meta.auteur_controle)),
            ("Témoin virtuel", _fmt_meta(meta.temoin_virtuel)),
        ],
    )


def _write_grille_table(doc, ctrl, *, cap: int = 40) -> None:
    g = ctrl.grille if ctrl else None
    if not g or not g.headers:
        doc.add_paragraph(NOT_AVAILABLE)
        return
    # Largeurs adaptées uniquement pour la grille standard 6 colonnes.
    widths = _GRILLE_COL_WIDTHS if len(g.headers) == len(_GRILLE_COL_WIDTHS) else None
    tbl = _docx_header_table(doc, g.headers, col_widths=widths)
    for rowvals in g.rows[:cap]:
        cells = tbl.add_row().cells
        for i in range(len(cells)):
            v = rowvals[i] if i < len(rowvals) else None
            cells[i].text = "" if v in (None, "") else str(_cell(v))
            if widths and i < len(widths):
                cells[i].width = Cm(widths[i])
    if len(g.rows) > cap:
        doc.add_paragraph(
            f"… {len(g.rows) - cap} lignes supplémentaires dans le fichier "
            "« Contrôle Maquettes AVP ».",
            style="Intense Quote",
        )


def _write_stats_annex(doc, ctrl) -> None:
    stats = (ctrl.stats if ctrl else {}) or {}
    if not any(stats.values()):
        doc.add_paragraph(NOT_AVAILABLE)
        return
    for name, st in stats.items():
        if not st:
            continue
        p = doc.add_paragraph()
        p.add_run(name).bold = True
        if "conforme" in st:
            detail = (
                f"total {st.get('total')} · conformes {st.get('conforme')} "
                f"({_pct(st.get('conforme_ratio'))}) · non conformes {st.get('non_conforme')}"
            )
        else:
            detail = (
                f"total {st.get('total')} · sans matériau {st.get('non_conforme')} "
                f"({_pct(st.get('non_conforme_ratio'))})"
            )
        doc.add_paragraph(f"• {detail}", style="List Bullet")


def _write_ecarts(doc, result, sources, snap=None) -> None:
    env = sources.enveloppe if sources else None
    # SHAB snapshot avec le **même repli** que les annexes (BaseQuantities
    # puis « Superficie calculée ») — sinon l'écart reste NOT_AVAILABLE alors
    # que les annexes sont justes. Snapshot explicite prioritaire (audit non
    # encore lancé), sinon celui de l'``AuditResult``.
    eff_snap = snap if snap is not None else (result.snapshot if result is not None else None)
    ifc_shab = snapshot_shab_total(eff_snap)
    if ifc_shab is None and env is not None:
        ifc_shab = env.shab
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    for i, txt in enumerate(["Indicateur", "IFC OpenShell", "Méthode", "Observation"]):
        cell = tbl.rows[0].cells[i]
        cell.text = txt
        _shade_cell(cell, BIMDATA_PRIMARY)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
    row = tbl.add_row().cells
    row[0].text = "SHAB totale (m²)"
    row[1].text = f"{ifc_shab:.2f}" if isinstance(ifc_shab, (int, float)) else NOT_AVAILABLE
    row[2].text = "Espaces IFC : BaseQuantities, puis Superficie calculée"
    row[3].text = "Valeur issue de la maquette ; aucun écart d'outil externe n'est intégré."
    doc.add_paragraph(
        "Les valeurs métier intégrées au pack sont extraites de la maquette "
        "ou calculées via la chaîne IFC OpenShell ; les colonnes d'outil "
        "externe ne sont pas utilisées comme source de données.",
        style="Intense Quote",
    )
    for note in (_note_methodologie(env), _note_menuiseries(env)):
        if note:
            doc.add_paragraph(note, style="Intense Quote")


def _points_bloquants(ctrl, env, ratio, seuil) -> list[str]:
    out: list[str] = []
    if isinstance(ratio, (int, float)) and isinstance(seuil, (int, float)) and ratio < seuil:
        out.append(
            f"Ratio FAC/SHAB {ratio:.3f} inférieur au Seuil 3F 2026 ({seuil}) — enveloppe à revoir."
        )
    # Points de contrôle évalués 0 (Non fourni / non trouvé) dans la grille.
    if ctrl and ctrl.grille:
        try:
            eval_idx = ctrl.grille.headers.index("EVALUATION")
            pts_idx = ctrl.grille.headers.index("POINTS DE CONTROLE")
        except ValueError:
            eval_idx = pts_idx = None
        if eval_idx is not None:
            zeros = [
                r[pts_idx] for r in ctrl.grille.rows if str(r[eval_idx]).strip() in ("0", "0.0")
            ]
            for pt in zeros[:8]:
                out.append(f"Point de contrôle non satisfait (éval. 0) : {pt}")
    return out


def _recommandations(pieces, zones, materiau, ratio, seuil) -> list[str]:
    recs: list[str] = []
    for label, stat in (("pièces", pieces), ("zones", zones)):
        nc = stat.get("non_conforme") if stat else None
        if isinstance(nc, (int, float)) and nc > 0:
            recs.append(
                f"Reprendre le nommage de {int(nc)} {label} non conformes (CCH BIM I3F chap. 6.3)."
            )
    if (
        materiau
        and isinstance(materiau.get("non_conforme"), (int, float))
        and materiau["non_conforme"] > 0
    ):
        recs.append(
            f"Compléter le matériau sur {int(materiau['non_conforme'])} éléments ARC sans matériau."
        )
    if isinstance(ratio, (int, float)) and isinstance(seuil, (int, float)) and ratio < seuil:
        recs.append(
            "Revoir la modélisation de l'enveloppe pour atteindre le ratio FAC/SHAB attendu."
        )
    if not recs:
        recs.append("Aucune action corrective majeure identifiée à partir des livrables fournis.")
    return recs
