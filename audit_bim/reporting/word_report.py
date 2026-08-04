"""Rapport d'audit Word (livrable AMO BIM I3F).

Structure du document (refondue en 0.3 — modèle de rapport de conformité
de maquette numérique) :

1. Page de garde
   (Titre, Projet, Maquette auditée, Version, Date, Auteur, Référence CCBIM)
2. Synthèse exécutive (objectif, niveau de conformité, décision, indicateurs)
3. Périmètre de l'audit (documents de référence + maquette auditée)
4. Méthodologie (contrôles réalisés)
5. Résultats globaux (synthèse par domaine : Conforme / Avertissement / Non conforme)
6. Résultats détaillés
   6.1 Structure de la maquette
   6.2 Qualité des données
   6.3 Classification
   6.4 Conventions de nommage
   6.5 Contrôles géométriques
   6.6 Cohérence métier
   6.7 Détection des conflits
7. Liste des non-conformités
8. Recommandations (par priorité : Critique / Majeure / Mineure)
9. Conclusion (conformité globale, points bloquants, décision finale)
10. Annexes

Les sections contextuelles sont alimentées par
:class:`audit_bim.reporting.context.ReportProjectContext`. Si aucune
information n'est disponible pour une section donnée, la mention
« Information non disponible dans les documents fournis. » est
affichée — **on n'invente jamais**. Les contrôles non couverts par
l'audit automatisé (géométrie fine, détection de conflits, cohérence
métier détaillée) sont explicitement signalés comme hors périmètre
plutôt que présentés comme conformes.

Les graphes sont générés via matplotlib et insérés en PNG.
"""

from __future__ import annotations

import io
from collections import Counter
from collections.abc import Iterable, Sequence
from datetime import date
from pathlib import Path

import bim_reporting.charts as _bc
import bim_reporting.word as _bw
from bim_reporting.sections import (
    BrandSpec,
    cover_page,
    data_table,
    findings_table,
)
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from ..audit.engine import AuditResult
from ..audit.findings import Finding, Severity, Theme
from ..classifier import suggestions_map
from ..profiles import DEFAULT_PROFILE_ID, get_profile
from ..profiles.models import ReportNarrativeSpec
from .bimdata_brand import WORDMARK, find_logo
from .context import ReportProjectContext, build_report_context
from .theming import (
    BIMDATA_FONT_FALLBACK,
    BIMDATA_FONT_PRIMARY,
    BIMDATA_GRANITE,
    BIMDATA_PRIMARY,
    SEVERITY_COLORS,
    THEME_COLORS,
)

# ── Primitives déléguées au socle générique ``bim-reporting`` ───────────
# Ré-exports directs : mêmes objets, aucun changement de comportement.
_hex_to_rgb = _bw.hex_to_rgb
_shade_cell = _bw.shade_cell
_add_heading = _bw.add_heading
_section_break = _bw.section_break
_kpi_table = _bw.kpi_table
_para_intro = _bw.para_intro
_model_meta = _bw.model_meta
_plt = _bc.plt

# Phrase de fallback : utilisée chaque fois qu'une donnée contextuelle
# manque, pour éviter toute hallucination et garder un ton AMO BIM.
NOT_AVAILABLE = "Information non disponible dans les documents fournis."

# Mention pour les familles de contrôle non couvertes par l'audit
# automatisé (géométrie fine, clash detection, cohérence métier
# détaillée). On ne prétend JAMAIS qu'un contrôle non réalisé est conforme.
OUT_OF_SCOPE = (
    "Contrôle non réalisé dans le périmètre de cet audit automatisé "
    "(hors champ des données exposées par l'API BIMData)."
)

# Titre principal du livrable (page de garde).
REPORT_TITLE = "Rapport d'audit de conformité de la maquette numérique"

# Suffixe affiché en fin de valeur pour les données extraites des
# sources documentaires sans validation utilisateur. Indique au
# lecteur que la valeur est issue d'une déduction automatique et
# doit être confirmée par la MOA / MOE.
SOURCE_SUFFIX_EXTRACTED = "(déduit de la maquette — à confirmer)"
SOURCE_SUFFIX_DEDUCED = "(déduit par heuristique — à confirmer)"


def _render_with_source(value: str, source: str) -> str:
    """Suffixe de traçabilité selon la source du champ (socle générique).

    Les suffixes restent définis ici : ce sont des formulations qui s'impriment
    dans le livrable I3F.
    """
    return _bw.render_with_source(
        value,
        source,
        suffix_extracted=SOURCE_SUFFIX_EXTRACTED,
        suffix_deduced=SOURCE_SUFFIX_DEDUCED,
    )


MAX_FINDINGS_PER_THEME = 25  # cap par thème pour garder un rendu équilibré
MAX_NONCONFORMITIES = 80  # cap de la table « Liste des non-conformités »
PIE_OTHER_THRESHOLD = 0.02  # tranches < 2 % regroupées en « Autres »

# ── Mapping sévérité (5 niveaux) → gravité métier (4 niveaux) ──────────────
# L'échelle métier française du rapport est plus grossière que l'échelle
# technique du moteur ; on agrège HIGH/MEDIUM côté « Majeure » et LOW côté
# « Mineure ».
GRAVITY_FR = {
    Severity.CRITICAL: "Critique",
    Severity.HIGH: "Majeure",
    Severity.MEDIUM: "Majeure",
    Severity.LOW: "Mineure",
    Severity.INFO: "Information",
}

# Une non-conformité « opposable » est une anomalie de gravité au moins
# MEDIUM. Les LOW/INFO relèvent de l'avertissement qualité.
NONCONFORMITY_SEVERITIES = {Severity.CRITICAL, Severity.HIGH, Severity.MEDIUM}

# ── Mapping thèmes du moteur → domaines de la synthèse « Résultats globaux »
# Chaque domaine agrège un ou plusieurs ``Theme`` du moteur d'audit.
# Thèmes déclarés en TUPLES, pas en sets : ces collections ne servent
# aujourd'hui qu'à construire des mappings (ordre indifférent), mais un set qui
# traîne finit toujours par être itéré pour produire de la sortie — c'est
# exactement ce qui s'est passé dans ``_theme_block``.
DOMAINS: list[tuple[str, tuple[Theme, ...]]] = [
    ("Structure IFC / hiérarchie spatiale", (Theme.SPATIAL_HIERARCHY,)),
    (
        "Conventions de nommage",
        (Theme.NAMING_SITE_BAT_ETAGE, Theme.NAMING_ZONE, Theme.NAMING_SPACE),
    ),
    ("Classification", (Theme.CLASSIFICATION,)),
    ("Propriétés (Psets)", (Theme.PROPERTY_MISSING, Theme.PROPERTY_INVALID)),
    ("Quantités / géométrie", (Theme.QUANTITY,)),
    ("Documents attendus", (Theme.DOCUMENT,)),
]

# Statut domaine → libellé + couleur (charte feux tricolores).
_STATUS_LABEL = {
    "conforme": "✔ Conforme",
    "avertissement": "⚠ Avertissement",
    "non_conforme": "✖ Non conforme",
}
_STATUS_COLOR = {
    "conforme": "28A745",  # vert
    "avertissement": "FF8C00",  # orange
    "non_conforme": "DC3545",  # rouge
}

# Clés candidates pour extraire les métadonnées modèle depuis le dict
# BIMData ``get_model`` (les noms varient selon la version de l'API).
_MODEL_SOFTWARE_KEYS = ("source", "application", "software", "authoring_tool")
_MODEL_SCHEMA_KEYS = ("schema", "ifc_schema", "ifc_version", "version")
_MODEL_AUTHOR_KEYS = ("creator", "author", "created_by", "owner")
_MODEL_DATE_KEYS = ("created_at", "creation_date", "modified_date", "date")
_MODEL_DISCIPLINE_KEYS = ("type", "discipline", "domain")


def _pie_chart(values: dict[str, int], colors_map: dict[str, str], title: str) -> io.BytesIO:
    """Camembert (socle générique) — seuil et libellés du livrable figés ici."""
    return _bc.pie_chart(
        values,
        colors_map,
        title,
        other_label="Autres",
        empty_label="Aucune anomalie",
        other_threshold=PIE_OTHER_THRESHOLD,
    )


def _bar_chart(values: dict[str, int], colors_map: dict[str, str], title: str) -> io.BytesIO:
    """Barres (socle générique) — libellé d'axe du livrable figé ici."""
    return _bc.bar_chart(values, colors_map, title, y_label="Nb anomalies")


# ── Page de garde ─────────────────────────────────────────────────────────


def _empty_note(doc: Document, text: str) -> None:
    """Mention en italique quand une sous-section n'a rien à lister."""
    run = doc.add_paragraph().add_run(text)
    run.italic = True


def _findings_block(
    doc: Document,
    items: Iterable[Finding],
    suggestions_map: dict | None = None,
) -> None:
    """Projette des ``Finding`` I3F en tableau, rendu par le socle.

    L'adaptateur reste ici : il connaît le modèle ``Finding``, le jeu de
    colonnes du rapport I3F et les troncatures retenues. Le socle, lui, ne voit
    que des lignes de chaînes et une table de couleurs — c'est ce qui lui permet
    de servir un autre AMO sans hériter de ces choix.
    """
    items = list(items)
    with_sug = suggestions_map is not None
    headers = ["Sév.", "Classe IFC", "Élément", "Attendu", "Réel"]
    if with_sug:
        headers += ["Suggestion", "Conf."]

    rows: list[list[str]] = []
    for f in items:
        exp = f.expected
        if isinstance(exp, list):
            exp = ", ".join(map(str, exp[:5])) + ("…" if len(exp) > 5 else "")
        row = [
            f.severity.value,
            f.ifc_type or "",
            (f.name or f.element_uuid or "")[:40],
            str(exp or "")[:80],
            str(f.actual or "")[:60],
        ]
        if with_sug:
            sug = suggestions_map.get(f.element_uuid) if f.element_uuid else None
            row += (
                [f"{sug['code']} — {sug['label']}", f"{sug['confidence']:.2f}"] if sug else ["", ""]
            )
        rows.append(row)

    findings_table(
        doc,
        headers,
        rows,
        status_column=0,
        status_colors={f.severity.value: SEVERITY_COLORS[f.severity.value] for f in items},
        empty_text="Aucune anomalie pour ce thème.",
    )


def _narrative(profile_id: str | None = None) -> ReportNarrativeSpec | None:
    """Narratif du profil actif, ou ``None`` s'il n'en déclare pas.

    ``None`` n'est pas un cas dégradé à corriger : un profil préparatoire n'a
    pas encore écrit ses phrases, et lui prêter celles d'I3F imprimerait le
    référentiel d'un autre AMO dans son rapport.
    """
    return get_profile(profile_id or DEFAULT_PROFILE_ID).report_narrative


def _narrative_text(profile_id: str | None, field: str, fallback: str = "") -> str:
    """Un texte de narratif, avec repli neutre si le profil n'en déclare pas."""
    spec = _narrative(profile_id)
    return getattr(spec, field) if spec else fallback


def _framework_long_label(context: ReportProjectContext) -> str:
    """Forme développée du référentiel pour les phrases narratives.

    Reproduit l'ancienne tournure « Cahier des Charges BIM I3F V3.6 » sans citer
    de maître d'ouvrage : la forme développée vient du profil actif.
    """
    fw = context.reference_framework
    base = fw.long_name or fw.name
    if not base:
        return "référentiel —"
    return f"{base} V{fw.version}" if fw.version else base


# ── Helpers de calcul (décision, statuts domaine, métadonnées modèle) ──────


def _decision(result: AuditResult) -> tuple[str, str]:
    """Décision d'acceptation de la maquette selon les anomalies + le taux.

    Returns:
        ``(décision, justification)`` — décision parmi *Acceptée*,
        *Acceptée sous réserve*, *Refusée*.
    """
    by_sev = result.count_by_severity()
    n_crit = by_sev.get("CRITICAL", 0)
    n_high = by_sev.get("HIGH", 0)
    conf = result.conformity_rate() * 100
    if n_crit == 0 and n_high == 0 and conf >= 90:
        return ("Acceptée", "Maquette conforme aux exigences contrôlées.")
    if n_crit == 0 and conf >= 70:
        return (
            "Acceptée sous réserve",
            "Conforme sous réserve de correction des anomalies signalées.",
        )
    return ("Refusée", "Non conforme — corrections requises avant acceptation.")


def _domain_status(findings: list[Finding]) -> str:
    """Statut d'un domaine d'après la gravité max de ses anomalies."""
    sevs = {f.severity for f in findings}
    if Severity.CRITICAL in sevs or Severity.HIGH in sevs:
        return "non_conforme"
    if sevs:
        return "avertissement"
    return "conforme"


# ── Assemblage du rapport ──────────────────────────────────────────────────


def write_word_report(
    result: AuditResult,
    output_path: str | Path,
    auditor: str = "AMO BIM (audit automatisé)",
    xlsx_annex_path: str | Path | None = None,
    context: ReportProjectContext | None = None,
    *,
    profile_id: str | None = None,
) -> Path:
    """Génère le rapport Word d'audit (modèle de conformité de maquette).

    Args:
        result: ``AuditResult`` complet (snapshot + catalog + findings).
        output_path: Destination ``.docx`` (parents créés si nécessaire).
        auditor: Nom affiché sur la page de garde.
        xlsx_annex_path: Chemin de l'annexe XLSX (référencé en annexe).
        context: ``ReportProjectContext`` enrichi. Si ``None`` (défaut),
            on appelle :func:`build_report_context` pour le construire
            automatiquement depuis ``result``.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = BIMDATA_FONT_PRIMARY  # Roboto (cf. charte BIMData)
    style.font.size = Pt(10)
    style.font.color.rgb = _hex_to_rgb(BIMDATA_GRANITE)
    # rFonts pour propager la police à tous les scripts (ASCII, hAnsi, CS).
    style_rpr = style.element.get_or_add_rPr()
    rfonts = style_rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        style_rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), BIMDATA_FONT_PRIMARY)
    rfonts.set(qn("w:hAnsi"), BIMDATA_FONT_PRIMARY)
    rfonts.set(qn("w:cs"), BIMDATA_FONT_FALLBACK)

    # Contexte projet enrichi : auto-build si non fourni par le caller.
    if context is None:
        context = build_report_context(result, profile_id=profile_id)

    project_name = context.project_name or (result.snapshot.project or {}).get("name", "?")
    model_name = context.model_name or (result.snapshot.model or {}).get("name", "?")

    # L'auditeur affiché vient en priorité du contexte (fourni
    # explicitement par l'utilisateur), sinon du paramètre ``auditor``.
    display_auditor = context.auditor_name or auditor
    if not context.auditor_name and auditor and auditor != "AMO BIM (audit automatisé)":
        new_sources = dict(context.field_sources)
        new_sources["auditor_name"] = "user"
        context = context.model_copy(update={"auditor_name": auditor, "field_sources": new_sources})

    # ── 1. Page de garde ────────────────────────────────────────────────
    # Les libellés sont assemblés ICI : le socle ne connaît ni « maquette
    # auditée », ni « phase BIM », ni « référence du CCBIM ». Il reçoit des
    # couples (intitulé, valeur) et les met en page.
    cover_page(
        doc,
        title=REPORT_TITLE,
        subtitle=project_name,
        supertitle="AUDIT BIM",
        meta_title="Identification du livrable",
        meta_rows=[
            ("Projet", project_name),
            ("Maquette auditée", model_name),
            ("Version", f"Phase BIM {result.phase.value}"),
            ("Date", date.today().isoformat()),
            ("Auteur", display_auditor),
            (
                _narrative_text(profile_id, "cover_reference_label", "Référentiel"),
                context.reference_framework.label or "—",
            ),
        ],
        brand=BrandSpec(logo_path=find_logo("light"), wordmark=WORDMARK),
    )
    _section_break(doc)

    # ── 2. Synthèse exécutive ───────────────────────────────────────────
    _write_section_executive_summary(doc, result, context, project_name, model_name)
    _section_break(doc)

    # ── 3. Périmètre de l'audit ─────────────────────────────────────────
    _write_section_scope(doc, context, result, profile_id=profile_id)
    _section_break(doc)

    # ── 4. Méthodologie ─────────────────────────────────────────────────
    _write_section_methodology(doc, context)
    _section_break(doc)

    # ── 5. Résultats globaux ────────────────────────────────────────────
    _write_section_global_results(doc, result)
    _section_break(doc)

    # ── 6. Résultats détaillés ──────────────────────────────────────────
    _write_section_detailed_results(doc, result, profile_id=profile_id)
    _section_break(doc)

    # ── 7. Liste des non-conformités ────────────────────────────────────
    _write_section_nonconformities(doc, result)
    _section_break(doc)

    # ── 8. Recommandations ──────────────────────────────────────────────
    _write_section_recommendations(doc, result, profile_id=profile_id)
    _section_break(doc)

    # ── 9. Conclusion ───────────────────────────────────────────────────
    _write_section_conclusion(doc, result, context)
    _section_break(doc)

    # ── 10. Annexes ─────────────────────────────────────────────────────
    _write_section_annexes(doc, xlsx_annex_path, context, profile_id=profile_id)

    doc.save(str(output_path))
    return output_path


# ── Sections ───────────────────────────────────────────────────────────────


def _kv_or_na(
    doc: Document,
    label: str,
    value: str | None,
    *,
    source: str = "user",
) -> None:
    """Bullet « Label : valeur » avec fallback NOT_AVAILABLE (socle générique).

    ``NOT_AVAILABLE`` est passé explicitement : le socle refuse un défaut, cette
    phrase s'imprimant telle quelle dans le livrable.
    """
    _bw.kv_or_na(
        doc,
        label,
        value,
        source=source,
        not_available=NOT_AVAILABLE,
        suffix_extracted=SOURCE_SUFFIX_EXTRACTED,
        suffix_deduced=SOURCE_SUFFIX_DEDUCED,
    )


def _write_section_executive_summary(
    doc: Document,
    result: AuditResult,
    context: ReportProjectContext,
    project_name: str,
    model_name: str,
) -> None:
    """Section 2 — Synthèse exécutive."""
    _add_heading(doc, "2. Synthèse exécutive", level=1)

    by_sev = result.count_by_severity()
    by_theme = result.count_by_theme()
    conf = result.conformity_rate() * 100
    decision, justification = _decision(result)

    n_crit = by_sev.get("CRITICAL", 0)
    n_high = by_sev.get("HIGH", 0)
    n_med = by_sev.get("MEDIUM", 0)
    n_low = by_sev.get("LOW", 0)
    n_info = by_sev.get("INFO", 0)
    n_nonconf = n_crit + n_high + n_med
    n_warn = n_low + n_info
    n_rules = context.n_property_specs + context.n_naming_rules

    # Points de vigilance = thèmes les plus impactés (top 3).
    top_themes = sorted(by_theme.items(), key=lambda kv: -kv[1])[:3]
    vigilance = (
        ", ".join(f"{t} ({c})" for t, c in top_themes)
        if top_themes
        else "aucun écart significatif détecté"
    )

    doc.add_paragraph(
        f"L'audit vise à vérifier la conformité de la maquette « {model_name} » "
        f"(programme {project_name}, phase {result.phase.value}) au "
        f"{_framework_long_label(context)}. "
        f"Le niveau global de conformité (pondéré) s'établit à {conf:.0f} %. "
        f"Principaux points de vigilance : {vigilance}. "
        f"Décision : {decision} — {justification}",
        style="Intense Quote",
    )

    # Tableau d'indicateurs synthétiques.
    _kpi_table(
        doc,
        [
            ("Taux de conformité (pondéré)", f"{conf:.0f} %"),
            ("Éléments audités", str(context.n_elements)),
            ("Règles de conformité contrôlées (catalogue)", str(n_rules)),
            ("Non-conformités (Critique / Majeure)", str(n_nonconf)),
            ("Avertissements (Mineure / Information)", str(n_warn)),
            ("Décision", decision),
        ],
    )

    doc.add_paragraph(
        "Les deux figures ci-dessous synthétisent le profil global des "
        "anomalies : répartition par thème (quels domaines concentrent les "
        "écarts) et par sévérité (points bloquants vs améliorations de qualité).",
        style="Intense Quote",
    )
    doc.add_paragraph()
    doc.add_picture(
        _pie_chart(by_theme, THEME_COLORS, "Répartition des anomalies par thème"),
        width=Cm(13),
    )
    doc.add_picture(
        _bar_chart(
            {s.value: by_sev.get(s.value, 0) for s in Severity.ordered()},
            SEVERITY_COLORS,
            "Anomalies par sévérité",
        ),
        width=Cm(15),
    )


def _write_section_scope(
    doc: Document,
    context: ReportProjectContext,
    result: AuditResult,
    *,
    profile_id: str | None = None,
) -> None:
    """Section 3 — Périmètre de l'audit (documents de référence + maquette)."""
    _add_heading(doc, "3. Périmètre de l'audit", level=1)
    _para_intro(
        doc,
        "Cette section précise les documents de référence opposables et "
        "identifie la maquette auditée. Elle garantit la traçabilité du "
        "périmètre contrôlé.",
    )

    # 3.1 Documents de référence
    _add_heading(doc, "Documents de référence", level=2)
    src_cch = context.reference_framework.source or "non précisé"
    src_data = context.data_spec_source or "non précisé"
    src_naming = context.naming_spec_source or "non précisé"
    doc.add_paragraph(
        f"• {_narrative_text(profile_id, 'applied_reference_label', 'Référentiel')} : "
        f"{context.reference_framework.label or '—'} (Cahier des annexes — {src_cch}).",
        style="List Bullet",
    )
    doc.add_paragraph(
        f"• Convention / exigences BIM du maître d'ouvrage : annexe "
        f"« Spécification des données » ({src_data}) et annexe « Nommage » "
        f"({src_naming}).",
        style="List Bullet",
    )
    if context.n_property_specs or context.n_naming_rules:
        doc.add_paragraph(
            f"• Catalogue d'exigences chargé : {context.n_property_specs} "
            f"spécification(s) de propriétés et {context.n_naming_rules} "
            "règle(s) de nommage.",
            style="List Bullet",
        )

    # 3.2 Maquette auditée
    _add_heading(doc, "Maquette auditée", level=2)
    model = result.snapshot.model or {}
    _kv_or_na(doc, "Nom du modèle", context.model_name, source=context.source_of("model_name"))
    _kv_or_na(
        doc,
        "Discipline",
        _model_meta(model, _MODEL_DISCIPLINE_KEYS),
        source="extracted",
    )
    _kv_or_na(
        doc,
        "Auteur / producteur",
        _model_meta(model, _MODEL_AUTHOR_KEYS),
        source="extracted",
    )
    _kv_or_na(
        doc,
        "Date du modèle",
        _model_meta(model, _MODEL_DATE_KEYS),
        source="extracted",
    )
    _kv_or_na(
        doc,
        "Logiciel de production",
        _model_meta(model, _MODEL_SOFTWARE_KEYS),
        source="extracted",
    )
    _kv_or_na(
        doc,
        "Version IFC (schéma)",
        _model_meta(model, _MODEL_SCHEMA_KEYS),
        source="extracted",
    )
    _kv_or_na(
        doc,
        "Périmètre extrait",
        (
            f"{context.n_elements} éléments / {context.n_storeys} étage(s) / "
            f"{context.n_spaces} espace(s) / {context.n_zones} zone(s) — "
            "extraction BIMData"
        )
        if context.n_elements
        else None,
        source="extracted",
    )
    # Adresse / MOA si disponibles (utile pour le contexte projet).
    _kv_or_na(doc, "Adresse du projet", context.address, source=context.source_of("address"))
    moa_value = context.client_name or context.owner_name
    moa_source = (
        context.source_of("client_name") if context.client_name else context.source_of("owner_name")
    )
    _kv_or_na(doc, "Maîtrise d'ouvrage", moa_value, source=moa_source)
    _kv_or_na(
        doc,
        "Description du projet",
        context.project_description,
        source=context.source_of("project_description"),
    )


def _write_section_methodology(doc: Document, context: ReportProjectContext) -> None:
    """Section 4 — Méthodologie (description + tableau des contrôles)."""
    _add_heading(doc, "4. Méthodologie", level=1)
    _para_intro(
        doc,
        "L'audit est exécuté de façon automatisée à partir des données "
        "exposées par l'API BIMData et du catalogue d'exigences chargé. "
        "Les familles de contrôles réalisés sont décrites ci-dessous.",
    )
    doc.add_paragraph(
        "Contrôles réalisés : structure IFC et hiérarchie spatiale, "
        "conventions de nommage, classification, propriétés obligatoires "
        "(Psets par phase), validation des valeurs, quantités (surfaces / "
        "volumes), unicité des identifiants d'équipement et couverture des "
        "typologies attendues."
    )
    doc.add_paragraph(
        "Hors périmètre de cet audit automatisé : le contrôle géométrique "
        "fin (objets dupliqués, géométrie invalide), la cohérence métier "
        "détaillée et la détection de conflits (clash detection), qui "
        "requièrent l'analyse de la géométrie 3D non exposée par l'API."
    )

    if not context.controls_performed:
        doc.add_paragraph(NOT_AVAILABLE)
        return
    data_table(
        doc,
        ["Thème de contrôle", "Objectif", "Données contrôlées", "Source de la règle"],
        [
            [c.theme, c.objective, c.checked_items, c.rule_source or "—"]
            for c in context.controls_performed
        ],
    )


def _write_section_global_results(doc: Document, result: AuditResult) -> None:
    """Section 5 — Résultats globaux (synthèse par domaine)."""
    _add_heading(doc, "5. Résultats globaux", level=1)
    _para_intro(
        doc,
        "Vue d'ensemble du statut de conformité par domaine de contrôle. "
        "Un domaine est « Non conforme » s'il présente au moins une "
        "anomalie critique ou majeure, « Avertissement » pour des écarts "
        "mineurs, « Conforme » en l'absence d'anomalie.",
    )

    # Regrouper les findings par domaine.
    by_domain: dict[str, list[Finding]] = {label: [] for label, _ in DOMAINS}
    theme_to_domain: dict[Theme, str] = {}
    for label, themes in DOMAINS:
        for th in themes:
            theme_to_domain[th] = label
    for f in result.findings:
        label = theme_to_domain.get(f.theme)
        if label is not None:
            by_domain[label].append(f)

    statuses = [_domain_status(by_domain[label]) for label, _ in DOMAINS]
    rows = [
        [
            label,
            _STATUS_LABEL[status],
            str(len(by_domain[label])),
            str(
                sum(1 for f in by_domain[label] if f.severity in (Severity.CRITICAL, Severity.HIGH))
            ),
        ]
        for (label, _), status in zip(DOMAINS, statuses, strict=True)
    ]
    # La table de couleurs est indexée par le LIBELLÉ affiché : le socle ne
    # connaît ni nos statuts ni notre échelle de gravité.
    colors = {_STATUS_LABEL[s]: _STATUS_COLOR[s] for s in set(statuses)}
    findings_table(
        doc,
        ["Domaine", "Statut", "Nb anomalies", "Dont critiques/majeures"],
        rows,
        status_column=1,
        status_colors=colors,
    )


def _write_section_detailed_results(
    doc: Document, result: AuditResult, *, profile_id: str | None = None
) -> None:
    """Section 6 — Résultats détaillés (6.1 → 6.7)."""
    _add_heading(doc, "6. Résultats détaillés", level=1)
    _para_intro(
        doc,
        "Détail des écarts par famille de contrôle. Constitue la base de "
        "travail pour les corrections à mener dans la maquette ou les "
        "données sources. Le détail est limité aux anomalies les plus "
        "sévères par thème ; l'annexe Excel contient l'exhaustif.",
    )

    # Findings groupés par thème.
    by_theme_all: dict[Theme, list[Finding]] = {}
    for f in result.findings:
        by_theme_all.setdefault(f.theme, []).append(f)

    def _theme_block(themes: Sequence[Theme], *, with_suggestions: bool = False) -> None:
        """Rend le bloc d'un ou plusieurs thèmes, dans l'ordre DÉCLARÉ.

        ``themes`` est une séquence, pas un ensemble : l'itération d'un
        ``set[Theme]`` dépend du hash de ses membres, donc du processus. Deux
        rendus du même audit sortaient les lignes dans un ordre différent —
        sans erreur, sans que rien ne le signale. L'ordre d'écriture au site
        d'appel devient la référence, et il est lisible en relecture.
        """
        items: list[Finding] = []
        for th in themes:
            items.extend(by_theme_all.get(th, []))
        # Tri par sévérité (CRITICAL d'abord) puis cap. ``sort`` est stable :
        # à sévérité égale, l'ordre des thèmes ci-dessus est conservé.
        order = {s: i for i, s in enumerate(Severity.ordered())}
        items.sort(key=lambda f: order.get(f.severity, 99))
        smap = _suggestions_map(result) if with_suggestions else None
        _findings_block(doc, items[:MAX_FINDINGS_PER_THEME], suggestions_map=smap)

    # 6.1 Structure de la maquette
    _add_heading(doc, "6.1 Structure de la maquette", level=2)
    doc.add_paragraph(
        "Organisation IFC (Site → Bâtiment → Niveau → Espace) et présence "
        "des entités spatiales attendues."
    )
    _theme_block((Theme.SPATIAL_HIERARCHY,))

    # 6.2 Qualité des données
    _add_heading(doc, "6.2 Qualité des données", level=2)
    doc.add_paragraph(
        "Contrôle des propriétés obligatoires (Psets par phase) et de la "
        "validité des valeurs (présence, type, valeurs non vides)."
    )
    _theme_block((Theme.PROPERTY_MISSING, Theme.PROPERTY_INVALID))

    # 6.3 Classification
    _add_heading(doc, "6.3 Classification", level=2)
    doc.add_paragraph(
        _narrative_text(
            profile_id,
            "classification_intro",
            "Présence et cohérence de la classification IFC.",
        )
    )
    _theme_block((Theme.CLASSIFICATION,), with_suggestions=True)

    # 6.4 Conventions de nommage
    _add_heading(doc, "6.4 Conventions de nommage", level=2)
    doc.add_paragraph(
        _narrative_text(
            profile_id,
            "naming_intro",
            "Contrôle du nommage des objets, niveaux, zones et espaces.",
        )
    )
    _theme_block((Theme.NAMING_SITE_BAT_ETAGE, Theme.NAMING_ZONE, Theme.NAMING_SPACE))

    # 6.5 Contrôles géométriques
    _add_heading(doc, "6.5 Contrôles géométriques", level=2)
    quantity_items = by_theme_all.get(Theme.QUANTITY, [])
    doc.add_paragraph(
        "Présence des quantités géométriques (surfaces, volumes / "
        "BaseQuantities) sur les éléments quantifiables. Les contrôles "
        "géométriques fins (objets dupliqués, objets isolés, géométrie "
        "invalide, objets sans volume, intersections anormales) ne sont "
        "pas couverts par cet audit automatisé."
    )
    if quantity_items:
        order = {s: i for i, s in enumerate(Severity.ordered())}
        quantity_items = sorted(quantity_items, key=lambda f: order.get(f.severity, 99))
        _findings_block(doc, quantity_items[:MAX_FINDINGS_PER_THEME])
    else:
        # Italique désormais réelle : `Paragraph.italic` était un no-op, python-docx
        # ne l'écrivant jamais dans le XML. Correction d'une intention déjà là.
        _empty_note(doc, "Aucune anomalie de quantité détectée.")

    # 6.6 Cohérence métier
    _add_heading(doc, "6.6 Cohérence métier", level=2)
    doc.add_paragraph(
        "Cohérence métier par discipline (espaces fermés, portes dans les "
        "murs, fenêtres ; poteaux / poutres / dalles ; réseaux / "
        "équipements / connexions MEP). " + OUT_OF_SCOPE
    )

    # 6.7 Détection des conflits
    _add_heading(doc, "6.7 Détection des conflits", level=2)
    doc.add_paragraph(
        "Détection de conflits inter-disciplines (hard clash, soft clash, "
        "clearance). " + OUT_OF_SCOPE
    )


def _write_section_nonconformities(doc: Document, result: AuditResult) -> None:
    """Section 7 — Liste des non-conformités (tableau détaillé)."""
    _add_heading(doc, "7. Liste des non-conformités", level=1)
    _para_intro(
        doc,
        "Liste des anomalies opposables (gravité Critique ou Majeure). Les "
        "écarts mineurs et informationnels figurent dans l'annexe Excel.",
    )

    order = {s: i for i, s in enumerate(Severity.ordered())}
    ncs = sorted(
        (f for f in result.findings if f.severity in NONCONFORMITY_SEVERITIES),
        key=lambda f: order.get(f.severity, 99),
    )
    if not ncs:
        doc.add_paragraph("Aucune non-conformité critique ou majeure détectée.")
        return

    if len(ncs) > MAX_NONCONFORMITIES:
        doc.add_paragraph(
            f"⚠ {len(ncs)} non-conformités détectées — tableau limité aux "
            f"{MAX_NONCONFORMITIES} plus sévères ; l'exhaustif figure dans "
            "l'annexe Excel.",
            style="Intense Quote",
        )

    nc_rows: list[list[str]] = []
    for i, f in enumerate(ncs[:MAX_NONCONFORMITIES], start=1):
        exp = f.expected
        if isinstance(exp, list):
            exp = ", ".join(map(str, exp[:3])) + ("…" if len(exp) > 3 else "")
        comment = f"Attendu : {exp or '—'} / Réel : {f.actual or '—'}"
        nc_rows.append(
            [
                f"NC-{i:03d}",
                (f.ref_cch or f.error_type.value or "")[:40],
                f.short_label()[:40],
                GRAVITY_FR.get(f.severity, f.severity.value),
                comment[:90],
                (f.recommended_action or "—")[:90],
            ]
        )

    # Couleurs indexées par le libellé de gravité FR affiché, pas par l'enum :
    # le socle ne connaît pas notre échelle.
    gravity_colors = {
        GRAVITY_FR.get(sev, sev.value): SEVERITY_COLORS[sev.value]
        for sev in {f.severity for f in ncs[:MAX_NONCONFORMITIES]}
    }
    findings_table(
        doc,
        ["ID", "Règle", "Objet", "Gravité", "Commentaire", "Action"],
        nc_rows,
        status_column=3,
        status_colors=gravity_colors,
    )


def _write_section_recommendations(
    doc: Document, result: AuditResult, *, profile_id: str | None = None
) -> None:
    """Section 8 — Recommandations classées par priorité."""
    _add_heading(doc, "8. Recommandations", level=1)
    _para_intro(
        doc,
        "Actions correctives priorisées à mener avant le prochain dépôt de "
        "maquette. Les recommandations sont déduites des anomalies détectées.",
    )
    buckets = _recommendations_by_priority(result, profile_id=profile_id)
    any_rec = False
    for priority in ("Critique", "Majeure", "Mineure"):
        recs = buckets.get(priority, [])
        if not recs:
            continue
        any_rec = True
        _add_heading(doc, priority, level=2)
        for r in recs:
            doc.add_paragraph(r, style="List Bullet")
    if not any_rec:
        doc.add_paragraph("Aucune action corrective majeure ne semble nécessaire à ce stade.")


def _write_section_conclusion(
    doc: Document, result: AuditResult, context: ReportProjectContext
) -> None:
    """Section 9 — Conclusion (conformité globale, points bloquants, décision)."""
    _add_heading(doc, "9. Conclusion", level=1)
    by_sev = result.count_by_severity()
    conf = result.conformity_rate() * 100
    decision, justification = _decision(result)
    n_crit = by_sev.get("CRITICAL", 0)
    n_high = by_sev.get("HIGH", 0)
    n_blocking = n_crit + n_high

    # Domaines conformes (sans CRITICAL/HIGH) pour valoriser l'acquis.
    theme_to_domain: dict[Theme, str] = {}
    for label, themes in DOMAINS:
        for th in themes:
            theme_to_domain[th] = label
    severe_domains: set[str] = set()
    for f in result.findings:
        if f.severity in (Severity.CRITICAL, Severity.HIGH):
            d = theme_to_domain.get(f.theme)
            if d:
                severe_domains.add(d)
    conform_domains = [label for label, _ in DOMAINS if label not in severe_domains]

    if n_blocking:
        blocking_txt = (
            f"{n_blocking} point(s) bloquant(s) (anomalies critiques ou "
            "majeures) doivent être levés avant la prochaine livraison."
        )
    else:
        blocking_txt = "Aucun point bloquant (anomalie critique ou majeure) n'a été détecté."

    conform_txt = (
        f"Les domaines suivants sont conformes ou ne présentent que des "
        f"écarts mineurs : {', '.join(conform_domains)}. "
        if conform_domains
        else ""
    )

    doc.add_paragraph(
        f"La maquette « {context.model_name or '—'} » présente un niveau de "
        f"conformité (pondéré) de {conf:.0f} % au regard du "
        f"{context.reference_framework.label or '—'} pour la phase {result.phase.value}. "
        f"{conform_txt}{blocking_txt}"
    )
    doc.add_paragraph(
        "Actions avant la prochaine livraison : corriger en priorité les "
        "non-conformités critiques et majeures (cf. § 7 et § 8), puis "
        "ré-itérer un audit pour valider la reprise."
    )

    # Décision finale mise en valeur.
    p = doc.add_paragraph()
    run = p.add_run(f"Décision finale : {decision}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = _hex_to_rgb(BIMDATA_PRIMARY)
    doc.add_paragraph(justification)


def _write_section_annexes(
    doc: Document,
    xlsx_annex_path: str | Path | None,
    context: ReportProjectContext,
    *,
    profile_id: str | None = None,
) -> None:
    """Section 10 — Annexes."""
    _add_heading(doc, "10. Annexes", level=1)
    doc.add_paragraph(
        "• Liste complète des règles contrôlées et export exhaustif des "
        "résultats : voir l'annexe Excel.",
        style="List Bullet",
    )
    if xlsx_annex_path:
        doc.add_paragraph(
            f"• Annexe détaillée (Excel) : « {Path(xlsx_annex_path).name} » — "
            "intégralité des anomalies par type d'erreur, avec GUID IFC des "
            "objets concernés, exploitable directement par les équipes projet.",
            style="List Bullet",
        )
    doc.add_paragraph(
        f"• Paramètres d'exécution : phase BIM {context.project_phase or '—'}, "
        f"référentiel {context.reference_framework.label or '—'}, "
        f"{context.n_property_specs} spécification(s) de propriétés et "
        f"{context.n_naming_rules} règle(s) de nommage.",
        style="List Bullet",
    )
    reference_line = _narrative_text(profile_id, "reference_documents_line")
    if reference_line:
        doc.add_paragraph(reference_line, style="List Bullet")
    # Limites de l'audit (rattachées aux annexes).
    if context.assumptions:
        _add_heading(doc, "Limites et hypothèses de l'audit", level=2)
        for a in context.assumptions:
            doc.add_paragraph(f"• {a}", style="List Bullet")
    if context.missing_information:
        _add_heading(doc, "Informations non disponibles", level=2)
        doc.add_paragraph(
            "Éléments contextuels non extraits des sources analysées (ne "
            "constituent pas des anomalies de la maquette)."
        )
        for item in context.missing_information:
            doc.add_paragraph(f"• {item}", style="List Bullet")


# ── Génération des recommandations ─────────────────────────────────────────


def _suggestions_map(result: AuditResult) -> dict[str, dict]:
    """Suggestions de classification (1 par élément) pour le thème dédié."""
    return suggestions_map(result.findings, result.snapshot)


# Sévérité → priorité métier de la recommandation.
_SEV_TO_PRIORITY = {
    Severity.CRITICAL: "Critique",
    Severity.HIGH: "Critique",
    Severity.MEDIUM: "Majeure",
    Severity.LOW: "Mineure",
    Severity.INFO: "Mineure",
}


def _recommendations_by_priority(
    result: AuditResult, *, profile_id: str | None = None
) -> dict[str, list[str]]:
    """Recommandations correctives groupées par priorité (Critique / Majeure / Mineure).

    Pour chaque (priorité, thème), on agrège le nombre d'anomalies et on
    produit une action concrète à partir des ``theme_hints`` du profil.
    """
    spec = _narrative(profile_id)
    theme_hints = spec.theme_hints if spec else {}

    # (priority, theme) → count
    agg: dict[tuple[str, Theme], int] = Counter()
    for f in result.findings:
        priority = _SEV_TO_PRIORITY.get(f.severity, "Mineure")
        agg[(priority, f.theme)] += 1

    buckets: dict[str, list[str]] = {"Critique": [], "Majeure": [], "Mineure": []}
    # Tri stable : par priorité puis nombre décroissant.
    for (priority, theme), count in sorted(agg.items(), key=lambda kv: -kv[1]):
        hint = theme_hints.get(theme.value, "corriger les écarts identifiés")
        label = f"{count} anomalie{'s' if count > 1 else ''} — {hint}."
        buckets[priority].append(label[0].upper() + label[1:])

    # Recommandation transverse si conformité faible.
    if result.conformity_rate() < 0.7:
        transverse = _narrative_text(profile_id, "low_conformity_recommendation")
        if transverse:
            buckets["Critique"].append(transverse)
    return buckets
