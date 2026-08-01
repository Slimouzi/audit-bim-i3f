"""Pack de livrables AVP I3F (Tarare 0546L) — génération BIMData.

Produit, à partir d'un ``AuditResult`` (snapshot/audit BIMData) et des
quantités IFC extraites/calculées depuis la maquette, le pack de livrables AVP :

1. ``… Contrôle Maquettes.xlsx`` — grille de contrôle + stats conformité.
2. ``… AVP - export SHAB maquette.xlsx``.
3. ``… Export Zones et Espaces.xlsx``.
4. ``… Extraction surface enveloppe.xlsx`` (+ ratio FAC/SHAB, Seuil 3F).
5. ``… export Menuiseries.xlsx``.
6. ``… export plancher.xlsx``.
7. ``… Rapport analyse BIM.docx`` (+ ``.pdf`` best-effort) — rapport consolidé.

Principes :

- **Réutilise** l'infra de reporting existante : ``write_safe`` protège les
  cellules Excel contre l'injection ; les helpers ``word_report`` restent la
  base du Word consolidé. Pas de stack parallèle.
- **Ne jamais inventer** : donnée absente du snapshot / des calculs IFC →
  ``NOT_AVAILABLE``.
- **Maquette-first** pour les exports : SHAB, zones/espaces, enveloppe,
  menuiseries et plancher utilisent les valeurs extraites de la maquette ou
  calculées via la chaîne IFC/OpenShell. Les .xlsx MOA éventuellement fournis
  servent au contexte documentaire (identité projet, seuils, template futur),
  pas de source autoritaire pour les surfaces issues d'outils externes.
- **Fidélité « tables à plat »** : mêmes onglets, colonnes, ordre, unités
  et vocabulaire métier, avec colonnes IFC OpenShell explicites.
"""

from __future__ import annotations

from collections import Counter
from copy import copy
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path

import openpyxl
import xlsxwriter
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..audit.engine import AuditResult
from ..audit.findings import ErrorType, Theme
from ..extraction.model_data import ModelSnapshot
from ..requirements._openpyxl_compat import patch_openpyxl
from .avp_snapshot import (
    _SRC_COMPUTED,
    build_sources_from_snapshot,
    count_envelope_walls,
    count_menuiseries,
    count_planchers,
    snapshot_shab_total,
)
from .avp_sources import (
    ENVELOPPE_MOA_HEADERS,
    AvpSourcePaths,
    AvpSources,
    ControleMaquettesSource,
    SheetTable,
    load_sources,
)
from .context import ReportProjectContext
from .pdf_export import docx_to_pdf
from .theming import (
    BIMDATA_FONT_FALLBACK,
    BIMDATA_PRIMARY,
)
from .word_report import NOT_AVAILABLE, _kpi_table, _shade_cell
from .xlsx_annex import _build_formats, write_safe

# Les annexes AVP fournies par la MOA peuvent porter les mêmes CustomFilter
# invalides que les fichiers I3F — patch appliqué explicitement (plus d'effet
# de bord d'import depuis que les parseurs requirements sont paresseux).
patch_openpyxl()

# Mention normalisée (note méthodo) apposée sous les tables contenant des
# quantités calculées : distingue le natif BIMData du calcul IfcOpenShell.
_COMPUTED_METHODO_NOTE = (
    "Quantités partiellement calculées par analyse géométrique IfcOpenShell "
    "(colonne « Source quantité » = « Calculée (IfcOpenShell) ») — valeurs NON "
    "contractuelles, en attente d'un ré-export maquette avec BaseQuantities natives. "
    "Les quantités « Maquette » proviennent des BaseQuantities BIMData."
)


def _rows_have_computed(rows) -> bool:
    """Vrai si une table contient au moins une valeur « Calculée (IfcOpenShell) »."""
    return any(c == _SRC_COMPUTED for r in (rows or []) for c in (r or []))


# Convention de nommage documentaire I3F, **générée à partir de données
# projet confirmées** :
#
#     YYMMDD <NomProjet> <CodeProjet> <Phase> - <TypeLivrable>.<ext>
#
# ``YYMMDD`` = date de génération du livrable. Chaque livrable a un libellé
# de type et une extension fixes ; le nom du projet, le code (ESI) et la
# phase sont injectés depuis les valeurs confirmées par l'utilisateur.
_DELIVERABLE_LABELS: dict[str, tuple[str, str]] = {
    "controle": ("Contrôle Maquettes", "xlsx"),
    "shab": ("export SHAB maquette", "xlsx"),
    "zones_espaces": ("Export Zones et Espaces", "xlsx"),
    "enveloppe": ("Extraction surface enveloppe", "xlsx"),
    "menuiseries": ("export Menuiseries", "xlsx"),
    "plancher": ("export plancher", "xlsx"),
    "analyse": ("Rapport analyse BIM", "docx"),
}

# Caractères interdits / risqués dans un nom de fichier (séparateurs de
# chemin, caractères réservés Windows). Remplacés par un espace.
_FILENAME_BAD = '/\\:*?"<>|\r\n\t'


def _sanitize_filename_part(value: str | None) -> str:
    """Nettoie un fragment de nom de fichier (séparateurs, espaces)."""
    if not value:
        return ""
    out = "".join(" " if c in _FILENAME_BAD else c for c in str(value))
    return " ".join(out.split()).strip()


def _deliverable_filename(
    key: str, *, date: str, project_name: str, project_code: str, phase: str
) -> str:
    """Construit le nom d'un livrable selon la convention I3F.

    ``YYMMDD Nom Code Phase - TypeLivrable.ext`` — les fragments vides
    (code / phase absents) sont simplement omis (jamais inventés).
    """
    label, ext = _DELIVERABLE_LABELS[key]
    head_parts = [
        _sanitize_filename_part(date),
        _sanitize_filename_part(project_name),
        _sanitize_filename_part(project_code),
        _sanitize_filename_part(phase),
    ]
    head = " ".join(p for p in head_parts if p)
    label = _sanitize_filename_part(label)
    return f"{head} - {label}.{ext}"


_CONTROLE_STATS_SHEETS = (
    "Zones Nommage",
    "Pièces Nommage",
    "ARC bsence de matériau",
    "Zones ObjectType",
)


@dataclass
class AvpMeta:
    # Défauts **génériques** : aucune identité client codée en dur (le nom
    # et le code réels viennent des données confirmées / des sources I3F).
    project_name: str = "Projet"
    project_code: str = ""
    phase: str = "AVP"
    auditor: str = "AMO BIM"
    # Métadonnées opérationnelles du contrôle (issues du rapport I3F de
    # référence, fournies par l'appelant). Absentes → NOT_AVAILABLE, jamais
    # inventées.
    usages_bim: list[str] | None = None
    nombre_logements: str | None = None
    temoin_virtuel: str | None = None
    date_controle: str | None = None
    auteur_controle: str | None = None


@dataclass
class AvpReportPack:
    controle_xlsx: Path
    shab_xlsx: Path
    zones_espaces_xlsx: Path
    enveloppe_xlsx: Path
    menuiseries_xlsx: Path
    plancher_xlsx: Path
    analyse_docx: Path
    analyse_pdf: Path | None = None

    def paths(self) -> list[Path]:
        out = [
            self.controle_xlsx,
            self.shab_xlsx,
            self.zones_espaces_xlsx,
            self.enveloppe_xlsx,
            self.menuiseries_xlsx,
            self.plancher_xlsx,
            self.analyse_docx,
        ]
        if self.analyse_pdf is not None:
            out.append(self.analyse_pdf)
        return out


class AvpQaError(RuntimeError):
    """Livrable(s) client vide(s) alors que la maquette contient des données.

    Levée par la QA gate post-génération : un export sort sans aucune ligne
    métier alors que le snapshot expose des espaces / murs / zones
    exploitables. On refuse de livrer un fichier qui ne contient que le
    bandeau.
    """

    def __init__(self, empty: list[str]):
        self.empty = empty
        super().__init__(
            "Annexe(s) vide(s) malgré des données exploitables dans la maquette : "
            + ", ".join(empty)
            + ". Livraison refusée (ni sources I3F ni extraction snapshot n'ont "
            "produit de lignes)."
        )


# Marqueurs d'échafaudage à ignorer lors du comptage des lignes métier.
_QA_SCAFFOLD = {
    _n
    for _n in (
        NOT_AVAILABLE.strip().lower(),
        "(onglet vide dans la source i3f)",
        "synthèse",
    )
}


def _count_business_rows(path: Path) -> int:
    """Ouvre une annexe et compte ses **lignes métier**.

    Ignore le bandeau éventuel, la ligne d'en-tête de chaque onglet, les
    marqueurs d'échafaudage (``NOT_AVAILABLE``, onglet vide) et les blocs KPI.
    Sert de garde qualité anti-livrable vide.
    """
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return 0
    total = 0
    try:
        for ws in wb.worksheets:
            header_seen = False
            for row in ws.iter_rows(values_only=True):
                cells = [c for c in row if c not in (None, "")]
                if not cells:
                    continue
                first = str(cells[0]).strip().lower()
                if first == "synthèse":
                    break  # début du bloc KPI → stop pour cet onglet
                if not header_seen and any(str(c).strip().lower() == "composant" for c in cells):
                    header_seen = True
                    continue
                if not header_seen:
                    continue
                if first in _QA_SCAFFOLD:
                    continue
                if row[0] in (None, ""):
                    continue
                total += 1
    finally:
        wb.close()
    return total


# ── Helpers Excel (charte BIMData réutilisée) ──────────────────────────────


def _cell(v):
    """Valeur cellule sûre : blanc pour vide, date ISO, sinon brut."""
    if v is None:
        return ""
    if isinstance(v, datetime):
        return v.strftime("%Y-%m-%d")
    return v


def _write_banner(ws, fmts, supertitle: str, title: str) -> int:
    """Bannière BIMData (supertitle + filet jaune + titre). Renvoie la
    prochaine ligne libre."""
    write_safe(ws, 0, 0, f"BIMDATA — {supertitle}", fmts["supertitle"])
    write_safe(ws, 1, 0, "", fmts["accent_filet"])
    ws.set_row(1, 4)  # filet jaune fin
    write_safe(ws, 2, 0, title, fmts["title"])
    return 4


def _write_flat_table(ws, fmts, table: SheetTable | None, *, start_row: int) -> int:
    """Écrit une table à plat (en-têtes brandés + lignes zébrées).

    ``table is None`` → mention ``NOT_AVAILABLE``. Renvoie la ligne suivante.
    """
    if table is None or not table.headers:
        write_safe(ws, start_row, 0, NOT_AVAILABLE, fmts["row"])
        return start_row + 1
    for c, h in enumerate(table.headers):
        write_safe(ws, start_row, c, h, fmts["header"])
        ws.set_column(c, c, max(12, min(42, len(str(h)) + 3)))
    ws.set_row(start_row, 28)
    r = start_row
    for i, rowvals in enumerate(table.rows):
        r = start_row + 1 + i
        fmt = fmts["row_alt"] if i % 2 == 0 else fmts["row"]
        for c, v in enumerate(rowvals):
            write_safe(ws, r, c, _cell(v), fmt)
    ws.freeze_panes(start_row + 1, 0)
    return r + 1


def _new_workbook(path: Path):
    wb = xlsxwriter.Workbook(str(path), {"strings_to_formulas": False})
    return wb, _build_formats(wb)


# ── Builders des 5 Excel ───────────────────────────────────────────────────


def _build_controle_maquettes_xlsx(
    path, result, sources, meta, snap: ModelSnapshot | None = None
) -> Path:
    src = sources.controle if sources else None
    template_path = getattr(src, "template_path", None)
    if template_path:
        template = Path(template_path)
        if template.exists():
            return _build_controle_maquettes_template_xlsx(path, template, result, src, meta, snap)

    wb = xlsxwriter.Workbook(str(path), {"strings_to_formulas": False})
    fmts = _moa_formats(wb)
    ws = wb.add_worksheet("Grille de contrôle")
    _write_controle_moa_header(ws, fmts, src, meta)

    grille = _audit_controle_table(result)
    if grille is None and src is not None:
        grille = src.grille
    _write_controle_moa_table(ws, fmts, grille)

    # Onglets de stats conformité : synthèse KPI depuis l'audit.
    for name in _CONTROLE_STATS_SHEETS:
        ws_s = wb.add_worksheet(name[:31])
        grid = _controle_grid(name, src) if result is None else None
        if grid and grid.rows:
            _write_moa_grid(ws_s, fmts, grid.rows, start_row=0)
        else:
            rows = _controle_stat_rows(name, _controle_stats(name, result, src), meta)
            _write_moa_grid(ws_s, fmts, rows, start_row=0)
    wb.close()
    return path


def _build_controle_maquettes_template_xlsx(
    path: Path,
    template_path: Path,
    result: AuditResult | None,
    src: ControleMaquettesSource | None,
    meta: AvpMeta,
    snap: ModelSnapshot | None = None,
) -> Path:
    """Réutilise le classeur MOA Contrôle comme template vivant.

    Les onglets statistiques du fichier MOA contiennent des formules, listes de
    conformité et zones de copie. Le générateur ne doit pas les remplacer par un
    squelette : on conserve le template et on injecte uniquement l'identité
    projet + les listes maquette/audit disponibles.
    """
    wb = openpyxl.load_workbook(template_path, data_only=False)
    try:
        _ensure_control_template_sheets(wb)
        _refresh_template_metadata(wb, src, meta)
        effective_snap = (
            snap if snap is not None else (result.snapshot if result is not None else None)
        )
        if effective_snap is not None:
            _clear_template_grid_assessments(wb["Grille de contrôle"])
            _refresh_template_stats_tabs(wb, effective_snap)
        if result is not None:
            _append_audit_summary_to_template_grid(wb, result)
        _force_workbook_recalc(wb)
        wb.save(path)
    finally:
        wb.close()
    return path


def _ensure_control_template_sheets(wb) -> None:
    if "Grille de contrôle" not in wb.sheetnames:
        wb.create_sheet("Grille de contrôle", 0)
    for name in _CONTROLE_STATS_SHEETS:
        if name not in wb.sheetnames:
            wb.create_sheet(name)
    expected = ["Grille de contrôle", *_CONTROLE_STATS_SHEETS]
    ordered = [wb[name] for name in expected]
    ordered.extend(ws for ws in wb.worksheets if ws.title not in expected)
    wb._sheets = ordered


def _force_workbook_recalc(wb) -> None:
    calc = getattr(wb, "calculation", None)
    if calc is None:
        return
    try:
        calc.fullCalcOnLoad = True
        calc.forceFullCalc = True
    except AttributeError:
        return


def _openpyxl_safe_value(value):
    value = _cell(_moa_text(value))
    if isinstance(value, str) and value.startswith(("=", "+", "-", "@", "\t", "\r")):
        return "'" + value
    return value


def _text_or_none(value) -> str | None:
    if value in (None, ""):
        return None
    text = str(value).strip()
    return text or None


def _dict_text(item: dict, *keys: str) -> str | None:
    for key in keys:
        value = item.get(key)
        text = _text_or_none(value)
        if text:
            return text
    return None


def _metadata_cell(ws, label: str, fallback: str):
    wanted = label.strip().lower()
    for row in range(1, min(ws.max_row, 12) + 1):
        cell_label = ws.cell(row, 2).value
        if isinstance(cell_label, str) and cell_label.strip().lower().startswith(wanted):
            return ws.cell(row, 3)
    return ws[fallback]


def _refresh_template_metadata(wb, src: ControleMaquettesSource | None, meta: AvpMeta) -> None:
    if "Grille de contrôle" not in wb.sheetnames:
        return
    ws = wb["Grille de contrôle"]
    header = (src.header if src else {}) or {}
    values = {
        "Projet": meta.project_name or header.get("projet"),
        "ESI": meta.project_code or header.get("esi"),
        "Phase": meta.phase or header.get("phase"),
    }
    for label, value in values.items():
        if value not in (None, ""):
            _metadata_cell(
                ws, label, {"Projet": "C5", "ESI": "C6", "Phase": "C7"}[label]
            ).value = _openpyxl_safe_value(value)
    if meta.date_controle:
        _metadata_cell(ws, "Date d'analyse", "C9").value = _openpyxl_safe_value(meta.date_controle)


def _copy_cell_style(src, dst) -> None:
    if src.has_style:
        dst._style = copy(src._style)
    if src.number_format:
        dst.number_format = src.number_format
    if src.alignment:
        dst.alignment = copy(src.alignment)
    if src.protection:
        dst.protection = copy(src.protection)


def _copy_template_row_style(ws, src_row: int, dst_row: int, *, max_col: int = 6) -> None:
    for col in range(1, max_col + 1):
        _copy_cell_style(ws.cell(src_row, col), ws.cell(dst_row, col))


def _find_control_header_row(ws) -> int | None:
    for row in range(1, min(ws.max_row, 40) + 1):
        if any(
            isinstance(ws.cell(row, col).value, str)
            and ws.cell(row, col).value.strip().lower() == "points de controle"
            for col in range(1, ws.max_column + 1)
        ):
            return row
    return None


def _clear_template_grid_assessments(ws) -> None:
    header_row = _find_control_header_row(ws)
    if header_row is None:
        return
    for row in range(header_row + 1, ws.max_row + 1):
        if ws.cell(row, 2).value in (None, ""):
            continue
        for col in (5, 6):
            value = ws.cell(row, col).value
            if isinstance(value, str) and value.startswith("="):
                continue
            ws.cell(row, col).value = None


def _append_audit_summary_to_template_grid(wb, result: AuditResult) -> None:
    if "Grille de contrôle" not in wb.sheetnames:
        return
    ws = wb["Grille de contrôle"]
    header_row = _find_control_header_row(ws)
    if header_row is None:
        return
    _clear_template_grid_assessments(ws)
    grille = _audit_controle_table(result)
    if grille is None or not grille.rows:
        return
    start = ws.max_row + 2
    _copy_template_row_style(ws, header_row, start)
    ws.cell(start, 1).value = "Synthèse audit MCP"
    for col in range(2, 7):
        ws.cell(start, col).value = None
    headers = [
        "CODE 3F",
        "POINTS DE CONTROLE",
        "EXIGENCE CCH BIM 3F",
        "Outil utilisé",
        "EVALUATION",
        "Commentaires CdP Bim",
    ]
    _copy_template_row_style(ws, header_row, start + 1)
    for col, value in enumerate(headers, start=1):
        ws.cell(start + 1, col).value = value
    data_style_row = header_row + 1
    for idx, row_values in enumerate(_controle_rows_for_moa(grille), start=start + 2):
        _copy_template_row_style(ws, data_style_row, idx)
        for col, value in enumerate(row_values[: len(headers)], start=1):
            ws.cell(idx, col).value = _openpyxl_safe_value(value)


def _copy_zone_start_row(ws) -> int:
    for row in range(1, min(ws.max_row, 40) + 1):
        values = [ws.cell(row, col).value for col in range(1, min(ws.max_column, 5) + 1)]
        if any(isinstance(value, str) and "coller ici" in value.lower() for value in values):
            return row + 1
    return 13


def _clear_template_copy_zone(ws, *, start_row: int, end_row: int) -> None:
    for row in range(start_row, end_row + 1):
        for col in (1, 2, 3):
            ws.cell(row, col).value = None


def _write_template_counts(
    wb,
    sheet_name: str,
    counts: Counter[str],
    row_label: str,
    *,
    clear_when_empty: bool,
    end_row: int = 200,
) -> None:
    if sheet_name not in wb.sheetnames:
        return
    if not counts and not clear_when_empty:
        return
    ws = wb[sheet_name]
    start = _copy_zone_start_row(ws)
    clear_to = max(end_row, start + len(counts) + 2)
    _clear_template_copy_zone(ws, start_row=start, end_row=min(clear_to, max(ws.max_row, end_row)))
    for idx, (label, count) in enumerate(sorted(counts.items(), key=lambda item: item[0].lower())):
        row = start + idx
        ws.cell(row, 1).value = row_label if idx == 0 else None
        ws.cell(row, 2).value = _openpyxl_safe_value(label)
        ws.cell(row, 3).value = count


def _counter_from_dicts(
    items: list[dict] | None, *keys: str, empty_label: str = "(vide)"
) -> Counter[str]:
    counter: Counter[str] = Counter()
    for item in items or []:
        label = _dict_text(item, *keys) or empty_label
        counter[label] += 1
    return counter


def _refresh_template_stats_tabs(wb, snap: ModelSnapshot) -> None:
    zones = snap.zones or []
    spaces = snap.spaces or []
    elements = snap.elements or []
    _write_template_counts(
        wb,
        "Zones Nommage",
        _counter_from_dicts(zones, "name", "Name", empty_label="(Name vide)"),
        "Zones",
        clear_when_empty=True,
    )
    _write_template_counts(
        wb,
        "Pièces Nommage",
        _counter_from_dicts(
            spaces,
            "longname",
            "long_name",
            "LongName",
            "name",
            "Name",
            empty_label="(Name vide)",
        ),
        "Pièces",
        clear_when_empty=True,
    )
    _write_template_counts(
        wb,
        "Zones ObjectType",
        _counter_from_dicts(
            zones,
            "object_type",
            "objectType",
            "ObjectType",
            "predefined_type",
            "PredefinedType",
            empty_label="(ObjectType vide)",
        ),
        "Zones",
        clear_when_empty=True,
    )
    missing_materials = Counter(
        _dict_text(element, "type", "ifc_type", "ifc_class", "IFCType") or "(classe IFC absente)"
        for element in elements
        if not _has_material(element)
    )
    _write_template_counts(
        wb,
        "ARC bsence de matériau",
        missing_materials,
        "ObjT Pièces",
        clear_when_empty=True,
        end_row=206,
    )
    if "ARC bsence de matériau" in wb.sheetnames:
        ws = wb["ARC bsence de matériau"]
        if ws.max_row >= 7 and ws.max_column >= 7:
            ws.cell(7, 7).value = len(elements)


def _write_controle_moa_header(ws, fmts, src, meta) -> None:
    """Entête proche du classeur MOA « Contrôle Maquettes »."""
    header = (src.header if src else {}) or {}
    legend = (src.legend if src else {}) or {
        0: "Non fourni / non trouvé",
        1: "Insuffisant : à reprendre ou compléter",
        2: "Satisfaisant",
    }
    ws.set_column(0, 0, 12)
    ws.set_column(1, 1, 34)
    ws.set_column(2, 2, 54)
    ws.set_column(3, 3, 22)
    ws.set_column(4, 4, 14)
    ws.set_column(5, 5, 48)
    write_safe(ws, 1, 0, "2.         Grille de contrôle des exigences du CCH BIM 3F", fmts["title"])
    write_safe(
        ws,
        2,
        0,
        "Contrôle automatisé depuis l'audit BIM et les données extraites de la maquette.",
        fmts["note"],
    )
    fallbacks = {"projet": meta.project_name, "esi": meta.project_code, "phase": meta.phase}
    for offset, (label, key) in enumerate(
        (("Projet", "projet"), ("ESI", "esi"), ("Phase", "phase")),
        start=4,
    ):
        val = header.get(key)
        if val in (None, ""):
            val = fallbacks[key]
        write_safe(ws, offset, 1, label, fmts["meta_label"])
        write_safe(ws, offset, 2, _cell(val), fmts["meta_value"])
    for idx, code in enumerate(sorted(legend), start=6):
        write_safe(ws, idx, 4, code, fmts["center"])
        write_safe(ws, idx, 5, legend[code], fmts["note"])


def _write_controle_moa_table(ws, fmts, grille: SheetTable | None) -> None:
    headers = [
        "CODE 3F",
        "POINTS DE CONTROLE",
        "EXIGENCE CCH BIM 3F",
        "Outil utilisé",
        "EVALUATION",
        "Commentaires CdP Bim",
    ]
    start = 12
    _write_moa_grid(ws, fmts, [headers], start_row=start)
    if grille is None or not grille.rows:
        write_safe(ws, start + 1, 0, NOT_AVAILABLE, fmts["data"])
        return
    rows = _controle_rows_for_moa(grille)
    for i, row in enumerate(rows, start=start + 1):
        for c, value in enumerate(row[: len(headers)]):
            fmt = fmts["center"] if c == 4 else fmts["data"]
            _write_moa_value(ws, i, c, value, fmt)


def _controle_rows_for_moa(grille: SheetTable) -> list[list]:
    """Normalise une grille source/audit vers les 6 colonnes MOA."""
    headers = [str(h or "").strip() for h in grille.headers]
    if headers == [
        "CODE 3F",
        "POINTS DE CONTROLE",
        "EXIGENCE CCH BIM 3F",
        "Outil utilisé",
        "EVALUATION",
        "Commentaires CdP Bim",
    ]:
        return grille.rows
    rows: list[list] = []
    for raw in grille.rows:
        point = raw[0] if len(raw) > 0 else ""
        total = raw[1] if len(raw) > 1 else ""
        conformes = raw[2] if len(raw) > 2 else ""
        non_conformes = raw[3] if len(raw) > 3 else ""
        ratio = raw[4] if len(raw) > 4 else ""
        try:
            evaluation = 2 if float(non_conformes or 0) == 0 else 1
        except (TypeError, ValueError):
            evaluation = ""
        rows.append(
            [
                "",
                point,
                "Contrôle automatisé MCP audit-bim-i3f",
                "Audit BIM / IFC OpenShell",
                evaluation,
                (
                    f"Total: {total}; conformes: {conformes}; "
                    f"non conformes: {non_conformes}; taux: {ratio}"
                ),
            ]
        )
    return rows


def _controle_stat_rows(name: str, stats: dict | None, meta) -> list[list]:
    """Onglet statistique proche MOA, alimenté par l'audit quand disponible."""
    rows = [
        [],
        ["", meta.project_name],
        ["", meta.project_code],
        [],
        ["", name],
        [],
        ["", "Indicateur", "Total", "Conforme", "Taux Conforme", "Non Conforme", "Taux"],
    ]
    if not stats:
        rows.append(["", NOT_AVAILABLE])
        return rows
    rows.append(
        [
            "",
            stats.get("label") or name,
            stats.get("total"),
            stats.get("conforme"),
            stats.get("conforme_ratio"),
            stats.get("non_conforme"),
            stats.get("non_conforme_ratio"),
        ]
    )
    return rows


def _controle_grid(name: str, src):
    """Grille détaillée d'un onglet de contrôle (matching nom normalisé)."""
    if not src or not src.stat_grids:
        return None
    for key, grid in src.stat_grids.items():
        if _norm(key).replace("bsence", "absence") == _norm(name):
            return grid
    return None


def _controle_stats(name: str, result: AuditResult | None, src) -> dict | None:
    """Stats conformité d'un onglet — audit, puis source de repli sans snapshot."""
    if result is not None:
        return _audit_stats(name, result)
    if src and src.stats:
        # La clé source peut porter la faute d'origine ("ARC bsence…").
        for key, val in src.stats.items():
            if _norm(key) == _norm(name) or _norm(key).replace("bsence", "absence") == _norm(name):
                if val:
                    return val
    return None


def _zone_finding_kind(f) -> str | None:
    """Classe un finding de nommage d'IfcZone : ``"objecttype"`` ou ``"name"``.

    Le nommage de zone produit **deux contrôles distincts** partageant le même
    thème (``NAMING_ZONE``) : le **Name** (pattern XXXXL-YYYY, présence) et
    l'**ObjectType** (typologie dans la liste I3F, présence). On ne peut donc PAS
    agréger par thème.

    Discrimination **prioritaire** par le champ structuré ``field_path``
    (``"IfcZone.ObjectType"`` / ``"IfcZone.Name"``, bim-core ≥ 0.1.1) — fiable et
    **indépendant du libellé** du finding. Repli sur l'heuristique de wording
    uniquement pour les findings historiques sans ``field_path``.
    """
    if f.ifc_type != "IfcZone" or f.theme != Theme.NAMING_ZONE:
        return None
    # 1) Champ structuré (source de vérité) : un reformulage du wording des
    #    règles ne peut plus fausser silencieusement le classement.
    fp = (getattr(f, "field_path", None) or "").strip().lower()
    if fp.endswith(".objecttype"):
        return "objecttype"
    if fp.endswith(".name"):
        return "name"
    # 2) Repli heuristique (findings sans field_path).
    text = f"{f.recommended_action or ''} {f.expected or ''}".lower()
    if f.error_type == ErrorType.NAMING_NOT_IN_LIST or "objecttype" in text:
        return "objecttype"
    return "name"


def _material_name(item) -> str | None:
    """Nom de matériau exploitable d'un item de ``material_list`` (formes variées :
    ``{"material": {"name": …}}`` BIMData, ``{"name": …}``, ou chaîne)."""
    if isinstance(item, dict):
        mat = item.get("material")
        if isinstance(mat, dict) and mat.get("name") and str(mat["name"]).strip():
            return str(mat["name"]).strip()
        if item.get("name") and str(item["name"]).strip():
            return str(item["name"]).strip()
    elif isinstance(item, str) and item.strip():
        return item.strip()
    return None


def _has_material(e: dict) -> bool:
    """Vrai si l'élément porte un **vrai nom de matériau**.

    Lit ``material_list`` (forme produite par ``bimdata-read`` :
    ``[{"material": {"name": "Béton"}}]``) avec repli ``materials``. La simple
    présence de la clé ne suffit pas : il faut un nom non vide.
    """
    for key in ("material_list", "materials"):
        for item in e.get(key) or []:
            if _material_name(item):
                return True
    return False


def _naming_stat(total: int, nc: int) -> dict:
    conf = total - nc
    return {
        "label": "Nombre de Noms",
        "total": total,
        "conforme": conf,
        "conforme_ratio": conf / total if total else None,
        "non_conforme": nc,
        "non_conforme_ratio": nc / total if total else None,
    }


def _audit_stats(name: str, result: AuditResult) -> dict | None:
    snap = result.snapshot
    # Nommage zones : Name et ObjectType sont des contrôles SÉPARÉS (mêmes
    # zones, anomalies distinctes) — on ne compte pas le même ensemble deux fois.
    if name == "Zones Nommage":
        total = len(snap.zones or [])
        if total == 0:
            return None
        nc = len(
            {
                f.element_uuid
                for f in result.findings
                if _zone_finding_kind(f) == "name" and f.element_uuid
            }
        )
        return _naming_stat(total, nc)
    if name == "Zones ObjectType":
        total = len(snap.zones or [])
        if total == 0:
            return None
        nc = len(
            {
                f.element_uuid
                for f in result.findings
                if _zone_finding_kind(f) == "objecttype" and f.element_uuid
            }
        )
        return _naming_stat(total, nc)
    if name == "Pièces Nommage":
        total = len(snap.spaces or [])
        if total == 0:
            return None
        nc = len(
            {
                f.element_uuid
                for f in result.findings
                if f.theme == Theme.NAMING_SPACE and f.element_uuid
            }
        )
        return _naming_stat(total, nc)
    if "matériau" in name.lower():
        elements = snap.elements or []
        total = len(elements)
        if total == 0:
            return None
        sans = sum(1 for e in elements if not _has_material(e))
        conf = total - sans
        return {
            "label": "Nombre d'éléments sans matériau",
            "total": total,
            "conforme": conf,
            "conforme_ratio": conf / total if total else None,
            "non_conforme": sans,
            "non_conforme_ratio": sans / total if total else None,
        }
    return None


def _audit_controle_table(result: AuditResult | None) -> SheetTable | None:
    """Grille de contrôle **réelle dérivée de l'AuditResult** (aucune source I3F).

    Une ligne par point de contrôle effectivement évalué (nommage zones/pièces,
    ObjectType zones, matériaux), avec la conformité mesurée sur la maquette.
    Renvoie ``None`` si l'audit n'expose aucun point exploitable — auquel cas
    l'annexe Contrôle est vide et la QA gate lève (livrable non exploitable).
    """
    if result is None:
        return None
    rows: list[list] = []
    for name in _CONTROLE_STATS_SHEETS:
        stats = _audit_stats(name, result)
        if not stats:
            continue
        total = stats.get("total")
        nc = stats.get("non_conforme")
        conf = stats.get("conforme")
        if conf is None and total is not None and nc is not None:
            conf = total - nc
        ratio = stats.get("conforme_ratio")
        eval_value = 2 if isinstance(nc, (int, float)) and nc == 0 else 1
        rows.append(
            [
                "",
                name,
                "Contrôle automatisé MCP audit-bim-i3f",
                "Audit BIM / IFC OpenShell",
                eval_value,
                (
                    f"Total: {total if total is not None else ''}; "
                    f"conformes: {conf if conf is not None else ''}; "
                    f"non conformes: {nc if nc is not None else ''}; "
                    f"taux: {round(ratio, 3) if isinstance(ratio, (int, float)) else ''}"
                ),
            ]
        )
    if not rows:
        return None
    return SheetTable(
        title="Grille de contrôle",
        headers=[
            "CODE 3F",
            "POINTS DE CONTROLE",
            "EXIGENCE CCH BIM 3F",
            "Outil utilisé",
            "EVALUATION",
            "Commentaires CdP Bim",
        ],
        rows=rows,
    )


def _count_controle_rows(path: Path) -> int:
    """Compte les lignes de données sous l'en-tête MOA « POINTS DE CONTROLE ».

    Compteur **propre à l'annexe Contrôle** : contrairement à
    :func:`_count_business_rows`, il ignore l'entête projet, la légende, les
    titres et ``NOT_AVAILABLE`` — il ne compte que les points de contrôle réels
    de la grille. Sert de garde qualité fiable (0 = grille sans contrôle réel).
    """
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return 0
    try:
        if "Grille de contrôle" not in wb.sheetnames:
            return 0
        ws = wb["Grille de contrôle"]
        anchor = None
        for r_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if any(isinstance(c, str) and c.strip().lower() == "points de controle" for c in row):
                anchor = r_idx
                break
        if anchor is None:
            return 0
        total = 0
        for row in ws.iter_rows(min_row=anchor + 1, values_only=True):
            cells = [c for c in row if c not in (None, "")]
            if not cells:
                continue
            if str(cells[0]).strip() == NOT_AVAILABLE:
                continue
            if len(row) > 1 and row[1] not in (None, ""):
                total += 1
            elif row[0] not in (None, ""):
                # Repli pour une vieille grille 5 colonnes éventuelle.
                continue
        return total
    finally:
        wb.close()


def _write_stats_block(ws, fmts, stats: dict | None, *, start_row: int) -> None:
    if not stats:
        write_safe(ws, start_row, 0, NOT_AVAILABLE, fmts["row"])
        return
    # Structure « nommage » (avec conforme) ou « matériau » (sans conforme).
    if "conforme" in stats:
        labels = [
            ("Indicateur", "label"),
            ("Total", "total"),
            ("Conforme", "conforme"),
            ("Taux conforme", "conforme_ratio"),
            ("Non conforme", "non_conforme"),
            ("Taux non conforme", "non_conforme_ratio"),
        ]
    else:
        labels = [
            ("Indicateur", "label"),
            ("Total éléments", "total"),
            ("Sans matériau", "non_conforme"),
            ("Taux sans matériau", "non_conforme_ratio"),
        ]
    for c, (title, _key) in enumerate(labels):
        write_safe(ws, start_row, c, title, fmts["header"])
        ws.set_column(c, c, 20)
    for c, (_title, key) in enumerate(labels):
        v = stats.get(key)
        write_safe(ws, start_row + 1, c, "" if v is None else v, fmts["row_alt"])


def _looks_like_header(vals: list) -> bool:
    return sum(1 for v in vals if isinstance(v, str) and v.strip()) >= 3


def _write_grid(ws, fmts, rows: list[list], *, start_row: int) -> int:
    """Reproduit une grille brute (pivot/synthèse I3F) en table à plat.

    La 1re ligne « en-tête » (≥ 3 cellules texte) est stylée ; les autres
    sont zébrées. Préserve l'ordre et le contenu source.
    """
    if not rows:
        write_safe(ws, start_row, 0, NOT_AVAILABLE, fmts["row"])
        return start_row + 1
    header_idx = next((i for i, r in enumerate(rows) if _looks_like_header(r)), None)
    ncols = max(len(r) for r in rows)
    for c in range(ncols):
        ws.set_column(c, c, 18)
    r = start_row
    for i, rowvals in enumerate(rows):
        r = start_row + i
        if i == header_idx:
            fmt = fmts["header"]
            ws.set_row(r, 26)
        else:
            fmt = fmts["row_alt"] if i % 2 == 0 else fmts["row"]
        for c in range(ncols):
            v = rowvals[c] if c < len(rowvals) else None
            write_safe(ws, r, c, _cell(v), fmt)
    if header_idx is not None:
        ws.freeze_panes(start_row + header_idx + 1, 0)
    return r + 1


def _moa_formats(wb) -> dict[str, xlsxwriter.format.Format]:
    base = {"font_name": "Aptos Narrow", "font_size": 11}
    return {
        "title": wb.add_format({**base, "bold": True, "font_size": 13}),
        "note": wb.add_format({**base, "text_wrap": True}),
        "meta_label": wb.add_format({**base, "bold": True}),
        "meta_value": wb.add_format({**base}),
        "header": wb.add_format(
            {
                **base,
                "bold": True,
                "border": 1,
                "align": "center",
                "valign": "vcenter",
                "text_wrap": True,
            }
        ),
        "data": wb.add_format(
            {**base, "border": 1, "valign": "top", "text_wrap": True, "num_format": "#,##0.00"}
        ),
        "center": wb.add_format({**base, "border": 1, "align": "center", "valign": "vcenter"}),
        "percent": wb.add_format({**base, "border": 1, "num_format": "0.00%"}),
    }


_MOA_FORMULA_PREFIXES = ("=IF(", "=SUBTOTAL(", "=COUNTA(", "=SUM(")


def _moa_text(value):
    if not isinstance(value, str):
        return value
    return (
        value.replace("Surface (Solibri)", "Surface IFC OpenShell")
        .replace("Surface\nSolibri", "Surface IFC OpenShell")
        .replace("Surface Solibri", "Surface IFC OpenShell")
        .replace("Solibri Surface des Fenêtres", "IFC OpenShell Surface des Fenêtres")
        .replace("Solibri Surface des Portes", "IFC OpenShell Surface des Portes")
    )


def _is_moa_formula(value) -> bool:
    return isinstance(value, str) and value.startswith(_MOA_FORMULA_PREFIXES)


def _write_moa_value(ws, row: int, col: int, value, fmt) -> None:
    if _is_moa_formula(value):
        ws.write_formula(row, col, value, fmt, "")
    else:
        write_safe(ws, row, col, _cell(_moa_text(value)), fmt)


def _write_moa_grid(ws, fmts, rows: list[list], *, start_row: int = 0) -> int:
    """Écrit une grille façon MOA depuis A1, sans bannière ni freeze panes."""
    if not rows:
        return start_row
    ncols = max(len(r) for r in rows)
    header_idx = next((i for i, r in enumerate(rows) if _looks_like_header(r)), 0)
    for c in range(ncols):
        sample = [r[c] for r in rows[:30] if c < len(r) and r[c] not in (None, "")]
        width = max([len(str(_moa_text(v))) for v in sample] + [10])
        ws.set_column(c, c, min(max(width + 2, 10), 42))
    for i, rowvals in enumerate(rows):
        xls_row = start_row + i
        fmt = fmts["header"] if i == header_idx else fmts["data"]
        if i == header_idx:
            ws.set_row(xls_row, 42 if i == 0 else 28)
        for c in range(ncols):
            value = rowvals[c] if c < len(rowvals) else None
            cell_fmt = fmts["percent"] if _is_moa_percent_formula(value) else fmt
            _write_moa_value(ws, xls_row, c, value, cell_fmt)
    return start_row + len(rows)


def _is_moa_percent_formula(value) -> bool:
    return isinstance(value, str) and "/D" in value and value.startswith("=IF(")


def _build_multisheet_export_xlsx(path, banner: str, title: str, multi, meta) -> Path:
    """Export reproduisant **tous** les onglets source (pivots + détail)."""
    wb = xlsxwriter.Workbook(str(path), {"strings_to_formulas": False})
    fmts = _moa_formats(wb)
    grids = (multi.grids if multi else None) or []
    if not grids:
        ws = wb.add_worksheet(_safe_sheet(title))
        write_safe(ws, 0, 0, NOT_AVAILABLE, fmts["data"])
        wb.close()
        return path
    for g in grids:
        ws = wb.add_worksheet(_safe_sheet(g.title))
        if g.rows:
            end = _write_moa_grid(ws, fmts, g.rows, start_row=0)
            if _rows_have_computed(g.rows):
                write_safe(ws, end + 1, 0, _COMPUTED_METHODO_NOTE, fmts["note"])
    wb.close()
    return path


def _build_enveloppe_xlsx(path, sources, meta) -> Path:
    src = sources.enveloppe if sources else None
    wb = xlsxwriter.Workbook(str(path), {"strings_to_formulas": False})
    ws = wb.add_worksheet(
        _safe_sheet((src.sheet_title if src else None) or "TDB 2022 04.2 - Extraction s...")
    )
    fmts = _enveloppe_moa_formats(wb)

    rows = (src.table.rows if src and src.table else []) or []
    headers = (src.table.headers if src and src.table and src.table.headers else None) or list(
        ENVELOPPE_MOA_HEADERS
    )

    widths = [10.86, 47.29, 14.86, 14.57, 11.29, 14.0, 11.43, 11.86, 13.0, 13.0]
    for c, width in enumerate(widths):
        ws.set_column(c, c, width)
    ws.set_row(0, 60)

    for c, h in enumerate(headers[:10]):
        write_safe(
            ws, 0, c, _enveloppe_output_header(h), fmts["header_calc" if c == 3 else "header"]
        )

    for r_idx, rowvals in enumerate(rows, start=1):
        for c in range(10):
            value = rowvals[c] if c < len(rowvals) else None
            write_safe(ws, r_idx, c, _cell(value), fmts["data"])

    n_data = len(rows)
    first_data = 2
    last_data = max(first_data, n_data + 1)
    formula_end_with_blank = last_data + 1
    summary_row = n_data + 2  # Tarare : 8 lignes → ligne Excel 11.

    d_total = _sum_table_col(rows, 3)
    e_total = _sum_table_col(rows, 4)
    f_total = _sum_table_col(rows, 5)
    men_total = _first_number(src.superficie_menuiseries if src else None, f_total)
    shab = _first_number(src.shab if src else None)
    ratio = _first_number(src.ratio_fac_shab if src else None)
    if ratio is None and shab:
        ratio = d_total / shab if d_total is not None else None
    ecart = (e_total / d_total - 1) if d_total else None

    write_safe(ws, summary_row, 2, "Superficie des façades : ", fmts["summary_label_top"])
    ws.write_formula(
        summary_row,
        3,
        f"=SUM(D{first_data}:D{formula_end_with_blank})",
        fmts["summary_value_top"],
        _formula_cached(d_total),
    )
    ws.write_formula(
        summary_row,
        4,
        f"=SUM(E{first_data}:E{formula_end_with_blank})",
        fmts["summary_value"],
        _formula_cached(e_total),
    )
    ws.write_formula(
        summary_row,
        5,
        f"=SUM(F{first_data}:F{last_data})",
        fmts["summary_value"],
        _formula_cached(f_total),
    )
    write_safe(
        ws,
        summary_row,
        6,
        "non pertinent : inclus portes et fenêtres",
        fmts["note"],
    )

    write_safe(ws, summary_row + 1, 3, "écart : ", fmts["label"])
    ws.write_formula(
        summary_row + 1,
        4,
        f"=E{summary_row + 1}/D{summary_row + 1}-1",
        fmts["percent_fill"],
        _formula_cached(ecart),
    )
    write_safe(ws, summary_row + 2, 6, "écart calcul IFC OpenShell à contrôler", fmts["note"])

    write_safe(ws, summary_row + 3, 2, "Superficie des menuiseries : ", fmts["summary_label"])
    write_safe(
        ws,
        summary_row + 3,
        3,
        NOT_AVAILABLE if men_total is None else men_total,
        fmts["summary_value_top"],
    )
    write_safe(ws, summary_row + 3, 6, "(murs complexes : murs non découpés ", fmts["note"])
    write_safe(ws, summary_row + 4, 6, "sur 100% de la périphérie )", fmts["note"])

    write_safe(ws, summary_row + 5, 2, "SHAB : ", fmts["summary_label"])
    write_safe(
        ws, summary_row + 5, 3, NOT_AVAILABLE if shab is None else shab, fmts["summary_value_top"]
    )
    write_safe(ws, summary_row + 6, 2, "ratio FAC/SHAB : ", fmts["summary_label_plain"])
    ws.write_formula(
        summary_row + 6,
        3,
        f"=D{summary_row + 1}/D{summary_row + 6}",
        fmts["ratio"],
        _formula_cached(ratio),
    )
    write_safe(ws, summary_row + 7, 2, "Seuil 3F 2026 : ", fmts["summary_label_plain"])
    write_safe(
        ws,
        summary_row + 7,
        3,
        NOT_AVAILABLE if not src or src.seuil_3f is None else src.seuil_3f,
        fmts["threshold"],
    )
    wb.close()
    return path


def _enveloppe_moa_formats(wb) -> dict[str, xlsxwriter.format.Format]:
    base = {"font_name": "Aptos Narrow", "font_size": 11}
    return {
        "header": wb.add_format(
            {**base, "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True}
        ),
        "header_calc": wb.add_format(
            {
                **base,
                "bold": True,
                "border": 1,
                "align": "center",
                "valign": "vcenter",
                "text_wrap": True,
                "bg_color": "#D9EAD3",
            }
        ),
        "data": wb.add_format({**base, "left": 1, "right": 1, "top": 1, "bottom": 7}),
        "summary_label_top": wb.add_format(
            {**base, "bold": True, "align": "right", "top": 1, "bottom": 1}
        ),
        "summary_label": wb.add_format(
            {**base, "bold": True, "align": "right", "right": 1, "top": 1, "bottom": 1}
        ),
        "summary_label_plain": wb.add_format({**base, "bold": True, "align": "right"}),
        "summary_value_top": wb.add_format(
            {**base, "bold": True, "border": 1, "bg_color": "#D9EAD3", "num_format": "#,##0.00"}
        ),
        "summary_value": wb.add_format({**base, "border": 1, "num_format": "#,##0.00"}),
        "label": wb.add_format({**base}),
        "percent_fill": wb.add_format(
            {**base, "border": 1, "bg_color": "#D9EAD3", "num_format": "0.00%"}
        ),
        "note": wb.add_format({**base}),
        "ratio": wb.add_format(
            {**base, "bold": True, "font_color": "#C00000", "num_format": "#,##0.00"}
        ),
        "threshold": wb.add_format({**base, "num_format": "0.00"}),
    }


def _enveloppe_output_header(value) -> str:
    text = str(value or "")
    return (
        text.replace("Surface Solibri", "Surface IFC OpenShell")
        .replace("Solibri Surface des Fenêtres", "IFC OpenShell Surface des Fenêtres")
        .replace("Solibri Surface des Portes", "IFC OpenShell Surface des Portes")
    )


def _first_number(*values) -> float | None:
    for value in values:
        if isinstance(value, (int, float)):
            return float(value)
    return None


def _formula_cached(value):
    return value if isinstance(value, (int, float)) else ""


def _sum_table_col(rows, idx: int) -> float:
    return round(sum(r[idx] for r in rows if len(r) > idx and isinstance(r[idx], (int, float))), 2)


def _env_ecart(src) -> float | None:
    """Écart Σ « Surface IFC OpenShell » (col E) − Σ « Archicad BQ NetSideArea »
    (col D) de la table enveloppe. Nul quand une source unique alimente D et E."""
    if not src or not getattr(src, "table", None) or not src.table.rows:
        return None
    d = sum(r[3] for r in src.table.rows if len(r) > 3 and isinstance(r[3], (int, float)))
    e = sum(r[4] for r in src.table.rows if len(r) > 4 and isinstance(r[4], (int, float)))
    return round(e - d, 2)


def _build_menuiseries_xlsx(path, sources, meta) -> Path:
    src = sources.menuiseries if sources else None
    wb = xlsxwriter.Workbook(str(path), {"strings_to_formulas": False})
    fmts = _moa_formats(wb)
    # Proximité I3F : conserver le nom d'onglet source (« TDB 2022 05.1… »).
    ws = wb.add_worksheet(_safe_sheet((src.sheet_title if src else None) or "Menuiseries"))
    table = src.table if src else None
    if table is None or not table.headers:
        write_safe(ws, 0, 0, NOT_AVAILABLE, fmts["data"])
        wb.close()
        return path

    rows = [list(table.headers), *table.rows]
    ncols = max(len(r) for r in rows) if rows else len(table.headers)
    rows.append([""] * ncols)
    summary = [""] * ncols
    if ncols > 1:
        summary[1] = "Nombre de types de menuiseries"
    if ncols > 2:
        last = max(2, len(table.rows) + 1)
        summary[2] = f"=COUNTA(D2:D{last})" if table.rows else (src.nombre_types if src else None)
    rows.append(summary)
    end = _write_moa_grid(ws, fmts, rows, start_row=0)
    if _rows_have_computed(table.rows):
        write_safe(ws, end + 1, 0, _COMPUTED_METHODO_NOTE, fmts["note"])
    wb.close()
    return path


def _build_plancher_xlsx(path, sources, meta) -> Path:
    """Export plancher (dalles ``IfcSlab``) — **multi-onglets** comme SHAB/Zones.

    Reproduit **tous** les onglets de la source I3F (« … Dalles Ok », « Planchers »
    avec totaux/écarts) ; repli maquette (un seul onglet « Planchers ») câblé par
    l'orchestrateur. Aucune valeur inventée : onglet source vide préservé tel quel.
    """
    return _build_multisheet_export_xlsx(
        path,
        "EXPORT PLANCHER",
        "Export plancher",
        (sources.plancher if sources else None),
        meta,
    )


# ── Consolidé « Analyse BIM AVP » (.docx, helpers word_report réutilisés) ───


_MOA_FONT = "Calibri"
_MOA_TABLE_HEADER = "D9EAD3"
_MOA_TABLE_SECTION = "EAF2F8"


def _setup_docx() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _MOA_FONT
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _MOA_FONT)
    rfonts.set(qn("w:hAnsi"), _MOA_FONT)
    rfonts.set(qn("w:cs"), BIMDATA_FONT_FALLBACK)
    _configure_section(doc.sections[0], WD_ORIENT.PORTRAIT)
    return doc


def _pct(v) -> str:
    return f"{v * 100:.0f} %" if isinstance(v, (int, float)) else NOT_AVAILABLE


def _build_analyse_bim_avp_docx(
    path, result, sources, meta, snap=None, controle_xlsx: Path | None = None
) -> Path:
    doc = _setup_docx()
    ctrl = sources.controle if sources else None
    env = sources.enveloppe if sources else None
    pieces = _stat_lookup(ctrl, "Pièces Nommage")
    zones = _stat_lookup(ctrl, "Zones Nommage")
    materiau = _stat_lookup(ctrl, "ARC absence de matériau")

    _write_moa_cover_page(doc, ctrl, meta)
    doc.add_page_break()
    _write_moa_contents_and_inputs(doc, ctrl, meta, result, sources, snap)
    _write_audit_synthese(doc, result)
    _write_computed_coverage(doc, snap)
    _write_ecarts(doc, result, sources, snap)
    _write_moa_indicateurs(doc, pieces, zones, materiau, env)

    _set_orientation(doc, WD_ORIENT.LANDSCAPE)
    _write_moa_control_grid_pages(doc, ctrl, meta, controle_xlsx)
    _write_moa_control_annex_pages(
        doc,
        ctrl,
        meta,
        controle_xlsx,
        use_source_cached_values=(snap is None and result is None),
    )

    doc.save(str(path))
    return path


def _configure_section(section, orient) -> None:
    section.orientation = orient
    if orient == WD_ORIENT.LANDSCAPE:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.left_margin = Cm(0.9)
        section.right_margin = Cm(0.9)
        section.top_margin = Cm(0.8)
        section.bottom_margin = Cm(0.8)
    else:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(1.7)
        section.right_margin = Cm(1.7)
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)


def _set_docx_run_font(run, *, size: float = 10, bold: bool | None = None) -> None:
    run.font.name = _MOA_FONT
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _MOA_FONT)
    rfonts.set(qn("w:hAnsi"), _MOA_FONT)
    rfonts.set(qn("w:cs"), BIMDATA_FONT_FALLBACK)


def _add_moa_paragraph(
    doc: Document,
    text: str = "",
    *,
    size: float = 10,
    bold: bool = False,
    align=None,
    space_after: float = 3,
):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_docx_run_font(run, size=size, bold=bold)
    return p


def _write_moa_title(doc: Document, text: str, *, size: float = 14) -> None:
    p = _add_moa_paragraph(doc, text, size=size, bold=True, space_after=6)
    p.paragraph_format.keep_with_next = True


def _write_moa_cover_page(doc: Document, ctrl, meta: AvpMeta) -> None:
    _add_moa_paragraph(
        doc,
        "Rapport d’analyse des maquettes numériques",
        size=12,
        bold=True,
        space_after=38,
    )
    _add_moa_paragraph(
        doc,
        "Rapport d’analyse\ndes maquettes numériques – 3F",
        size=22,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=22,
    )
    operation = _operation_label(meta)
    _add_moa_paragraph(
        doc,
        operation,
        size=13,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=8,
    )
    _add_moa_paragraph(
        doc,
        f"Programme {meta.project_code}" if meta.project_code else "Programme",
        size=12,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=28,
    )
    _add_moa_paragraph(
        doc,
        f"Les maquettes IFC remises pour le rendu {meta.phase} ont permis une analyse "
        f"partielle de la phase {meta.phase} :",
        size=10,
        space_after=6,
    )
    _write_moa_usages(doc, meta)
    _write_moa_phase_table(doc, ctrl, meta)


def _operation_label(meta: AvpMeta) -> str:
    if meta.nombre_logements and meta.project_name:
        return f"Opération de construction de {meta.nombre_logements} à {meta.project_name}"
    if meta.project_name and meta.project_code:
        return f"Opération {meta.project_name} – Programme {meta.project_code}"
    if meta.project_name:
        return f"Opération {meta.project_name}"
    return "Opération"


def _write_moa_usages(doc: Document, meta: AvpMeta) -> None:
    if meta.usages_bim:
        for usage in meta.usages_bim:
            _add_moa_paragraph(doc, str(usage), size=10, space_after=2)
    else:
        _add_moa_paragraph(doc, "Usages BIM 3F : " + NOT_AVAILABLE + ".", size=10, space_after=2)
    if meta.temoin_virtuel:
        _add_moa_paragraph(doc, str(meta.temoin_virtuel), size=10, space_after=12)


def _write_moa_phase_table(doc: Document, ctrl, meta: AvpMeta) -> None:
    h = (ctrl.header if ctrl else {}) or {}
    date_controle = meta.date_controle or h.get("date d'analyse")
    auteur = meta.auteur_controle or meta.auditor
    rows = [
        ["Phase", "Date du contrôle", "Contrôle effectué par :"],
        [meta.phase, _fmt_meta(date_controle), _fmt_meta(auteur)],
    ]
    _write_docx_table(doc, rows, widths=(3.0, 5.0, 8.0), font_size=10, header_rows=1)


def _write_moa_contents_and_inputs(
    doc: Document, ctrl, meta: AvpMeta, result, sources, snap=None
) -> None:
    _add_moa_paragraph(doc, "Rapport d’analyse des maquettes numériques", size=11, bold=True)
    _write_moa_title(doc, "Table des matières", size=16)
    for label, page in (
        ("1. Données d’entrées", "2"),
        ("2. Grille de contrôle", "3"),
        ("3. Annexes de contrôle", "4"),
    ):
        _add_moa_paragraph(doc, f"{label} {' .' * 42} {page}", size=10, space_after=1)
    _write_moa_title(doc, "1. Données d’entrées", size=14)
    h = (ctrl.header if ctrl else {}) or {}
    n_models = _snapshot_model_count(
        snap if snap is not None else (result.snapshot if result else None)
    )
    _add_moa_paragraph(doc, _models_transmitted_label(n_models), size=10, space_after=3)
    _add_moa_paragraph(doc, "Exports IFC réalisés le :", size=10, space_after=2)
    export_date = h.get("maquettes ifc transmises le") or h.get("date d'analyse")
    model_name = _snapshot_model_name(
        snap if snap is not None else (result.snapshot if result else None)
    )
    if export_date not in (None, ""):
        _add_moa_paragraph(doc, f"{_fmt_meta(export_date)} : {model_name}", size=10, space_after=6)
    else:
        _add_moa_paragraph(doc, NOT_AVAILABLE, size=10, space_after=6)
    _write_donnees_entree(doc, ctrl, meta)
    _write_moa_title(doc, "Synthèse technique MCP", size=12)


def _snapshot_model_count(snap) -> int | None:
    if snap is None:
        return None
    models = getattr(snap, "models", None)
    if isinstance(models, list) and models:
        return len(models)
    model = getattr(snap, "model", None)
    return 1 if model else None


def _models_transmitted_label(n_models: int | None) -> str:
    if not isinstance(n_models, int):
        return "Maquettes transmises : " + NOT_AVAILABLE
    return (
        f"{n_models} maquette{'s' if n_models > 1 else ''} transmise{'s' if n_models > 1 else ''}"
    )


def _snapshot_model_name(snap) -> str:
    model = getattr(snap, "model", None) if snap is not None else None
    if isinstance(model, dict):
        name = model.get("name")
        if name:
            return str(name)
    return "maquette IFC"


def _write_moa_indicateurs(doc: Document, pieces, zones, materiau, env) -> None:
    ratio = env.ratio_fac_shab if env else None
    seuil = env.seuil_3f if env else None
    if isinstance(ratio, (int, float)) and isinstance(seuil, (int, float)):
        ratio_ok = "Conforme" if ratio >= seuil else "Non conforme"
    else:
        ratio_ok = NOT_AVAILABLE
    rows = [
        ["Indicateur", "Valeur"],
        [
            "Taux de conformité nommage pièces",
            _pct(pieces.get("conforme_ratio")) if pieces else NOT_AVAILABLE,
        ],
        [
            "Taux de conformité nommage zones",
            _pct(zones.get("conforme_ratio")) if zones else NOT_AVAILABLE,
        ],
        [
            "Éléments sans matériau (taux)",
            _pct(materiau.get("non_conforme_ratio")) if materiau else NOT_AVAILABLE,
        ],
        ["Ratio FAC/SHAB", f"{ratio:.3f}" if isinstance(ratio, (int, float)) else NOT_AVAILABLE],
        [
            f"Seuil 3F 2026 (≥ {seuil})" if isinstance(seuil, (int, float)) else "Seuil 3F 2026",
            ratio_ok,
        ],
    ]
    _write_docx_table(doc, rows, widths=(7.5, 4.0), font_size=8.5, header_rows=1)


def _write_moa_control_grid_pages(
    doc: Document, ctrl, meta: AvpMeta, controle_xlsx: Path | None
) -> None:
    _write_moa_title(doc, "3. Grille de contrôle des exigences du CCH BIM 3F", size=13)
    _add_moa_paragraph(
        doc,
        "Les maquettes numériques ont été analysées suivant la liste des critères suivants :",
        size=8,
        space_after=4,
    )
    _write_moa_grid_metadata(doc, ctrl, meta)
    rows = _control_grid_rows_from_xlsx(controle_xlsx) if controle_xlsx else []
    if not rows and ctrl and ctrl.grille:
        rows = [ctrl.grille.headers, *_controle_rows_for_moa(ctrl.grille)]
    if not rows:
        _add_moa_paragraph(doc, NOT_AVAILABLE, size=9)
        return
    _write_docx_table(
        doc,
        rows,
        widths=_GRILLE_COL_WIDTHS,
        font_size=6.2,
        header_rows=1,
        max_rows=85,
        repeat_header=True,
    )


def _write_moa_grid_metadata(doc: Document, ctrl, meta: AvpMeta) -> None:
    h = (ctrl.header if ctrl else {}) or {}
    left = [
        ["Projet", meta.project_name or _fmt_meta(h.get("projet"))],
        ["ESI", meta.project_code or _fmt_meta(h.get("esi"))],
        ["Phase", meta.phase or _fmt_meta(h.get("phase"))],
        ["Maquettes IFC transmises le", _fmt_meta(h.get("maquettes ifc transmises le"))],
        ["Date d'analyse", _fmt_meta(h.get("date d'analyse"))],
        ["Version d'analyse", _fmt_meta(h.get("version d'analyse"))],
    ]
    legend = (ctrl.legend if ctrl else {}) or {
        0: "Non fourni / non trouvé",
        1: "Insuffisant : à reprendre ou compléter",
        2: "Satisfaisant",
    }
    rows = []
    for idx in range(max(len(left), len(legend))):
        left_row = left[idx] if idx < len(left) else ["", ""]
        code = sorted(legend)[idx] if idx < len(legend) else ""
        rows.append(
            [left_row[0], left_row[1], "" if idx else "Légende :", code, legend.get(code, "")]
        )
    _write_docx_table(doc, rows, widths=(4.0, 4.0, 2.0, 1.2, 6.0), font_size=7.2)


def _write_moa_control_annex_pages(
    doc: Document,
    ctrl,
    meta: AvpMeta,
    controle_xlsx: Path | None,
    *,
    use_source_cached_values: bool = False,
) -> None:
    for sheet_name in _CONTROLE_STATS_SHEETS:
        doc.add_page_break()
        _write_moa_title(doc, f"onglet {sheet_name}", size=12)
        rows = _control_annex_rows_from_xlsx(controle_xlsx, sheet_name) if controle_xlsx else []
        grid = _controle_grid(sheet_name, ctrl)
        if use_source_cached_values and grid and grid.rows:
            rows = grid.rows
        elif rows:
            rows = _inject_annex_summary_values(rows, _stat_lookup(ctrl, sheet_name), sheet_name)
        if not rows:
            rows = grid.rows if grid else []
        if not rows:
            rows = _fallback_control_annex_rows(sheet_name, _stat_lookup(ctrl, sheet_name), meta)
        _write_docx_table(
            doc,
            rows,
            widths=_annex_widths(sheet_name),
            font_size=5.8,
            header_rows=0,
            max_rows=_annex_row_cap(sheet_name),
        )


def _control_grid_rows_from_xlsx(controle_xlsx: Path | None) -> list[list]:
    if not controle_xlsx or not Path(controle_xlsx).exists():
        return []
    wb = openpyxl.load_workbook(controle_xlsx, data_only=False, read_only=True)
    try:
        if "Grille de contrôle" not in wb.sheetnames:
            return []
        ws = wb["Grille de contrôle"]
        header_row = _find_control_header_row(ws)
        if header_row is None:
            return []
        rows: list[list] = []
        for row in ws.iter_rows(min_row=header_row, max_col=6, values_only=True):
            values = [_docx_excel_value(v) for v in row]
            if not any(values):
                continue
            rows.append(values)
        return rows
    finally:
        wb.close()


def _control_annex_rows_from_xlsx(controle_xlsx: Path | None, sheet_name: str) -> list[list]:
    if not controle_xlsx or not Path(controle_xlsx).exists():
        return []
    wb = openpyxl.load_workbook(controle_xlsx, data_only=True, read_only=True)
    try:
        if sheet_name not in wb.sheetnames:
            return []
        ws = wb[sheet_name]
        max_cols = min(ws.max_column, 11)
        max_rows = min(ws.max_row, _annex_row_cap(sheet_name))
        rows: list[list] = []
        for row in ws.iter_rows(min_row=1, max_row=max_rows, max_col=max_cols, values_only=True):
            values = [_docx_excel_value(v) for v in row]
            if any(values):
                rows.append(values)
        return rows
    finally:
        wb.close()


def _docx_excel_value(value) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    if isinstance(value, float):
        return f"{value:.2f}".rstrip("0").rstrip(".")
    if isinstance(value, str) and value.startswith("="):
        return ""
    text = str(value)
    if "openpyxl.worksheet.formula" in text:
        return ""
    return text


def _inject_annex_summary_values(
    rows: list[list], stats: dict | None, sheet_name: str
) -> list[list]:
    if not stats:
        return rows
    out = [list(row) for row in rows]
    for row in out:
        label = str(row[1] if len(row) > 1 else "").strip().lower()
        if "matériau" in sheet_name.lower() or "materiau" in _norm(sheet_name):
            if "sans" in label and "mat" in label:
                _set_row_value(row, 2, stats.get("non_conforme"))
                _set_row_value(row, 3, _pct(stats.get("non_conforme_ratio")))
                _set_row_value(row, 5, "Nombre d'élements :")
                _set_row_value(row, 6, stats.get("total"))
            elif "type" in label:
                _set_row_value(row, 2, _count_annex_types(out))
            continue
        if "noms" in label and "type" not in label:
            _set_row_value(row, 2, stats.get("total"))
            _set_row_value(row, 3, stats.get("conforme"))
            _set_row_value(row, 4, _pct(stats.get("conforme_ratio")))
            _set_row_value(row, 5, stats.get("non_conforme"))
            _set_row_value(row, 6, _pct(stats.get("non_conforme_ratio")))
        elif "type" in label:
            _set_row_value(row, 2, _count_annex_types(out))
    return out


def _set_row_value(row: list, idx: int, value) -> None:
    while len(row) <= idx:
        row.append("")
    if value not in (None, ""):
        row[idx] = str(value)


def _count_annex_types(rows: list[list]) -> int:
    values = {
        str(row[1]).strip()
        for row in rows
        if len(row) > 2
        and str(row[1]).strip()
        and str(row[2]).strip()
        and not str(row[1]).strip().lower().startswith(("name", "object", "ifc"))
    }
    return len(values)


def _fallback_control_annex_rows(sheet_name: str, stats: dict | None, meta: AvpMeta) -> list[list]:
    rows = [[meta.project_code], [meta.project_code], [f"onglet {sheet_name}"]]
    if not stats:
        rows.append([NOT_AVAILABLE])
        return rows
    if "matériau" in sheet_name.lower():
        rows.extend(
            [
                ["", "MN"],
                [
                    stats.get("label"),
                    stats.get("non_conforme"),
                    _pct(stats.get("non_conforme_ratio")),
                ],
                ["Total éléments", stats.get("total")],
            ]
        )
        return rows
    rows.extend(
        [
            ["", "", "MN", "", "Conforme", "", "Non Conforme"],
            [
                "",
                stats.get("label"),
                stats.get("total"),
                stats.get("conforme"),
                _pct(stats.get("conforme_ratio")),
                stats.get("non_conforme"),
                _pct(stats.get("non_conforme_ratio")),
            ],
        ]
    )
    return rows


def _annex_row_cap(sheet_name: str) -> int:
    if sheet_name == "Zones Nommage":
        return 48
    if sheet_name == "Pièces Nommage":
        return 62
    if "matériau" in sheet_name.lower():
        return 42
    return 58


def _annex_widths(sheet_name: str) -> tuple[float, ...]:
    if "matériau" in sheet_name.lower():
        return (2.4, 5.2, 1.4, 4.0, 1.0, 5.2, 1.8, 1.0)
    if sheet_name == "Pièces Nommage":
        return (2.0, 3.6, 1.1, 2.3, 1.0, 3.0, 1.0, 1.0, 3.0, 4.0, 1.0)
    return (2.0, 4.2, 1.1, 2.3, 1.0, 3.0, 1.0, 1.0, 4.2, 4.4)


def _write_docx_table(
    doc: Document,
    rows: list[list],
    *,
    widths: tuple[float, ...] | None = None,
    font_size: float = 8,
    header_rows: int = 0,
    max_rows: int | None = None,
    repeat_header: bool = False,
):
    if max_rows is not None:
        rows = rows[:max_rows]
    if not rows:
        return None
    ncols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        if repeat_header and row_idx == 0:
            _set_repeat_table_header(table.rows[0])
        for col_idx, cell in enumerate(cells):
            value = values[col_idx] if col_idx < len(values) else ""
            cell.text = "" if value is None else str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths and col_idx < len(widths):
                cell.width = Cm(widths[col_idx])
            if row_idx < header_rows or _looks_like_moa_section_row(values):
                _shade_cell(
                    cell, _MOA_TABLE_HEADER if row_idx < header_rows else _MOA_TABLE_SECTION
                )
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_docx_run_font(run, size=font_size, bold=row_idx < header_rows)
    return table


def _looks_like_moa_section_row(values: list) -> bool:
    non_empty = [v for v in values if v not in (None, "")]
    if len(non_empty) > 2:
        return False
    first = str(values[0] if values else "").strip()
    return first.isdigit() or first in {"Zones", "Pièces", "ObjT Pièces"}


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _fmt_meta(v) -> str:
    if v in (None, ""):
        return NOT_AVAILABLE
    if isinstance(v, datetime):
        return v.strftime("%Y-%m-%d")
    return str(v)


def _set_orientation(doc, orient) -> None:
    """Nouvelle section avec orientation portrait/paysage (lisibilité des
    tableaux larges en livrable client)."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(sec, orient)


# Largeurs (cm) des 6 colonnes de la grille de contrôle en paysage.
_GRILLE_COL_WIDTHS = (2.0, 6.5, 5.0, 3.0, 3.0, 7.0)


def _docx_header_table(doc, headers, *, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    tbl.allow_autofit = False
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = str(h)
        if col_widths and i < len(col_widths):
            cell.width = Cm(col_widths[i])
        _shade_cell(cell, BIMDATA_PRIMARY)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
    return tbl


def _write_computed_coverage(doc, snap) -> None:
    """Couverture des quantités **calculées** (fusion IfcOpenShell, Lot 3), si
    présente sur le snapshot — sinon rien. Ne masque aucun gap : affiche les
    valeurs conservées / ignorées / uuid inconnus."""
    cov = getattr(snap, "computed_coverage", None) if snap is not None else None
    if not cov:
        return
    p = doc.add_paragraph()
    p.add_run("Quantités calculées (IfcOpenShell) — couverture").bold = True
    doc.add_paragraph(
        "Valeurs géométriques calculées par IfcOpenShell et fusionnées en gap-only "
        "(les BaseQuantities natives BIMData ne sont jamais écrasées). Valeurs "
        "**non contractuelles** — un ré-export maquette avec BaseQuantities natives "
        "reste recommandé.",
        style="Intense Quote",
    )
    _kpi_table(
        doc,
        [
            ("Quantités fusionnées (comblées)", str(cov.get("n_merged", 0))),
            ("BaseQuantities natives conservées", str(cov.get("n_gap_kept", 0))),
            ("Entrées ignorées (skipped/failed)", str(cov.get("n_skipped_status", 0))),
            ("GlobalId inconnus du snapshot", str(cov.get("n_unknown_uuid", 0))),
        ],
    )


def _write_audit_synthese(doc, result) -> None:
    """Synthèse de l'audit BIMData réel (sévérité, thèmes, quantités
    manquantes) — le consolidé ne doit pas ignorer l'``AuditResult``."""
    if result is None:
        doc.add_paragraph("Audit automatisé : " + NOT_AVAILABLE + " (aucun audit chargé).")
        return
    by_sev = result.count_by_severity()
    by_theme = result.count_by_theme()
    by_type = result.count_by_error_type()
    p = doc.add_paragraph()
    p.add_run("Audit automatisé de la maquette active").bold = True
    # Répartition par sévérité COMPLÈTE (CRITICAL→INFO) pour que le total
    # se réconcilie côté client.
    _kpi_table(
        doc,
        [
            ("Anomalies détectées", str(len(result.findings))),
            ("Taux de conformité (pondéré)", f"{result.conformity_rate() * 100:.0f} %"),
            ("CRITICAL", str(by_sev.get("CRITICAL", 0))),
            ("HIGH", str(by_sev.get("HIGH", 0))),
            ("MEDIUM", str(by_sev.get("MEDIUM", 0))),
            ("LOW", str(by_sev.get("LOW", 0))),
            ("INFO", str(by_sev.get("INFO", 0))),
            ("Quantités manquantes", str(by_type.get("spatial_missing_quantity", 0))),
        ],
    )
    top = sorted(by_theme.items(), key=lambda kv: -kv[1])[:5]
    if top:
        doc.add_paragraph("Principaux thèmes en écart :")
        for theme, count in top:
            doc.add_paragraph(f"• {theme} : {count}", style="List Bullet")


def _write_donnees_entree(doc, ctrl, meta) -> None:
    h = (ctrl.header if ctrl else {}) or {}
    _kpi_table(
        doc,
        [
            ("Opération", _fmt_meta(meta.nombre_logements)),
            ("Maquettes IFC transmises le", _fmt_meta(h.get("maquettes ifc transmises le"))),
            ("Date d'analyse", _fmt_meta(h.get("date d'analyse"))),
            ("Version d'analyse", _fmt_meta(h.get("version d'analyse"))),
            ("Date du contrôle", _fmt_meta(meta.date_controle)),
            ("Auteur du contrôle", _fmt_meta(meta.auteur_controle)),
            ("Témoin virtuel", _fmt_meta(meta.temoin_virtuel)),
        ],
    )


def _write_grille_table(doc, ctrl, *, cap: int = 40) -> None:
    g = ctrl.grille if ctrl else None
    if not g or not g.headers:
        doc.add_paragraph(NOT_AVAILABLE)
        return
    # Largeurs adaptées uniquement pour la grille standard 6 colonnes.
    widths = _GRILLE_COL_WIDTHS if len(g.headers) == len(_GRILLE_COL_WIDTHS) else None
    tbl = _docx_header_table(doc, g.headers, col_widths=widths)
    for rowvals in g.rows[:cap]:
        cells = tbl.add_row().cells
        for i in range(len(cells)):
            v = rowvals[i] if i < len(rowvals) else None
            cells[i].text = "" if v in (None, "") else str(_cell(v))
            if widths and i < len(widths):
                cells[i].width = Cm(widths[i])
    if len(g.rows) > cap:
        doc.add_paragraph(
            f"… {len(g.rows) - cap} lignes supplémentaires dans le fichier "
            "« Contrôle Maquettes AVP ».",
            style="Intense Quote",
        )


def _write_stats_annex(doc, ctrl) -> None:
    stats = (ctrl.stats if ctrl else {}) or {}
    if not any(stats.values()):
        doc.add_paragraph(NOT_AVAILABLE)
        return
    for name, st in stats.items():
        if not st:
            continue
        p = doc.add_paragraph()
        p.add_run(name).bold = True
        if "conforme" in st:
            detail = (
                f"total {st.get('total')} · conformes {st.get('conforme')} "
                f"({_pct(st.get('conforme_ratio'))}) · non conformes {st.get('non_conforme')}"
            )
        else:
            detail = (
                f"total {st.get('total')} · sans matériau {st.get('non_conforme')} "
                f"({_pct(st.get('non_conforme_ratio'))})"
            )
        doc.add_paragraph(f"• {detail}", style="List Bullet")


def _write_ecarts(doc, result, sources, snap=None) -> None:
    env = sources.enveloppe if sources else None
    # SHAB snapshot avec le **même repli** que les annexes (BaseQuantities
    # puis « Superficie calculée ») — sinon l'écart reste NOT_AVAILABLE alors
    # que les annexes sont justes. Snapshot explicite prioritaire (audit non
    # encore lancé), sinon celui de l'``AuditResult``.
    eff_snap = snap if snap is not None else (result.snapshot if result is not None else None)
    ifc_shab = snapshot_shab_total(eff_snap)
    if ifc_shab is None and env is not None:
        ifc_shab = env.shab
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    for i, txt in enumerate(["Indicateur", "IFC OpenShell", "Méthode", "Observation"]):
        cell = tbl.rows[0].cells[i]
        cell.text = txt
        _shade_cell(cell, BIMDATA_PRIMARY)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True
    row = tbl.add_row().cells
    row[0].text = "SHAB totale (m²)"
    row[1].text = f"{ifc_shab:.2f}" if isinstance(ifc_shab, (int, float)) else NOT_AVAILABLE
    row[2].text = "Espaces IFC : BaseQuantities, puis Superficie calculée"
    row[3].text = "Valeur issue de la maquette ; aucun écart d'outil externe n'est intégré."
    doc.add_paragraph(
        "Les valeurs métier intégrées au pack sont extraites de la maquette "
        "ou calculées via la chaîne IFC OpenShell ; les colonnes d'outil "
        "externe ne sont pas utilisées comme source de données.",
        style="Intense Quote",
    )


def _points_bloquants(ctrl, env, ratio, seuil) -> list[str]:
    out: list[str] = []
    if isinstance(ratio, (int, float)) and isinstance(seuil, (int, float)) and ratio < seuil:
        out.append(
            f"Ratio FAC/SHAB {ratio:.3f} inférieur au Seuil 3F 2026 ({seuil}) — enveloppe à revoir."
        )
    # Points de contrôle évalués 0 (Non fourni / non trouvé) dans la grille.
    if ctrl and ctrl.grille:
        try:
            eval_idx = ctrl.grille.headers.index("EVALUATION")
            pts_idx = ctrl.grille.headers.index("POINTS DE CONTROLE")
        except ValueError:
            eval_idx = pts_idx = None
        if eval_idx is not None:
            zeros = [
                r[pts_idx] for r in ctrl.grille.rows if str(r[eval_idx]).strip() in ("0", "0.0")
            ]
            for pt in zeros[:8]:
                out.append(f"Point de contrôle non satisfait (éval. 0) : {pt}")
    return out


def _recommandations(pieces, zones, materiau, ratio, seuil) -> list[str]:
    recs: list[str] = []
    for label, stat in (("pièces", pieces), ("zones", zones)):
        nc = stat.get("non_conforme") if stat else None
        if isinstance(nc, (int, float)) and nc > 0:
            recs.append(
                f"Reprendre le nommage de {int(nc)} {label} non conformes (CCH BIM I3F chap. 6.3)."
            )
    if (
        materiau
        and isinstance(materiau.get("non_conforme"), (int, float))
        and materiau["non_conforme"] > 0
    ):
        recs.append(
            f"Compléter le matériau sur {int(materiau['non_conforme'])} éléments ARC sans matériau."
        )
    if isinstance(ratio, (int, float)) and isinstance(seuil, (int, float)) and ratio < seuil:
        recs.append(
            "Revoir la modélisation de l'enveloppe pour atteindre le ratio FAC/SHAB attendu."
        )
    if not recs:
        recs.append("Aucune action corrective majeure identifiée à partir des livrables fournis.")
    return recs


# ── Utilitaires ────────────────────────────────────────────────────────────


def _norm(s: str) -> str:
    return "".join(ch for ch in s.lower() if ch.isalnum())


def _safe_sheet(title: str) -> str:
    bad = set(r"[]:*?/\\")
    return "".join(c for c in title if c not in bad)[:31] or "Feuille"


def _stat_lookup(ctrl, name: str) -> dict:
    if not ctrl or not ctrl.stats:
        return {}
    for key, val in ctrl.stats.items():
        if _norm(key).replace("bsence", "absence") == _norm(name):
            return val or {}
    return {}


def _controle_from_audit_or_metadata(
    result: AuditResult | None, src: ControleMaquettesSource | None
) -> ControleMaquettesSource | None:
    """Construit la source Contrôle utilisée par les livrables.

    Les données métier mesurables (grille + stats) viennent de l'``AuditResult``.
    La source MOA éventuelle fournit l'entête, la légende et surtout le classeur
    template à conserver pour rester conforme au livrable maître d'ouvrage.
    """
    header = dict(src.header) if src and src.header else {}
    legend = dict(src.legend) if src and src.legend else {}
    template_path = getattr(src, "template_path", None)
    stat_grids = dict(src.stat_grids) if src and src.stat_grids else {}
    grille = _audit_controle_table(result)
    stats: dict[str, dict] = {}
    if result is not None:
        for name in _CONTROLE_STATS_SHEETS:
            value = _audit_stats(name, result)
            if value:
                stats[name] = value
    if not header and not legend and grille is None and not stats and not template_path:
        return None
    return ControleMaquettesSource(
        template_path=template_path,
        header=header,
        legend=legend,
        grille=grille,
        stats=stats,
        stat_grids=stat_grids,
    )


def _ifc_first_sources(
    *,
    model_sources: AvpSources,
    source_inputs: AvpSources | None,
    result: AuditResult | None,
) -> AvpSources:
    """Assemble les sources de génération en donnant priorité à la maquette.

    Les cinq exports métriques sont ceux calculés depuis le snapshot IFC. Les
    fichiers MOA fournis ne sont conservés que pour les métadonnées non
    métriques qui ne se déduisent pas de la géométrie (ex. seuil 3F).
    """
    ctrl_src = source_inputs.controle if source_inputs else None
    input_env = source_inputs.enveloppe if source_inputs else None
    # Exception importante : l'enveloppe peut venir du contrat structuré
    # ``envelope.json`` produit par ifc-geometry (par_type / hors_filtre_type).
    # Ce n'est pas un XLS MOA externe : c'est la source IFC/OpenShell métier et
    # elle doit primer sur le repli snapshot élémentaire.
    env = input_env if _is_envelope_json_source(input_env) else model_sources.enveloppe
    if env is not None and input_env is not None and not _is_envelope_json_source(input_env):
        # Le seuil est une règle/paramètre documentaire, pas une mesure externe.
        if input_env.seuil_3f is not None:
            env.seuil_3f = input_env.seuil_3f
    return AvpSources(
        controle=_controle_from_audit_or_metadata(result, ctrl_src),
        shab=model_sources.shab,
        zones_espaces=model_sources.zones_espaces,
        enveloppe=env,
        menuiseries=model_sources.menuiseries,
        plancher=model_sources.plancher,
    )


def _is_envelope_json_source(src) -> bool:
    """Vrai pour une source issue du contrat ``envelope.json`` ifc-geometry."""
    if src is None:
        return False
    if getattr(src, "superficie_calque_total", None) is not None:
        return True
    if getattr(src, "hors_filtre_type", None):
        return True
    table = getattr(src, "table", None)
    headers = getattr(table, "headers", None) or []
    return "Surface IFC OpenShell" in headers and "IFC OpenShell Surface des Fenêtres" in headers


# ── Orchestrateur ──────────────────────────────────────────────────────────


def write_avp_i3f_report_pack(
    result: AuditResult | None,
    output_dir: str | Path,
    *,
    sources: AvpSourcePaths | AvpSources | None = None,
    snapshot: ModelSnapshot | None = None,
    project_name: str = "Projet",
    project_code: str = "",
    phase: str = "AVP",
    auditor: str = "AMO BIM",
    date: str | None = None,
    usages_bim: list[str] | None = None,
    nombre_logements: str | None = None,
    temoin_virtuel: str | None = None,
    date_controle: str | None = None,
    auteur_controle: str | None = None,
    export_pdf: bool = True,
    context: ReportProjectContext | None = None,  # noqa: ARG001 (compat future)
) -> AvpReportPack:
    """Génère le pack de livrables AVP I3F dans ``output_dir``.

    Les noms de livrables suivent la convention documentaire I3F,
    **générés à partir des données projet confirmées** :
    ``YYMMDD <NomProjet> <CodeProjet> <Phase> - <TypeLivrable>.<ext>``.

    Args:
        result: ``AuditResult`` BIMData (peut être ``None`` si ``snapshot`` est
            fourni explicitement).
        output_dir: dossier de sortie (créé si besoin).
        sources: chemins des .xlsx I3F (``AvpSourcePaths``) ou sources déjà
            chargées (``AvpSources``). Quand un snapshot existe, ces sources
            ne fournissent pas les surfaces : les exports métriques sont
            recalculés depuis la maquette IFC.
        project_name, project_code, phase: identité projet **confirmée**
            injectée dans les noms de livrables (et les entêtes).
        date: préfixe daté ``YYMMDD`` des noms de livrables. ``None`` →
            date de génération (aujourd'hui).
        usages_bim, nombre_logements, temoin_virtuel, date_controle,
            auteur_controle: métadonnées opérationnelles du contrôle (issues
            du rapport I3F de référence) pour les sections « Données
            d'entrée » et « Usages BIM 3F ». Absentes → ``NOT_AVAILABLE``.
            Exception : ``auteur_controle`` non fourni est aligné sur
            ``auditor`` (le rédacteur AMO est aussi l'auteur du contrôle
            par défaut) plutôt que ``NOT_AVAILABLE``.
        export_pdf: tente la conversion .docx → .pdf (best-effort).
    """
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    # Date de génération du livrable (YYMMDD) si non imposée par l'appelant.
    gen_date = (
        date.strip()
        if isinstance(date, str) and date.strip()
        else datetime.now().strftime("%y%m%d")
    )
    # « Auteur du contrôle » : champ opérationnel distinct du rédacteur
    # AMO (``auditor``). Sur le pack I3F il est facultatif ; plutôt que
    # d'écrire ``NOT_AVAILABLE`` quand il n'est pas précisé, on l'aligne
    # sur ``auditor`` (donnée fournie par l'utilisateur, jamais inventée).
    # L'appelant peut toujours le distinguer en passant ``auteur_controle``.
    eff_auteur_controle = (
        auteur_controle if auteur_controle and auteur_controle.strip() else auditor
    )
    meta = AvpMeta(
        project_name=project_name,
        project_code=project_code,
        phase=phase,
        auditor=auditor,
        usages_bim=usages_bim,
        nombre_logements=nombre_logements,
        temoin_virtuel=temoin_virtuel,
        date_controle=date_controle,
        auteur_controle=eff_auteur_controle,
    )

    # Noms de livrables générés depuis l'identité projet confirmée
    # (convention I3F uniforme). On n'hérite plus du basename des sources :
    # le livrable BIMData porte l'identité et la date de génération.
    def _name(key: str) -> str:
        return _deliverable_filename(
            key,
            date=gen_date,
            project_name=project_name,
            project_code=project_code,
            phase=phase,
        )

    fn_controle = _name("controle")
    fn_shab = _name("shab")
    fn_zones = _name("zones_espaces")
    fn_env = _name("enveloppe")
    fn_men = _name("menuiseries")
    fn_plancher = _name("plancher")
    fn_analyse = _name("analyse")

    if isinstance(sources, AvpSourcePaths):
        sources = load_sources(sources)
    source_inputs = sources

    # ── Maquette-first / IFC OpenShell ──────────────────────────────────
    # Le snapshot est pris explicitement (``snapshot=``, ex. après
    # ``verify_active_model`` sans audit), sinon depuis ``result.snapshot``.
    # Dès qu'il existe, il devient la source autoritaire des exports métriques.
    snap = snapshot if snapshot is not None else (result.snapshot if result is not None else None)
    if snap is not None:
        sources = _ifc_first_sources(
            model_sources=build_sources_from_snapshot(snap),
            source_inputs=source_inputs,
            result=result,
        )
    elif sources is None:
        sources = AvpSources()

    controle = _build_controle_maquettes_xlsx(out / fn_controle, result, sources, meta, snap)
    shab = _build_multisheet_export_xlsx(
        out / fn_shab,
        "EXPORT SHAB MAQUETTE",
        "AVP - export SHAB maquette",
        (sources.shab if sources else None),
        meta,
    )
    zones = _build_multisheet_export_xlsx(
        out / fn_zones,
        "EXPORT ZONES ET ESPACES",
        "Export Zones et Espaces",
        (sources.zones_espaces if sources else None),
        meta,
    )
    enveloppe = _build_enveloppe_xlsx(out / fn_env, sources, meta)
    menuiseries = _build_menuiseries_xlsx(out / fn_men, sources, meta)
    plancher = _build_plancher_xlsx(out / fn_plancher, sources, meta)
    analyse = _build_analyse_bim_avp_docx(
        out / fn_analyse, result, sources, meta, snap, controle_xlsx=controle
    )

    pdf = docx_to_pdf(analyse) if export_pdf else None

    pack = AvpReportPack(
        controle_xlsx=controle,
        shab_xlsx=shab,
        zones_espaces_xlsx=zones,
        enveloppe_xlsx=enveloppe,
        menuiseries_xlsx=menuiseries,
        plancher_xlsx=plancher,
        analyse_docx=analyse,
        analyse_pdf=pdf,
    )

    # ── QA gate : anti-livrable vide ────────────────────────────────────
    # On rouvre chaque annexe et on compte les lignes métier. Échec si un
    # export sort sans ligne alors que la maquette contient des entités
    # exploitables (espaces / murs / zones). On lève : le tool renverra un
    # statut d'erreur explicite plutôt qu'un fichier vide.
    empty = _qa_empty_deliverables(pack, snap, result)
    if empty:
        raise AvpQaError(empty)

    return pack


def _multisheet_is_empty(multi) -> bool:
    """Un ``MultiSheetSource`` est vide si aucun onglet ne porte de ligne."""
    if multi is None:
        return True
    grids = getattr(multi, "grids", None) or []
    return not any(getattr(g, "rows", None) for g in grids)


def _tabular_is_empty(src) -> bool:
    """Une source tabulaire (enveloppe/menuiseries) est vide sans lignes."""
    if src is None:
        return True
    table = getattr(src, "table", None)
    return table is None or not getattr(table, "rows", None)


def _qa_empty_deliverables(
    pack: AvpReportPack, snap, result: AuditResult | None = None
) -> list[str]:
    """Liste des annexes vides alors que la maquette a des données."""
    if snap is None:
        return []
    problems: list[str] = []
    # 5ᵉ annexe : quand un audit est disponible, la « Grille de contrôle » doit
    # porter des points de contrôle réels (comptés SOUS son titre, hors entête/
    # légende/NOT_AVAILABLE). Vide malgré un audit = livrable non exploitable.
    if result is not None:
        expected_controle = _audit_controle_table(result) is not None or not result.findings
        if expected_controle and _count_controle_rows(pack.controle_xlsx) == 0:
            problems.append("Contrôle")
    has_spaces_or_zones = bool(getattr(snap, "spaces", None)) or bool(getattr(snap, "zones", None))
    if has_spaces_or_zones:
        if _count_business_rows(pack.shab_xlsx) == 0:
            problems.append("SHAB")
        if _count_business_rows(pack.zones_espaces_xlsx) == 0:
            problems.append("Zones/Espaces")
    if count_envelope_walls(snap) > 0 and _count_business_rows(pack.enveloppe_xlsx) == 0:
        problems.append("Enveloppe")
    if count_menuiseries(snap) > 0 and _count_business_rows(pack.menuiseries_xlsx) == 0:
        problems.append("Menuiseries")
    if count_planchers(snap) > 0 and _count_business_rows(pack.plancher_xlsx) == 0:
        problems.append("Plancher")
    return problems
